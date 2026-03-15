#!/usr/bin/env python3
"""
Hotel Insurance Proposal - DOCX Generator
Generates complete branded DOCX proposals following HUB International design system.
23-section document with full compliance pages.
"""

import os
import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# Constants and lookup tables extracted to separate module
from proposal_constants import (
    ELECTRIC_BLUE, CLASSIC_BLUE, ARCTIC_GRAY, EGGSHELL, WHITE, CHARCOAL,
    ELECTRIC_BLUE_HEX, CLASSIC_BLUE_HEX, ARCTIC_GRAY_HEX, EGGSHELL_HEX,
    LOGO_PATH, AM_BEST_RATINGS,
)

logger = logging.getLogger(__name__)


def _clean_carrier_name(name):
    """Strip (Non-Adm), (Non-Admitted), (Surplus Lines) etc. from carrier names."""
    if not name:
        return name
    import re
    name = re.sub(r'\s*\(Non-Adm(?:itted)?\)', '', name, flags=re.IGNORECASE).strip()
    name = re.sub(r'\s*\(Surplus Lines?\)', '', name, flags=re.IGNORECASE).strip()
    name = re.sub(r'\s*\(E&S\)', '', name, flags=re.IGNORECASE).strip()
    return name


def lookup_am_best(carrier_name):
    """Look up AM Best rating for a carrier. Returns rating or None."""
    if not carrier_name:
        return None
    name_lower = carrier_name.lower().strip()
    # Direct match
    if name_lower in AM_BEST_RATINGS:
        return AM_BEST_RATINGS[name_lower]
    # Partial match - check if any key is contained in the carrier name
    for key, rating in AM_BEST_RATINGS.items():
        if key in name_lower or name_lower in key:
            return rating
    return None

# Default Service Team
SERVICE_TEAM = [
    {
        "role": "Hotel Franchise Practice Leader",
        "name": "Stefan Burkey",
        "phone": "O: 407-636-8133\nM: 407-782-1900",
        "email": "stefan.burkey@hubinternational.com"
    },
    {
        "role": "Account Executive",
        "name": "Maureen Harvey, CIC, CRM",
        "phone": "O: 407-893-3830\nF: 407-831-3063",
        "email": "maureen.harvey@hubinternational.com"
    },
    {
        "role": "Senior Franchise Claims Advocate",
        "name": "Sheena Callazo, RPLU",
        "phone": "O: 630-468-5674",
        "email": "sheena.callazo@hubinternational.com"
    }
]

OFFICE_LOCATIONS = [
    "HUB International Midwest Limited — 203 N LaSalle, Suite 2000, Chicago, IL 60601",
    "HUB International Midwest Limited — 1411 Opus Place, Suite 450, Downers Grove, IL 60515"
]

# California Licenses
CA_LICENSES = [
    ("Agency Two Insurance Marketing Group, LLC, d/b/a AgencyOne", "0H44808"),
    ("All World Insurance Services, Inc.", "0F69702"),
    ("Avant Brokerage LLC", "0I77138"),
    ("Avant Specialty Claims, LLC", "6003211"),
    ("Avant Underwriters, LLC", "0G67877"),
    ("Brokers' Service Marketing Group II LLC", "0E02001"),
    ("Business Underwriters Associates Agency Inc.", "0C26183"),
    ("Chun-Ha Insurance Services, Inc.", "0F71901"),
    ("Dale Barton Agency", "0137389"),
    ("FNA Insurance Services Inc", "0I72746"),
    ("HUB Heartland, LLC", "0H15020"),
    ("HUB International Insurance Services Inc.", "0757776"),
    ("HUB International Iowa, LLC", "0K02887"),
    ("HUB International Mid-Atlantic Inc.", "0D58520"),
    ("HUB International Midwest Limited (MWW)", "0D08495"),
    ("HUB International Mountain States Limited", "0A96371"),
    ("HUB International New England, LLC", "0F79381"),
    ("HUB International Northeast Limited", "0E16962"),
    ("HUB International Northwest LLC", "0D08450"),
    ("HUB International Texas, Inc.", "0E24644"),
    ("HUB International Transportation Insurance Services Inc", "0D43442"),
    ("ISR Marine Insurance Services LLC", "0I67807"),
    ("Program Brokerage Corporation", "0B27851"),
    ("Sadler & Company, Inc.", "0B57651"),
    ("SBR Services, LLC", "6007384"),
    ("Silverstone Group, LLC", "0D79635"),
    ("Specialty Program Group LLC", "0L09546"),
    ("Squaremouth LLC", "0H58357"),
    ("The HDH Group, Inc.", "0C66771"),
    ("VIU by HUB", "6010202"),
]


# ─── Helper Functions ─────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with values like '1pt'."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="{ARCTIC_GRAY_HEX}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def remove_cell_borders(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def set_cell_width(cell, inches):
    """Set cell width in inches using XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(inches * 1440)}" w:type="dxa"/>')
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)


def set_cell_vertical_alignment(cell, align="center"):
    """Set vertical alignment of cell content."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
    existing = tcPr.find(qn('w:vAlign'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(vAlign)


def add_formatted_paragraph(doc, text, size=11, color=CLASSIC_BLUE, bold=False,
                            italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            space_before=0, space_after=0):
    """Add a formatted paragraph to the document."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Calibri"
    return p


def add_section_header(doc, text):
    """Add a 22pt Classic Blue bold section header with enough space to clear page header."""
    return add_formatted_paragraph(doc, text, size=22, color=CLASSIC_BLUE, bold=True,
                                   space_before=36, space_after=12)


def add_subsection_header(doc, text):
    """Add a 14pt Electric Blue bold subsection header."""
    return add_formatted_paragraph(doc, text, size=14, color=ELECTRIC_BLUE, bold=True,
                                   space_before=12, space_after=8)


def create_styled_table(doc, headers, rows, col_widths=None, header_size=10, body_size=10,
                        total_width=7.5, col_alignments=None, header_alignments=None):
    """Create a table with HUB styling: Electric Blue header, alternating rows.
    
    Args:
        doc: Document object
        headers: List of header strings
        rows: List of row data lists
        col_widths: List of column widths in inches. If None, auto-calculated.
        header_size: Font size for header row (default 10pt)
        body_size: Font size for body rows (default 10pt)
        total_width: Total table width in inches (default 7.5)
        col_alignments: Dict or list of WD_ALIGN_PARAGRAPH values per column for body rows. If None, all left-aligned.
        header_alignments: Dict or list of WD_ALIGN_PARAGRAPH values per column for header row.
                          If None, all center-aligned (default behavior).
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Auto-layout off for fixed widths
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    existing_layout = tblPr.find(qn('w:tblLayout'))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    tblPr.append(tblLayout)
    
    # Set total table width
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{int(total_width * 1440)}" w:type="dxa"/>')
    existing_tblW = tblPr.find(qn('w:tblW'))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblPr.append(tblW)
    
    # Calculate column widths if not provided
    if not col_widths:
        col_widths = [total_width / len(headers)] * len(headers)
    
    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        # Determine header alignment: use header_alignments if specified, else center
        h_align = WD_ALIGN_PARAGRAPH.CENTER
        if header_alignments is not None:
            if isinstance(header_alignments, dict):
                if i in header_alignments and header_alignments[i] is not None:
                    h_align = header_alignments[i]
            elif isinstance(header_alignments, (list, tuple)):
                if i < len(header_alignments) and header_alignments[i] is not None:
                    h_align = header_alignments[i]
        p.alignment = h_align
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(header_size + 2)
        run = p.add_run(header)
        run.font.size = Pt(header_size)
        run.font.color.rgb = WHITE
        run.font.bold = True
        run.font.name = "Calibri"
        set_cell_shading(cell, ELECTRIC_BLUE_HEX)
        set_cell_width(cell, col_widths[i] if i < len(col_widths) else 1.0)
        set_cell_vertical_alignment(cell, "center")
    
    # Style data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = Pt(body_size + 2)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(body_size)
            run.font.color.rgb = CLASSIC_BLUE
            run.font.name = "Calibri"
            set_cell_width(cell, col_widths[col_idx] if col_idx < len(col_widths) else 1.0)
            set_cell_vertical_alignment(cell, "center")
            # Apply column alignment if specified
            if col_alignments is not None:
                if isinstance(col_alignments, dict):
                    if col_idx in col_alignments and col_alignments[col_idx] is not None:
                        p.alignment = col_alignments[col_idx]
                elif isinstance(col_alignments, (list, tuple)):
                    if col_idx < len(col_alignments) and col_alignments[col_idx] is not None:
                        p.alignment = col_alignments[col_idx]
            # Alternating row colors
            if row_idx % 2 == 1:
                set_cell_shading(cell, EGGSHELL_HEX)
    
    return table


def add_page_header(doc):
    """Add page header with logo left and text right."""
    section = doc.sections[-1]
    header = section.header
    header.is_linked_to_previous = False
    
    # Create a table for the header layout
    htable = header.add_table(1, 2, width=Inches(7))
    htable.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Logo cell (left)
    logo_cell = htable.rows[0].cells[0]
    logo_cell.width = Inches(2.5)
    if os.path.exists(LOGO_PATH):
        p = logo_cell.paragraphs[0]
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.8))
    
    # Text cell (right) — intentionally left blank per branding preferences
    text_cell = htable.rows[0].cells[1]
    text_cell.width = Inches(4.5)
    p = text_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Hotel Franchise Practice")
    run.font.size = Pt(12)
    run.font.color.rgb = ELECTRIC_BLUE
    run.font.bold = True
    run.font.name = "Calibri"
    
    # Remove borders from header table
    for row in htable.rows:
        for cell in row.cells:
            remove_cell_borders(cell)
    
    # Add page footer with automatic page numbers
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(4)
    fp.paragraph_format.space_after = Pt(0)
    # "Page " prefix
    run_prefix = fp.add_run("Page ")
    run_prefix.font.size = Pt(8)
    run_prefix.font.color.rgb = ARCTIC_GRAY
    run_prefix.font.name = "Calibri"
    # Auto page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_num = fp.add_run()
    run_num._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_num2 = fp.add_run()
    run_num2._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_num3 = fp.add_run()
    run_num3._r.append(fldChar2)
    # Style the page number runs
    for r in [run_num, run_num2, run_num3]:
        r.font.size = Pt(8)
        r.font.color.rgb = ARCTIC_GRAY
        r.font.name = "Calibri"


def add_callout_box(doc, text, size=10):
    """Add an eggshell background callout/disclaimer box."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, EGGSHELL_HEX)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = CLASSIC_BLUE
    run.font.name = "Calibri"
    run.font.italic = True
    return table


def add_page_break(doc):
    """Add a page break with a spacer paragraph to push content below the header.
    Word suppresses space_before at the top of a new page, so we add an empty
    spacer paragraph with a small font size and fixed spacing to create clearance."""
    doc.add_page_break()
    # Add invisible spacer paragraph - Word won't suppress this
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(12)
    spacer.paragraph_format.line_spacing = Pt(2)
    run = spacer.add_run()
    run.font.size = Pt(2)


def fmt_currency(amount):
    """Format a number as currency, preserving cents if present."""
    if isinstance(amount, (int, float)):
        if amount == int(amount):
            return f"${int(amount):,}"
        return f"${amount:,.2f}"
    if isinstance(amount, str):
        if amount.startswith("$"):
            return amount
        try:
            val = float(amount.replace(',', ''))
            if val == int(val):
                return f"${int(val):,}"
            return f"${val:,.2f}"
        except (ValueError, AttributeError):
            return amount
    return str(amount)


def _parse_currency(s):
    """Parse a currency string like '$168,411' or '168411' into a float."""
    import re
    if isinstance(s, (int, float)):
        return float(s)
    if isinstance(s, str):
        cleaned = re.sub(r'[^\d.]', '', s.replace(',', ''))
        if cleaned:
            try:
                return float(cleaned)
            except ValueError:
                return 0
    return 0


# ─── Section Generators ───────────────────────────────────────

def generate_cover_page(doc, data):
    """Section 1: Cover Page - fits on a single page, no page header."""
    ci = data.get("client_info", {})
    client_name = ci.get("named_insured", "Client Name")
    dba = ci.get("dba", "")
    address = ci.get("address", "")
    effective_date = ci.get("effective_date", "")
    proposal_date = datetime.date.today().strftime("%B %d, %Y")
    
    # Ensure cover page section has NO header or footer
    cover_section = doc.sections[0]
    cover_header = cover_section.header
    cover_header.is_linked_to_previous = False
    # Clear any existing header content
    for p in cover_header.paragraphs:
        p.clear()
    # Clear footer on cover page
    cover_footer = cover_section.footer
    cover_footer.is_linked_to_previous = False
    for p in cover_footer.paragraphs:
        p.clear()
    
    # Logo centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(6)
    if os.path.exists(LOGO_PATH):
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.5))
    
    # Electric Blue line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>' 
        f'<w:bottom w:val="single" w:sz="36" w:space="1" w:color="{ELECTRIC_BLUE_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    
    # Title - single line
    add_formatted_paragraph(doc, "Commercial Insurance Proposal", size=32, color=CLASSIC_BLUE,
                           bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=16, space_after=36)
    
    # Prepared For box
    box_table = doc.add_table(rows=1, cols=1)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in box_table.rows:
        for cell in row.cells:
            cell.width = Inches(5.5)
    tbl = box_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="7920" w:type="dxa"/>')
    existing_tblW = tblPr.find(qn('w:tblW'))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblPr.append(tblW)
    cell = box_table.rows[0].cells[0]
    set_cell_shading(cell, EGGSHELL_HEX)
    
    # Top and bottom blue borders
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="24" w:space="0" w:color="{ELECTRIC_BLUE_HEX}"/>'
        f'<w:bottom w:val="single" w:sz="24" w:space="0" w:color="{ELECTRIC_BLUE_HEX}"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    # "Prepared For" label
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Prepared For")
    run.font.size = Pt(14)
    run.font.color.rgb = CHARCOAL
    run.font.name = "Calibri"
    
    # Client name
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(2)
    run2 = p2.add_run(client_name)
    run2.font.size = Pt(28)
    run2.font.color.rgb = ELECTRIC_BLUE
    run2.font.bold = True
    run2.font.name = "Calibri"
    
    # Address if present
    if address:
        p_addr = cell.add_paragraph()
        p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_addr.paragraph_format.space_before = Pt(2)
        p_addr.paragraph_format.space_after = Pt(8)
        run_addr = p_addr.add_run(address)
        run_addr.font.size = Pt(11)
        run_addr.font.color.rgb = CLASSIC_BLUE
        run_addr.font.name = "Calibri"
    
    # DBA if present
    if dba:
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(8)
        run3 = p3.add_run(f"DBA: {dba}")
        run3.font.size = Pt(14)
        run3.font.color.rgb = CLASSIC_BLUE
        run3.font.name = "Calibri"
    
    # Dates - two column table
    date_table = doc.add_table(rows=2, cols=2)
    date_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Labels row
    for i, label in enumerate(["Proposal Date", "Effective Date"]):
        dc = date_table.rows[0].cells[i]
        p = dc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        run.font.size = Pt(10)
        run.font.color.rgb = ARCTIC_GRAY
        run.font.bold = True
        run.font.name = "Calibri"
        remove_cell_borders(dc)
    
    # Values row
    for i, val in enumerate([proposal_date, effective_date]):
        dc = date_table.rows[1].cells[i]
        p = dc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(val)
        run.font.size = Pt(12)
        run.font.color.rgb = CLASSIC_BLUE
        run.font.bold = True
        run.font.name = "Calibri"
        remove_cell_borders(dc)
    
    # Gray line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="{ARCTIC_GRAY_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    
    # Presented By
    add_formatted_paragraph(doc, "Presented By", size=12, color=CLASSIC_BLUE,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=2)
    add_formatted_paragraph(doc, "HUB International Midwest Limited", size=13, color=CLASSIC_BLUE,
                           bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
    add_formatted_paragraph(doc, "Hotel Franchise Practice", size=11,
                           color=ELECTRIC_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)


def generate_service_team(doc, data):
    """Section 2: Service Team - starts a new section with page header."""
    # Add a new section (not just a page break) so the cover page stays header-free
    new_section = doc.add_section()
    new_section.top_margin = Inches(0.7)
    new_section.bottom_margin = Inches(0.6)
    new_section.left_margin = Inches(0.75)
    new_section.right_margin = Inches(0.75)
    add_page_header(doc)
    add_section_header(doc, "Your Service Team")
    
    headers = ["Role", "Name", "Phone", "Email"]
    rows = []
    for member in SERVICE_TEAM:
        rows.append([member["role"], member["name"], member["phone"], member["email"]])
    
    create_styled_table(doc, headers, rows, col_widths=[1.8, 1.6, 1.4, 2.7],
                       header_size=11, body_size=10)
    
    # Office locations
    add_formatted_paragraph(doc, "", space_before=12)
    for loc in OFFICE_LOCATIONS:
        add_formatted_paragraph(doc, loc, size=10, color=CLASSIC_BLUE, space_after=2)


def generate_premium_summary(doc, data):
    """Section 3: Premium Summary - Expiring vs Proposed"""
    add_page_break(doc)
    add_section_header(doc, "Premium Summary")
    
    add_subsection_header(doc, "Coverage Premium Comparison")
    add_formatted_paragraph(doc,
        "Premiums shown include applicable taxes and fees. TRIA/Terrorism premiums are not included.",
        size=9, color=CHARCOAL, space_after=6)
    
    coverages = data.get("coverages", {})
    expiring = data.get("expiring_premiums", {})
    expiring_details = data.get("expiring_details", {})
    
    logger.info(f"Premium Summary - coverages keys: {list(coverages.keys())}")
    logger.info(f"Premium Summary - expiring keys: {list(expiring.keys())} values: {expiring}")
    logger.info(f"Premium Summary - expiring_details keys: {list(expiring_details.keys())}")
    
    coverage_names = {
        "property": "Property",
        "property_alt_1": "Property (Option 2)",
        "property_alt_2": "Property (Option 3)",
        "excess_property": "Excess Property (Layer 1)",
        "excess_property_2": "Excess Property (Layer 2)",
        "general_liability": "General Liability",
        "general_liability_alt_1": "General Liability (Option 2)",
        "general_liability_alt_2": "General Liability (Option 3)",
        "umbrella": "Umbrella / Excess 1",
        "umbrella_alt_1": "Umbrella / Excess 2",
        "umbrella_alt_2": "Umbrella / Excess 3",
        "umbrella_alt_3": "Umbrella / Excess 4",
        "umbrella_layer_2": "2nd Excess Layer",
        "umbrella_layer_3": "3rd Excess Layer",
        "umbrella_layer_4": "4th Excess Layer",
        "excess_liability": "Excess Liability",
        "excess": "Excess Liability",
        "workers_comp": "Workers Compensation",
        "workers_compensation": "Workers Compensation",
        "workers_compensation_alt_1": "Workers Comp (Option 2)",
        "commercial_auto": "Commercial Auto",
        "flood": "Flood",
        "epli": "EPLI",
        "cyber": "Cyber",
        "cyber_alt_1": "Cyber (Option 2)",
        "terrorism": "Terrorism / TRIA",
        "crime": "Crime",
        "inland_marine": "Inland Marine",
        "equipment_breakdown": "Equipment Breakdown",
        "liquor_liability": "Liquor Liability",
        "innkeepers_liability": "Innkeepers Liability",
        "environmental": "Environmental / Pollution",
        "workplace_violence": "Workplace Violence",
        "garage_keepers": "Garage Keepers",
        "enviro_pack": "Enviro Pack",
        "wind_deductible_buydown": "Wind Deductible Buy Down",
        "earthquake": "Earthquake",
        "pollution": "Pollution Liability",
        "abuse_molestation": "Sexual Abuse & Molestation",
        "active_assailant": "Active Assailant",
        "deductible_buydown": "Deductible Buy Down",
    }
    
    # Determine if we have expiring data
    has_expiring = bool(expiring) or bool(expiring_details)

    # Collect all coverage keys — deduplicate workers_comp/workers_compensation
    all_keys = list(coverage_names.keys())
    _has_wc = "workers_comp" in coverages
    _has_wc_long = "workers_compensation" in coverages
    if _has_wc and _has_wc_long:
        all_keys = [k for k in all_keys if k != "workers_compensation"]  # prefer short key
    elif _has_wc_long and not _has_wc:
        all_keys = [k for k in all_keys if k != "workers_comp"]  # keep only long key

    # Separate optional coverages from main coverages
    optional_rows = []
    rows = []
    total_proposed = 0
    total_expiring = 0

    # Expiring key mapping: coverage key -> expiring key (some may differ)
    _expiring_key_map = {
        "general_liability": ["general_liability", "gl"],
        "workers_comp": ["workers_comp", "workers_compensation"],
        "workers_compensation": ["workers_comp", "workers_compensation"],
        "commercial_auto": ["commercial_auto", "auto"],
        "equipment_breakdown": ["equipment_breakdown", "boiler_machinery"],
    }

    def _get_expiring_premium(key):
        """Look up expiring premium for a coverage key."""
        # First check expiring_details (has per-coverage structured data)
        if expiring_details:
            detail = expiring_details.get(key)
            if isinstance(detail, dict):
                ep = detail.get("premium") or detail.get("total_premium") or 0
                if ep:
                    return _parse_currency(ep) if isinstance(ep, str) else float(ep)
            # Try alternate key names
            for alt_key in _expiring_key_map.get(key, []):
                detail = expiring_details.get(alt_key)
                if isinstance(detail, dict):
                    ep = detail.get("premium") or detail.get("total_premium") or 0
                    if ep:
                        return _parse_currency(ep) if isinstance(ep, str) else float(ep)
        # Fallback: simple expiring dict
        if expiring:
            val = expiring.get(key)
            if val:
                return _parse_currency(val) if isinstance(val, str) else float(val)
            for alt_key in _expiring_key_map.get(key, []):
                val = expiring.get(alt_key)
                if val:
                    return _parse_currency(val) if isinstance(val, str) else float(val)
        return 0

    for key in all_keys:
        display_name = coverage_names[key]
        cov = coverages.get(key)
        is_optional = cov.get("optional", False) if isinstance(cov, dict) else False

        # Show each coverage line individually
        if not cov:
            continue
        carrier = _clean_carrier_name(cov.get("carrier", ""))
        carrier_short = carrier
        if len(carrier) > 30:
            carrier_short = carrier.replace("Insurance Company", "Ins Co").replace("Specialty ", "Spec ")
        premium = cov.get("premium", 0) or 0
        total_prem = cov.get("total_premium", 0) or 0
        # Use total_premium as the displayed proposed premium (includes taxes/fees)
        proposed = total_prem if total_prem else premium

        if has_expiring:
            exp_prem = _get_expiring_premium(key)
            dollar_change = proposed - exp_prem if (proposed and exp_prem) else 0
            pct_change = ((proposed - exp_prem) / exp_prem * 100) if exp_prem else 0

            row_data = [
                display_name,
                carrier_short,
                fmt_currency(exp_prem) if exp_prem else "—",
                fmt_currency(proposed) if proposed else "—",
                fmt_currency(dollar_change) if (proposed and exp_prem) else "—",
                f"{pct_change:+.1f}%" if (proposed and exp_prem) else "—",
            ]
        else:
            row_data = [
                display_name,
                carrier_short,
                fmt_currency(proposed) if proposed else "—",
            ]

        if is_optional:
            optional_rows.append(row_data)
        else:
            rows.append(row_data)
            # Exclude terrorism/TRIA from totals
            if key != "terrorism":
                total_proposed += proposed
                if has_expiring:
                    total_expiring += _get_expiring_premium(key)

    # Total row
    if has_expiring:
        total_dollar = total_proposed - total_expiring if (total_proposed and total_expiring) else 0
        total_pct = ((total_proposed - total_expiring) / total_expiring * 100) if total_expiring else 0
        rows.append([
            "TOTAL",
            "",
            fmt_currency(total_expiring) if total_expiring else "—",
            fmt_currency(total_proposed) if total_proposed else "—",
            fmt_currency(total_dollar) if (total_proposed and total_expiring) else "—",
            f"{total_pct:+.1f}%" if (total_proposed and total_expiring) else "—",
        ])
        headers = ["Coverage", "Carrier", "Expiring\nPremium", "Proposed\nPremium", "$ Change", "% Change"]
        col_widths = [1.4, 2.0, 1.0, 1.0, 0.9, 0.8]
        col_alignments = [None, None, WD_ALIGN_PARAGRAPH.RIGHT,
                          WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
                          WD_ALIGN_PARAGRAPH.RIGHT]
    else:
        rows.append([
            "TOTAL",
            "",
            fmt_currency(total_proposed) if total_proposed else "—",
        ])
        headers = ["Coverage", "Carrier", "Proposed Premium"]
        col_widths = [2.0, 2.5, 1.5]
        col_alignments = [None, None, WD_ALIGN_PARAGRAPH.RIGHT]

    table = create_styled_table(doc, headers, rows,
                               col_widths=col_widths,
                               header_size=10, body_size=10,
                               col_alignments=col_alignments)
    
    # Bold and shade the total row
    last_row = table.rows[-1]
    for col_idx, cell in enumerate(last_row.cells):
        set_cell_shading(cell, ELECTRIC_BLUE_HEX)
        for p in cell.paragraphs:
            if col_idx >= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
    
    # Optional coverages section below TOTAL
    if optional_rows:
        add_formatted_paragraph(doc, "", space_before=12)
        add_subsection_header(doc, "Optional Coverages")
        add_formatted_paragraph(doc,
            "The following coverages are presented for consideration and are not included in the total premium above.",
            size=9, color=CHARCOAL, space_after=6)
        
        # Optional coverages always show only 3 columns: Coverage, Carrier, Proposed Premium
        # No expiring, $ change, or % change since these are new coverage recommendations
        opt_simple_rows = []
        for orow in optional_rows:
            cov_name = orow[0]
            carrier_name = orow[1]
            # Proposed premium is at index 3 (with expiring) or index 2 (without)
            if has_expiring:
                proposed_val = orow[3] if len(orow) > 3 else "N/A"
            else:
                proposed_val = orow[2] if len(orow) > 2 else "N/A"
            opt_simple_rows.append([cov_name, carrier_name, proposed_val])
        
        opt_headers = ["Coverage", "Carrier", "Proposed Premium"]
        opt_table = create_styled_table(doc, opt_headers, opt_simple_rows,
                                      col_widths=[2.0, 2.5, 1.5],
                                      header_size=10, body_size=10,
                                      col_alignments=[None, None, WD_ALIGN_PARAGRAPH.RIGHT])
    
    add_formatted_paragraph(doc, "", space_before=6)
    add_callout_box(doc,
        "This comparison is for reference only. Actual coverage terms, conditions, and exclusions "
        "are governed by the policies as issued. Please review all policies carefully upon receipt.")


def generate_payment_options(doc, data):
    """Section 4: Payment Options"""
    add_page_break(doc)
    add_section_header(doc, "Payment Options")
    
    payment_opts = data.get("payment_options", [])
    if payment_opts:
        # Filter out commission-related entries and clean commission text from terms
        import re
        filtered_opts = []
        for po in payment_opts:
            terms = po.get("terms", "")
            carrier = po.get("carrier", "")
            # Skip entries that are purely about commission
            if carrier.lower().strip() in ("commission", "broker fee", "broker"):
                continue
            # Remove commission-related sentences from terms text
            terms = re.sub(r'[^.]*commission[^.]*\.?', '', terms, flags=re.IGNORECASE).strip()
            terms = re.sub(r'[^.]*broker fee[^.]*\.?', '', terms, flags=re.IGNORECASE).strip()
            if terms or po.get("mep"):
                filtered_opts.append({"carrier": carrier, "coverage_type": po.get("coverage_type", ""), "terms": terms, "mep": po.get("mep", "")})
        
        if filtered_opts:
            headers = ["Carrier — Policy Type", "Payment Terms", "Min. Earned Premium"]
            rows = []
            for po in filtered_opts:
                carrier_name = po.get("carrier", "")
                cov_type = po.get("coverage_type", "")
                # Append coverage type after carrier name (e.g., "Kinsale — Property")
                if cov_type:
                    carrier_display = f"{carrier_name} — {cov_type}"
                else:
                    carrier_display = carrier_name
                rows.append([carrier_display, po.get("terms", ""), po.get("mep", "")])
            create_styled_table(doc, headers, rows, col_widths=[2.8, 2.7, 2.0],
                               header_size=10, body_size=9,
                               col_alignments={2: WD_ALIGN_PARAGRAPH.CENTER})
        else:
            add_formatted_paragraph(doc, "Payment terms to be confirmed upon binding.", size=11)
    else:
        add_formatted_paragraph(doc, "Payment terms to be confirmed upon binding.", size=11)
    
    # Earned premium / cancellation disclaimer - small font, bold, red
    _add_earned_premium_disclaimer(doc)


def generate_subjectivities(doc, data):
    """Section 5: Binding Subjectivities"""
    add_page_break(doc)
    add_section_header(doc, "Binding Subjectivities")
    
    add_formatted_paragraph(doc, "The following items are required prior to or as a condition of binding:",
                           size=11, space_after=8)
    
    coverages = data.get("coverages", {})
    coverage_names = {
        "property": "Property",
        "property_alt_1": "Property (Option 2)",
        "property_alt_2": "Property (Option 3)",
        "excess_property": "Excess Property (Layer 1)",
        "excess_property_2": "Excess Property (Layer 2)",
        "general_liability": "General Liability",
        "general_liability_alt_1": "General Liability (Option 2)",
        "general_liability_alt_2": "General Liability (Option 3)",
        "umbrella": "Umbrella / Excess 1",
        "umbrella_alt_1": "Umbrella / Excess 2",
        "umbrella_alt_2": "Umbrella / Excess 3",
        "umbrella_alt_3": "Umbrella / Excess 4",
        "umbrella_layer_2": "2nd Excess Layer",
        "umbrella_layer_3": "3rd Excess Layer",
        "umbrella_layer_4": "4th Excess Layer",
        "excess_liability": "Excess Liability",
        "excess": "Excess Liability",
        "workers_comp": "Workers Compensation",
        "workers_compensation": "Workers Compensation",
        "workers_compensation_alt_1": "Workers Comp (Option 2)",
        "commercial_auto": "Commercial Auto",
        "terrorism": "Terrorism / TRIA",
        "cyber": "Cyber Liability",
        "cyber_alt_1": "Cyber (Option 2)",
        "epli": "Employment Practices Liability",
        "crime": "Crime",
        "flood": "Flood",
        "inland_marine": "Inland Marine",
        "equipment_breakdown": "Equipment Breakdown",
        "liquor_liability": "Liquor Liability",
        "innkeepers_liability": "Innkeepers Liability",
        "environmental": "Environmental / Pollution",
        "workplace_violence": "Workplace Violence",
        "garage_keepers": "Garage Keepers",
        "enviro_pack": "Enviro Pack",
        "wind_deductible_buydown": "Wind Deductible Buy Down",
        "earthquake": "Earthquake",
        "pollution": "Pollution Liability",
        "abuse_molestation": "Sexual Abuse & Molestation",
        "active_assailant": "Active Assailant",
        "deductible_buydown": "Deductible Buy Down",
    }
    
    has_subjectivities = False
    for key, display_name in coverage_names.items():
        cov = coverages.get(key)
        if cov and cov.get("subjectivities"):
            has_subjectivities = True
            # Add carrier info with the coverage name
            carrier = _clean_carrier_name(cov.get("carrier", ""))
            header_text = f"{display_name} — {carrier}" if carrier else display_name
            add_subsection_header(doc, header_text)
            for subj in cov["subjectivities"]:
                add_formatted_paragraph(doc, f"☐  {subj}", size=10, space_after=3)
    
    if not has_subjectivities:
        add_formatted_paragraph(doc, "No subjectivities noted. Please confirm with carrier.", size=11)


def _proper_case(name):
    """Convert a name to proper title case, handling special cases.
    ALL CAPS and all lowercase get converted; mixed case is preserved."""
    if not name or not name.strip():
        return name
    s = name.strip()
    # If ALL CAPS or all lowercase, convert to title case
    if s.isupper() or s.islower():
        s = s.title()
    # Fix common abbreviations that should stay uppercase
    for abbr in ["LLC", "LP", "LLP", "INC", "DBA", "II", "III", "IV",
                 "NW", "NE", "SW", "SE", "US", "CT", "NJ", "PA", "NY",
                 "FL", "TX", "CA", "VA", "MD", "GA", "NC", "SC", "OH"]:
        # Use word boundary replacement to avoid partial matches
        import re
        s = re.sub(r'\b' + abbr.title() + r'\b', abbr, s)
    return s


def generate_named_insureds(doc, data):
    """Section 6: Named Insureds"""
    add_page_break(doc)
    add_section_header(doc, "Named Insureds")
    
    # Deduplicate named insureds case-insensitively and apply proper case
    raw_named = data.get("named_insureds", [])
    seen = set()
    named = []
    
    # Hotel brand names that should NOT appear in named insured entries
    # (GPT sometimes concatenates brand names from quotes into named insured DBA fields)
    _brand_keywords = {"marriott", "hilton", "ihg", "wyndham", "best western", "choice",
                       "hampton inn", "hampton", "holiday inn", "holiday inn express",
                       "candlewood", "towneplace", "staybridge", "springhill",
                       "comfort inn", "comfort suites", "quality inn", "sleep inn",
                       "la quinta", "days inn", "super 8", "ramada", "baymont",
                       "microtel", "wingate", "hawthorn", "home2", "tru by hilton",
                       "embassy suites", "doubletree", "hyatt", "radisson", "crowne plaza"}
    
    def _sanitize_named_insured(name_str):
        """Remove hotel brand names that GPT may have concatenated into named insured."""
        if not name_str:
            return name_str
        # If the name contains multiple brand keywords, it's likely a GPT hallucination
        name_lower = name_str.lower()
        brand_count = sum(1 for b in _brand_keywords if b in name_lower)
        if brand_count >= 2:
            # Strip everything after the first DBA + entity name
            import re
            # Try to find "LLC DBA <brand>" and keep just the LLC part
            m = re.match(r'^(.+?\b(?:LLC|LP|LLP|Inc|Corp)\b)\s*(?:DBA\s+.+)?$', name_str, re.IGNORECASE)
            if m:
                return m.group(1).strip()
        return name_str
    
    for ni in raw_named:
        if isinstance(ni, dict):
            ni_name = ni.get("name", "")
            ni_dba = ni.get("dba", "")
        else:
            ni_name = str(ni)
            ni_dba = ""
        # Sanitize: remove hallucinated brand concatenations
        ni_name = _sanitize_named_insured(ni_name)
        if ni_dba:
            ni_dba = _sanitize_named_insured(ni_dba)
        key = ni_name.strip().upper()
        if key and key not in seen:
            seen.add(key)
            display = _proper_case(ni_name)
            if ni_dba:
                display += f" DBA {_proper_case(ni_dba)}"
            named.append(display)
    
    # Also check if DBA is available from client_info and append to first named insured
    ci = data.get("client_info", {})
    ci_dba = ci.get("dba", "")
    if named and ci_dba and "DBA" not in named[0].upper():
        named[0] = f"{named[0]} DBA {_proper_case(ci_dba)}"
    
    if named:
        headers = ["#", "Named Insured"]
        rows = [[str(i), ni] for i, ni in enumerate(named, 1)]
        create_styled_table(doc, headers, rows, col_widths=[0.5, 7.0],
                           header_size=10, body_size=10)
    else:
        ni_text = _proper_case(ci.get("named_insured", "TBD"))
        if ci_dba:
            ni_text += f" DBA {_proper_case(ci_dba)}"
        headers = ["#", "Named Insured"]
        rows = [["1", ni_text]]
        create_styled_table(doc, headers, rows, col_widths=[0.5, 7.0],
                           header_size=10, body_size=10)
    
    # Additional Named Insureds - only show if there are entries NOT already in named_insureds
    # Since the web pipeline merges additional_named_insureds into named_insureds,
    # this section is typically empty. Only render if there are truly new entries.
    addl_named = data.get("additional_named_insureds", [])
    if addl_named:
        # Check if any additional named insureds are NOT already in the named list
        named_upper = {n.strip().upper() for n in named}
        new_addl = []
        for ani in addl_named:
            if isinstance(ani, dict):
                name = ani.get("name", "")
                dba = ani.get("dba", "")
                display = f"{name} DBA {dba}" if dba else name
            else:
                display = str(ani)
            # Check if this entity is already in the named insureds table
            display_upper = display.strip().upper()
            # Also check partial matches (name without DBA)
            name_only = (ani.get("name", "") if isinstance(ani, dict) else str(ani)).strip().upper()
            already_present = False
            for existing in named_upper:
                if name_only and (name_only in existing or existing in name_only):
                    already_present = True
                    break
                if display_upper and (display_upper in existing or existing in display_upper):
                    already_present = True
                    break
            if not already_present:
                new_addl.append(display)
        if new_addl:
            add_subsection_header(doc, "Additional Named Insureds")
            headers = ["#", "Additional Named Insured"]
            rows = [[str(i), d] for i, d in enumerate(new_addl, 1)]
            create_styled_table(doc, headers, rows, col_widths=[0.5, 7.0],
                               header_size=10, body_size=10)
    
    # Additional Insureds
    addl_insureds = data.get("additional_insureds", [])
    if addl_insureds:
        add_subsection_header(doc, "Additional Insureds")
        headers = ["#", "Additional Insured", "Relationship"]
        rows = []
        for i, ai in enumerate(addl_insureds, 1):
            if isinstance(ai, dict):
                rows.append([str(i), ai.get("name", ""), ai.get("relationship", "")])
            else:
                rows.append([str(i), str(ai), ""])
        create_styled_table(doc, headers, rows, col_widths=[0.5, 4.5, 2.5],
                           header_size=10, body_size=10)
    
    # Scan ALL coverages for additional named insureds not already shown
    # This catches per-coverage ANI lists (GL, Crime, Umbrella, etc.)
    coverages = data.get("coverages", {})
    _all_shown_names = set()
    for n in named:
        _all_shown_names.add(n.strip().upper())
    for n in (new_addl if 'new_addl' in dir() else []):
        _all_shown_names.add(n.strip().upper())
    for ai in addl_insureds:
        if isinstance(ai, dict):
            _all_shown_names.add((ai.get("name", "") or "").strip().upper())
        else:
            _all_shown_names.add(str(ai).strip().upper())

    # Supplement named insureds from SOV corporate names (may have more entities)
    sov_data_ni = data.get("sov_data")
    if sov_data_ni and sov_data_ni.get("locations"):
        for loc in sov_data_ni["locations"]:
            corp = (loc.get("corporate_name", "") or "").strip()
            if not corp:
                continue
            corp_upper = corp.strip().upper()
            already = False
            for existing_key in seen:
                if corp_upper in existing_key or existing_key in corp_upper:
                    already = True
                    break
            if not already:
                seen.add(corp_upper)
                display = _proper_case(corp)
                dba = (loc.get("dba", "") or loc.get("hotel_flag", "") or "").strip()
                if dba:
                    display += f" DBA {_proper_case(dba)}"
                named.append(display)


    _coverage_ani_rows = []
    _coverage_display_names = {
        "general_liability": "General Liability",
        "crime": "Crime",
        "umbrella": "Umbrella / Excess",
        "workers_comp": "Workers Compensation",
        "workers_compensation": "Workers Compensation",
        "commercial_auto": "Commercial Auto",
        "cyber": "Cyber",
        "epli": "EPLI",
    }
    sov_locs = (data.get("sov_data") or {}).get("locations", []) if data.get("sov_data") else []

    for cov_key, cov_data in coverages.items():
        if not isinstance(cov_data, dict):
            continue
        cov_ani = cov_data.get("additional_named_insureds", []) or []
        if not cov_ani:
            cov_ani = cov_data.get("additional_insureds", []) or []
        if not cov_ani:
            cov_ani = cov_data.get("named_insureds", []) or []
        for ani in cov_ani:
            if isinstance(ani, dict):
                ani_name = ani.get("name", "") or ani.get("insured", "") or ""
                ani_loc = ani.get("location", "") or ani.get("hotel", "") or ""
            else:
                ani_name = str(ani)
                ani_loc = ""
            if not ani_name.strip():
                continue
            if ani_name.strip().upper() in _all_shown_names:
                continue
            _all_shown_names.add(ani_name.strip().upper())
            if not ani_loc and sov_locs:
                ani_lower = ani_name.lower()
                for sov_loc in sov_locs:
                    dba = (sov_loc.get("dba", "") or sov_loc.get("hotel_flag", "") or "").lower()
                    corp = (sov_loc.get("corporate_name", "") or "").lower()
                    if dba and dba in ani_lower:
                        ani_loc = f"{sov_loc.get('address', '')}, {sov_loc.get('city', '')}, {sov_loc.get('state', '')}"
                        break
                    elif corp and corp in ani_lower:
                        ani_loc = f"{sov_loc.get('address', '')}, {sov_loc.get('city', '')}, {sov_loc.get('state', '')}"
                        break
            cov_display = _coverage_display_names.get(cov_key, cov_key.replace("_", " ").title())
            _coverage_ani_rows.append([ani_name, ani_loc, cov_display])

    if _coverage_ani_rows:
        add_subsection_header(doc, "Additional Named Insureds by Coverage")
        headers = ["#", "Named Insured", "Location / Hotel", "Coverage"]
        rows = [[str(i), r[0], r[1], r[2]] for i, r in enumerate(_coverage_ani_rows, 1)]
        create_styled_table(doc, headers, rows, col_widths=[0.4, 2.8, 2.8, 1.5],
                           header_size=9, body_size=8)

    # Note box
    add_formatted_paragraph(doc, "", space_before=8)
    add_callout_box(doc, "Note: Additional named insureds may be added as required by franchise agreements or management contracts.")

    # Additional Interests
    interests = data.get("additional_interests", [])
    if interests:
        add_subsection_header(doc, "Additional Interests")
        headers = ["Type", "Name & Address", "Description"]
        rows = [[ai.get("type", ""), ai.get("name_address", ""), ai.get("description", "")] for ai in interests]
        create_styled_table(doc, headers, rows, col_widths=[1.5, 3.5, 2.5])


def generate_information_summary(doc, data):
    """Section 7: Information Summary"""
    add_page_break(doc)
    add_section_header(doc, "Information Summary")
    
    ci = data.get("client_info", {})
    
    headers = ["Item", "Details"]
    rows = [
        ["First Named Insured", ci.get("named_insured", "")],
        ["Mailing Address", ci.get("address", "")],
        ["Entity Type", ci.get("entity_type", "")],
        ["Effective Date", ci.get("effective_date", "")],
        ["Expiration Date", ci.get("expiration_date", "")],
    ]
    if ci.get("sales_exposure_basis"):
        # Recalculate proposed sales from sum of ALL GL exposures
        gl_cov = data.get("coverages", {}).get("general_liability", {})
        _gl_cls = gl_cov.get("schedule_of_classes", []) if isinstance(gl_cov, dict) else []
        _recalc = 0
        import re as _re_s
        for _e in _gl_cls:
            if isinstance(_e, dict):
                _exp = _e.get("exposure", "")
                if isinstance(_exp, (int, float)):
                    _recalc += _exp
                elif isinstance(_exp, str):
                    _c = _re_s.sub(r'[^\d.]', '', _exp.replace(',', ''))
                    if _c:
                        try: _recalc += float(_c)
                        except ValueError: pass
        if _recalc > 0:
            _basis = ci["sales_exposure_basis"]
            _upd = _re_s.sub(r'\$[\d,]+', fmt_currency(_recalc), _basis, count=1)
            if _upd == _basis:
                _upd = fmt_currency(_recalc) + " \u2013 " + _basis
            rows.append(["Proposed Sales/Exposure Basis", _upd])
        else:
            rows.append(["Proposed Sales/Exposure Basis", ci["sales_exposure_basis"]])
    if ci.get("dba"):
        rows.insert(1, ["DBA", _proper_case(ci["dba"])])
    
    # Calculate total sales from GL schedule_of_classes exposures
    coverages = data.get("coverages", {})
    gl_cov = coverages.get("general_liability", {})
    designated_premises = gl_cov.get("designated_premises", []) if isinstance(gl_cov, dict) else []
    if isinstance(gl_cov, dict):
        gl_classes = gl_cov.get("schedule_of_classes", [])
        total_sales = 0
        import re as _re
        for entry in gl_classes:
            if isinstance(entry, dict):
                exposure = entry.get("exposure", "")
                if isinstance(exposure, (int, float)):
                    total_sales += exposure
                elif isinstance(exposure, str):
                    # Parse dollar amounts like "$8,748,612" or "8748612"
                    cleaned = _re.sub(r'[^\d.]', '', exposure.replace(',', ''))
                    if cleaned:
                        try:
                            total_sales += float(cleaned)
                        except ValueError:
                            pass
        if total_sales > 0:
            rows.append(["Total Sales / Exposure", fmt_currency(total_sales)])
        # Also add total_sales from GL coverage if extracted
        elif gl_cov.get("total_sales"):
            rows.append(["Total Sales / Exposure", gl_cov["total_sales"]])
    
    # Add number of locations with property/liability breakdown
    # Count UNIQUE addresses using composite key (addr|city|state) — 1 address with 4 buildings = 1 location
    sov_data = data.get("sov_data")
    sov_locs = sov_data.get("locations", []) if sov_data else []
    
    # Helper to build composite key for dedup
    # State name -> abbreviation map for _loc_key normalization
    _state_abbrevs = {
        "ALABAMA": "AL", "ALASKA": "AK", "ARIZONA": "AZ", "ARKANSAS": "AR",
        "CALIFORNIA": "CA", "COLORADO": "CO", "CONNECTICUT": "CT", "DELAWARE": "DE",
        "FLORIDA": "FL", "GEORGIA": "GA", "HAWAII": "HI", "IDAHO": "ID",
        "ILLINOIS": "IL", "INDIANA": "IN", "IOWA": "IA", "KANSAS": "KS",
        "KENTUCKY": "KY", "LOUISIANA": "LA", "MAINE": "ME", "MARYLAND": "MD",
        "MASSACHUSETTS": "MA", "MICHIGAN": "MI", "MINNESOTA": "MN", "MISSISSIPPI": "MS",
        "MISSOURI": "MO", "MONTANA": "MT", "NEBRASKA": "NE", "NEVADA": "NV",
        "NEW HAMPSHIRE": "NH", "NEW JERSEY": "NJ", "NEW MEXICO": "NM", "NEW YORK": "NY",
        "NORTH CAROLINA": "NC", "NORTH DAKOTA": "ND", "OHIO": "OH", "OKLAHOMA": "OK",
        "OREGON": "OR", "PENNSYLVANIA": "PA", "RHODE ISLAND": "RI",
        "SOUTH CAROLINA": "SC", "SOUTH DAKOTA": "SD", "TENNESSEE": "TN", "TEXAS": "TX",
        "UTAH": "UT", "VERMONT": "VT", "VIRGINIA": "VA", "WASHINGTON": "WA",
        "WEST VIRGINIA": "WV", "WISCONSIN": "WI", "WYOMING": "WY",
        "DISTRICT OF COLUMBIA": "DC",
    }
    def _normalize_state(s):
        s = s.strip().upper()
        return _state_abbrevs.get(s, s)

    def _loc_key(loc):
        return (_normalize_addr(loc.get("address", "")) + "|" +
                _normalize_city(loc.get("city", "")) + "|" +
                _normalize_state(loc.get("state", "")))
    
    # Property locations: unique by composite key
    _prop_unique_keys = set()
    if sov_locs:
        for loc in sov_locs:
            key = _loc_key(loc)
            if key != "||":
                _prop_unique_keys.add(key)
    prop_loc_count = len(_prop_unique_keys) if _prop_unique_keys else 0
    if not prop_loc_count and coverages.get("property"):
        prop_loc_count = int(coverages["property"].get("num_locations", 0) or 0)
    
    # Count liability locations from schedule_of_classes (skip non-physical entries)
    _skip_gl_classes = {"hired auto", "non-owned auto", "loss control", "package store",
                        "category vi", "liquor", "sundry", "flat"}
    gl_loc_keys = set()
    # Also count unique location IDENTIFIERS (e.g., "Primary", "location#3") as reliable GL count
    _gl_loc_ids = set()
    if isinstance(gl_cov, dict):
        for entry in gl_cov.get("schedule_of_classes", []):
            if isinstance(entry, dict):
                classification = (entry.get("classification", "") or "").lower()
                if any(skip in classification for skip in _skip_gl_classes):
                    continue
                # Track unique location IDs (even if not real addresses)
                loc_id = (entry.get("location", "") or "").strip().lower()
                if loc_id and loc_id not in ("", "n/a", "all", "various"):
                    _gl_loc_ids.add(loc_id)
                addr = entry.get("address", "")
                if addr and addr.strip() and len(addr.strip()) > 5:
                    # Only use address if it looks like a real street address (not "Primary" or "location#3")
                    import re as _re2
                    # Real addresses contain numbers and letters (e.g., "5370 Clearwater Court")
                    if _re2.search(r'\d+\s+\w+', addr):
                        parts = [p.strip() for p in addr.split(",")]
                        street = parts[0] if parts else addr
                        city = ""
                        state = ""
                        if len(parts) >= 3:
                            city = parts[1]
                            st_m = _re2.match(r'([A-Z]{2})\s*\d*', parts[2].strip().upper())
                            if st_m: state = st_m.group(1)
                        elif len(parts) == 2:
                            st_m = _re2.match(r'([A-Z]{2})\s*\d*', parts[1].strip().upper())
                            if st_m:
                                state = st_m.group(1)
                            else:
                                city = parts[1]
                        key = (_normalize_addr(street) + "|" + _normalize_city(city) + "|" + state.strip().upper())
                        if key != "||":
                            gl_loc_keys.add(key)
    # Use GL location IDs count if it's higher than address-matched count (more reliable when addresses aren't extracted)
    liab_loc_count = max(
        len(gl_loc_keys) if gl_loc_keys else 0,
        len(_gl_loc_ids) if _gl_loc_ids else 0,
        len(designated_premises) if designated_premises else 0,
        int(gl_cov.get("num_locations", 0) or 0) if isinstance(gl_cov, dict) else 0
    )
    # If designated_premises covers >=50% of property locations, GL likely covers all
    if designated_premises and prop_loc_count > 0 and len(designated_premises) >= prop_loc_count * 0.5:
        liab_loc_count = max(liab_loc_count, prop_loc_count)
    
    # Calculate UNIQUE total location count
    # When SOV is available, it is the AUTHORITATIVE source for location count
    # Only supplement with GL locations that are genuinely new (not already in SOV)
    all_unique_keys = set()
    # Add property keys (from SOV)
    all_unique_keys.update(_prop_unique_keys)
    # Add liability keys ONLY if they don't fuzzy-match existing property keys
    for gl_key in gl_loc_keys:
        gl_parts = gl_key.split("|")
        already_matched = False
        for existing_key in all_unique_keys:
            ex_parts = existing_key.split("|")
            if len(gl_parts) == 3 and len(ex_parts) == 3:
                state_ok = (not gl_parts[2] or not ex_parts[2] or _normalize_state(gl_parts[2]) == _normalize_state(ex_parts[2]))
                if state_ok and _fuzzy_addr_match(gl_parts[0], ex_parts[0]):
                    already_matched = True
                    break
        if not already_matched:
            all_unique_keys.add(gl_key)
    # NOTE: Do NOT add raw GPT-extracted locations to the count.
    # GPT often extracts mailing addresses, carrier addresses, and other non-physical
    # addresses that inflate the location count. SOV + GL schedule are authoritative.
    # When GL uses location identifiers (Primary, location#3) instead of addresses,
    # the GL location ID count is more reliable than gl_loc_keys (address-based count).
    # Use the higher of: merged address count OR GL unique location IDs count
    # (since some GL-only locations may not have matched addresses)
    _merged_addr_count = len(all_unique_keys) if all_unique_keys else 0
    _gl_id_count = len(_gl_loc_ids) if _gl_loc_ids else 0
    # If GL IDs suggest more locations than we found by address matching, use GL count
    # because the extra GL locations may have location references instead of parseable addresses
    if _gl_id_count > _merged_addr_count:
        # GL has locations we couldn't match by address — compute: prop unique + GL-only delta
        total_loc_count = max(_merged_addr_count, prop_loc_count + max(0, _gl_id_count - prop_loc_count))
    else:
        total_loc_count = _merged_addr_count if _merged_addr_count else max(prop_loc_count, liab_loc_count)
    # Final check: total must be at least max(property, liability)
    total_loc_count = max(total_loc_count, prop_loc_count, liab_loc_count)
    if total_loc_count > 0:
        rows.append(["Total Number of Locations", str(total_loc_count)])
    if prop_loc_count > 0:
        rows.append(["Property Locations", str(prop_loc_count)])
    if liab_loc_count > 0:
        rows.append(["Liability Locations", str(liab_loc_count)])
    
    # Count location types from SOV/property locations and GL schedule_of_classes
    # Use SOV descriptions and GL classifications for PHYSICAL locations only
    # (skip GL exposure-only classes like Hired Auto, Loss Control, Package Stores)
    hotel_count = 0
    office_count = 0
    lro_count = 0
    vacant_count = 0
    other_types = {}
    seen_loc_addrs = []  # track normalized addresses for fuzzy dedup
    
    def _addr_already_seen(addr):
        """Check if addr fuzzy-matches any already-seen address."""
        for existing in seen_loc_addrs:
            if _fuzzy_addr_match(addr, existing):
                return True
        return False
    
    # Hotel brand keywords for classification
    _hotel_keywords = [
        "hotel", "motel", "inn", "suite", "lodge", "resort", "hampton", "holiday",
        "best western", "marriott", "hilton", "ihg", "wyndham", "choice", "comfort",
        "quality", "candlewood", "towneplace", "springhill", "hyatt", "latitude",
        "fairfield", "courtyard", "residence", "doubletree", "embassy", "homewood",
        "home2", "tru by", "avid", "la quinta", "sonesta", "days inn", "super 8",
        "ramada", "microtel", "wingate", "baymont", "americinn", "country inn",
        "red roof", "econo lodge", "rodeway", "clarion", "cambria", "ascend",
        "sleep inn", "mainstay", "suburban", "woodspring", "extended stay",
        "staybridge", "crowne plaza", "intercontinental", "kimpton", "indigo",
        "even hotel", "atwell", "avid hotel", "candlewood suites",
        "non-franchised", "motor lodge", "simply suites", "select",
    ]

    # First: count from SOV locations (most reliable source for property types)
    if sov_data and sov_data.get("locations"):
        for loc in sov_data["locations"]:
            addr = _normalize_addr(loc.get("address", ""))
            if _addr_already_seen(addr):
                continue
            seen_loc_addrs.append(addr)
            # Determine type: check ALL relevant fields (not just first truthy one)
            # Combine all descriptive fields for comprehensive matching
            all_desc_parts = [
                (loc.get("description", "") or "").lower(),
                (loc.get("hotel_flag", "") or "").lower(),
                (loc.get("occupancy", "") or "").lower(),
                (loc.get("dba", "") or "").lower(),
            ]
            combined_desc = " ".join(all_desc_parts)
            if any(kw in combined_desc for kw in _hotel_keywords):
                hotel_count += 1
            elif "office" in combined_desc or "corporate" in combined_desc:
                office_count += 1
            elif "vacant" in combined_desc or "land" in combined_desc:
                vacant_count += 1
            elif "lessor" in combined_desc or "lro" in combined_desc:
                lro_count += 1
            else:
                # Default: if it has a building value, it's likely a hotel
                bldg = loc.get("building_value", 0) or 0
                if bldg > 0:
                    hotel_count += 1
                else:
                    vacant_count += 1
    
    # Second: supplement from GL schedule_of_classes (only PHYSICAL location entries)
    # Skip non-location exposure classes
    _skip_classes = {"hired auto", "non-owned auto", "loss control", "package store",
                     "category vi", "liquor", "sundry", "flat"}
    if isinstance(gl_cov, dict):
        for entry in gl_cov.get("schedule_of_classes", []):
            if isinstance(entry, dict):
                addr = _normalize_addr(entry.get("address", "") or entry.get("location", ""))
                classification = (entry.get("classification", "") or "").lower()
                # Skip non-physical-location exposure classes
                if any(skip in classification for skip in _skip_classes):
                    continue
                if not addr or _addr_already_seen(addr):
                    continue
                seen_loc_addrs.append(addr)
                # Also check brand_dba for hotel identification
                brand_dba = (entry.get("brand_dba", "") or "").lower()
                combined_class = classification + " " + brand_dba
                if any(kw in combined_class for kw in _hotel_keywords):
                    hotel_count += 1
                elif "office" in classification or "building" in classification:
                    office_count += 1
                elif "lessor" in classification or "lro" in classification:
                    lro_count += 1
                elif "vacant" in classification:
                    vacant_count += 1
                elif classification.strip():
                    type_name = classification.split("-")[0].strip().title()
                    other_types[type_name] = other_types.get(type_name, 0) + 1
    
    type_parts = []
    if hotel_count: type_parts.append(f"{hotel_count} Hotel(s)")
    if office_count: type_parts.append(f"{office_count} Office(s)")
    if lro_count: type_parts.append(f"{lro_count} LRO(s)")
    if vacant_count: type_parts.append(f"{vacant_count} Vacant Land")
    for ot in sorted(other_types):
        type_parts.append(f"{other_types[ot]} {ot}")
    if type_parts:
        rows.append(["Location Types", ", ".join(type_parts)])
    
    # Add TIV from SOV or property quote
    _tiv_added = False
    if sov_data and sov_data.get("totals", {}).get("tiv"):
        rows.append(["Total Insured Value (TIV)", fmt_currency(sov_data["totals"]["tiv"])])
        _tiv_added = True
    elif sov_data and sov_data.get("locations"):
        # Calculate TIV from individual SOV locations
        _total_tiv = sum(loc.get("tiv", 0) or 0 for loc in sov_data["locations"])
        if _total_tiv > 0:
            rows.append(["Total Insured Value (TIV)", fmt_currency(_total_tiv)])
            _tiv_added = True
    if not _tiv_added:
        prop_cov = coverages.get("property", {})
        if isinstance(prop_cov, dict):
            prop_tiv = prop_cov.get("tiv", "")
            if prop_tiv:
                rows.append(["Total Insured Value (TIV)", prop_tiv if isinstance(prop_tiv, str) else fmt_currency(prop_tiv)])
                _tiv_added = True
            elif prop_cov.get("limits"):
                # Try to sum building + BPP + BI from property limits
                _prop_total = 0
                for lim in prop_cov["limits"]:
                    if isinstance(lim, dict):
                        lim_val = lim.get("limit", "")
                        if isinstance(lim_val, str):
                            import re as _re2
                            cleaned = _re2.sub(r'[^\d.]', '', lim_val.replace(',', ''))
                            if cleaned:
                                try:
                                    _prop_total += float(cleaned)
                                except ValueError:
                                    pass
                if _prop_total > 0:
                    rows.append(["Total Insured Value (TIV)", fmt_currency(_prop_total)])
    
    # Add Umbrella / Excess total limit
    # Include ALL umbrella/excess layer keys: primary, alt options (used as layers), and explicit layers
    _umbrella_total = 0
    _umbrella_seen_carriers = set()  # Track carriers to avoid double-counting competing options
    _umbrella_keys = [
        "umbrella", "umbrella_alt_1", "umbrella_alt_2", "umbrella_alt_3",
        "umbrella_layer_2", "umbrella_layer_3", "umbrella_layer_4",
        "excess_liability", "excess",
    ]
    for umb_key in _umbrella_keys:
        umb_cov = coverages.get(umb_key, {})
        if isinstance(umb_cov, dict) and umb_cov.get("carrier"):
            carrier_name = (umb_cov.get("carrier", "") or "").strip().lower()
            # Skip if same carrier already counted (avoid double-counting from alias keys)
            if carrier_name in _umbrella_seen_carriers:
                continue
            _umbrella_seen_carriers.add(carrier_name)
            for lim in umb_cov.get("limits", []):
                if isinstance(lim, dict):
                    desc = (lim.get("description", "") or "").lower()
                    if "occurrence" in desc or "each occurrence" in desc:
                        lim_val = lim.get("limit", "")
                        if isinstance(lim_val, str):
                            import re as _re3
                            cleaned = _re3.sub(r'[^\d.]', '', lim_val.replace(',', ''))
                            if cleaned:
                                try:
                                    _umbrella_total += float(cleaned)
                                except ValueError:
                                    pass
                        elif isinstance(lim_val, (int, float)):
                            _umbrella_total += lim_val
                        break  # Only count each occurrence limit once per layer
    if _umbrella_total > 0:
        rows.append(["Total Umbrella / Excess Limit", fmt_currency(_umbrella_total)])
    
    create_styled_table(doc, headers, rows, col_widths=[2.5, 5.0],
                       header_size=10, body_size=10)
    
    add_formatted_paragraph(doc, "", space_before=8)
    add_callout_box(doc, "The information contained in this proposal is based on data provided by the insured and/or their representatives. HUB International makes no warranty as to the accuracy of this information.")


def _normalize_addr(s):
    """Normalize street address for dedup: uppercase, strip, replace common variants.
    Handles U.S. 51 / US 51 / US-51 / Highway 51 / Hwy 51 all mapping to the same form.
    Also strips trailing zip codes."""
    import re as _re_norm
    s = s.strip().upper()
    # Remove periods, commas, and dashes ("Burlington - Mount Holly" -> "Burlington Mount Holly")
    s = s.replace(".", "").replace(",", "").replace(" - ", " ").replace("-", " ")
    # Normalize route designators: "U.S. 51" / "US 51" / "US-51" / "US HWY 51" -> "HWY 51"
    s = _re_norm.sub(r'\bUS\s*-?\s*(\d)', r'HWY \1', s)
    s = _re_norm.sub(r'\bU\s*S\s*-?\s*(\d)', r'HWY \1', s)
    # Normalize word-level abbreviations FIRST (before suffix replacements)
    # This handles both directions: "MOUNT" -> "MT", and ensures consistency
    word_replacements = {
        "MOUNT": "MT", "SAINT": "ST", "FORT": "FT",
        "TOWNSHIP": "TWP", "COUNTY": "CTY",
    }
    words = s.split()
    words = [word_replacements.get(w, w) for w in words]
    s = " ".join(words)
    # Normalize suffix-level replacements
    replacements = {
        " STREET": " ST", " AVENUE": " AVE", " BOULEVARD": " BLVD",
        " DRIVE": " DR", " ROAD": " RD", " LANE": " LN",
        " COURT": " CT", " PLACE": " PL", " CIRCLE": " CIR",
        " HIGHWAY": " HWY", " PARKWAY": " PKWY", " TERRACE": " TER",
        " NORTH": " N", " SOUTH": " S", " EAST": " E", " WEST": " W",
        " NORTHWEST": " NW", " NORTHEAST": " NE", " SOUTHWEST": " SW",
        " SOUTHEAST": " SE",
    }
    for old, new in replacements.items():
        s = s.replace(old, new)
    # Normalize state names to abbreviations (TEXAS -> TX, LOUISIANA -> LA, etc.)
    _state_names = {
        "ALABAMA": "AL", "ALASKA": "AK", "ARIZONA": "AZ", "ARKANSAS": "AR",
        "CALIFORNIA": "CA", "COLORADO": "CO", "CONNECTICUT": "CT", "DELAWARE": "DE",
        "FLORIDA": "FL", "GEORGIA": "GA", "HAWAII": "HI", "IDAHO": "ID",
        "ILLINOIS": "IL", "INDIANA": "IN", "IOWA": "IA", "KANSAS": "KS",
        "KENTUCKY": "KY", "LOUISIANA": "LA", "MAINE": "ME", "MARYLAND": "MD",
        "MASSACHUSETTS": "MA", "MICHIGAN": "MI", "MINNESOTA": "MN", "MISSISSIPPI": "MS",
        "MISSOURI": "MO", "MONTANA": "MT", "NEBRASKA": "NE", "NEVADA": "NV",
        "NEW HAMPSHIRE": "NH", "NEW JERSEY": "NJ", "NEW MEXICO": "NM", "NEW YORK": "NY",
        "NORTH CAROLINA": "NC", "NORTH DAKOTA": "ND", "OHIO": "OH", "OKLAHOMA": "OK",
        "OREGON": "OR", "PENNSYLVANIA": "PA", "RHODE ISLAND": "RI",
        "SOUTH CAROLINA": "SC", "SOUTH DAKOTA": "SD", "TENNESSEE": "TN", "TEXAS": "TX",
        "UTAH": "UT", "VERMONT": "VT", "VIRGINIA": "VA", "WASHINGTON": "WA",
        "WEST VIRGINIA": "WV", "WISCONSIN": "WI", "WYOMING": "WY",
        "DISTRICT OF COLUMBIA": "DC",
    }
    for state_name, state_abbr in _state_names.items():
        s = _re_norm.sub(r"\b" + state_name + r"\b", state_abbr, s)
    # Strip trailing country names
    s = _re_norm.sub(r"\s+(UNITED STATES|USA|US)\s*$", "", s)
    # Strip trailing zip codes (5-digit or 5+4)
    s = _re_norm.sub(r'\s+\d{5}(-\d{4})?\s*$', '', s)
    s = " ".join(s.split())
    return s


def _normalize_city(s):
    """Normalize city name for dedup: uppercase, remove spaces/punctuation.
    Handles 'La Place' vs 'LaPlace' vs 'LA PLACE' all mapping to 'LAPLACE'."""
    s = s.strip().upper()
    s = s.replace(".", "").replace(",", "").replace("-", "").replace("'", "")
    # Remove ALL spaces so 'LA PLACE' == 'LAPLACE' == 'LA  PLACE'
    s = s.replace(" ", "")
    return s


def _levenshtein(s1, s2):
    """Compute Levenshtein edit distance between two strings."""
    if len(s1) < len(s2):
        return _levenshtein(s2, s1)
    if len(s2) == 0:
        return len(s1)
    prev_row = range(len(s2) + 1)
    for i, c1 in enumerate(s1):
        curr_row = [i + 1]
        for j, c2 in enumerate(s2):
            insertions = prev_row[j + 1] + 1
            deletions = curr_row[j] + 1
            substitutions = prev_row[j] + (c1 != c2)
            curr_row.append(min(insertions, deletions, substitutions))
        prev_row = curr_row
    return prev_row[-1]


def _words_fuzzy_equal(w1, w2):
    """Check if two words are equal or differ by at most 2 edits (for words >= 5 chars)."""
    if w1 == w2:
        return True
    # For short words (< 5 chars), require exact match
    if len(w1) < 5 or len(w2) < 5:
        return False
    # Allow edit distance up to 2 for longer words (catches typos like PARRAMORE vs PARRAMOREE)
    max_dist = 2 if min(len(w1), len(w2)) >= 7 else 1
    return _levenshtein(w1, w2) <= max_dist


def _fuzzy_addr_match(addr1, addr2):
    """Check if two normalized addresses refer to the same location.
    Handles cases like '4288 HWY 51' vs '4285 HWY 51' by comparing
    the street name portion after stripping house numbers.
    Also handles typos like 'PARRAMORE RD' vs 'PARRAMOREE RD' via edit distance."""
    import re
    if not addr1 or not addr2:
        return False
    if addr1 == addr2:
        return True
    if addr1 in addr2 or addr2 in addr1:
        return True
    # Extract street name without house number for fuzzy match
    num1 = re.match(r'^(\d+)\s+(.+)', addr1)
    num2 = re.match(r'^(\d+)\s+(.+)', addr2)
    if num1 and num2:
        street1 = num1.group(2)
        street2 = num2.group(2)
        house1 = int(num1.group(1))
        house2 = int(num2.group(1))
        # Same house number (or within 20) and street names match
        if abs(house1 - house2) <= 20:
            if street1 == street2:
                return True
            # Word-level match with edit distance tolerance
            words1 = street1.split()
            words2 = street2.split()
            words1_set = set(words1)
            words2_set = set(words2)
            # Remove common suffixes for comparison
            _suffixes = {'ST', 'AVE', 'BLVD', 'DR', 'RD', 'LN', 'CT', 'PL', 'CIR', 'HWY', 'PKWY', 'TER', 'WAY'}
            core1 = words1_set - _suffixes
            core2 = words2_set - _suffixes
            # Exact subset check
            if core1 and core2 and (core1.issubset(core2) or core2.issubset(core1)):
                return True
            # Fuzzy word matching: check if each core word in the smaller set
            # has a fuzzy match in the larger set (catches typos)
            if core1 and core2:
                smaller, larger = (core1, core2) if len(core1) <= len(core2) else (core2, core1)
                all_match = True
                for sw in smaller:
                    found = False
                    for lw in larger:
                        if _words_fuzzy_equal(sw, lw):
                            found = True
                            break
                    if not found:
                        all_match = False
                        break
                if all_match:
                    return True
            # Also check if one is a substring of the other at word level
            if street1 in street2 or street2 in street1:
                return True
            # Edit distance on full street name (for very similar streets)
            if len(street1) >= 5 and len(street2) >= 5:
                max_dist = 2 if min(len(street1), len(street2)) >= 10 else 1
                if _levenshtein(street1, street2) <= max_dist:
                    return True
    return False


def _dedup_locations(raw_locations):
    """Deduplicate locations by normalized address."""
    seen_addrs = set()
    locations = []
    for loc in raw_locations:
        addr_key = (_normalize_addr(loc.get("address", "")) + "|" + 
                    _normalize_city(loc.get("city", "")) + "|" +
                    loc.get("state", "").strip().upper())
        if addr_key not in seen_addrs:
            seen_addrs.add(addr_key)
            locations.append(loc)
    return locations


def generate_locations(doc, data):
    """Section 8: Locations — unified schedule with Property/Liability coverage checkmarks."""
    import re
    add_page_break(doc)
    add_section_header(doc, "Schedule of Locations")
    
    raw_locations = data.get("locations", [])
    locations = _dedup_locations(raw_locations)
    sov_data = data.get("sov_data")
    coverages = data.get("coverages", {})
    
    # Build a master list of all locations from all sources
    # Each entry: {name, address, city, state, tiv, on_property, on_liability}
    master_locations = []
    seen_addr_keys = set()
    
    # --- Determine which addresses are on the PROPERTY policy ---
    property_addr_keys = set()
    if sov_data and sov_data.get("locations"):
        for loc in sov_data["locations"]:
            addr_key = (_normalize_addr(loc.get("address", "")) + "|" +
                       _normalize_city(loc.get("city", "")) + "|" +
                       loc.get("state", "").strip().upper())
            property_addr_keys.add(addr_key)
    elif "property" in coverages:
        # If no SOV, try to use property quote's schedule_of_values for property addresses
        prop_cov = coverages.get("property", {})
        prop_sov = prop_cov.get("schedule_of_values", []) if isinstance(prop_cov, dict) else []
        if prop_sov:
            for s in prop_sov:
                if isinstance(s, dict):
                    addr = s.get("address", "") or s.get("location", "")
                    if addr:
                        parts = [p.strip() for p in addr.split(",")]
                        street = parts[0] if parts else addr
                        city = parts[1] if len(parts) >= 2 else ""
                        state = ""
                        if len(parts) >= 3:
                            st_m = re.match(r'([A-Z]{2})', parts[2].strip().upper())
                            if st_m: state = st_m.group(1)
                        addr_key = (_normalize_addr(street) + "|" +
                                   _normalize_city(city) + "|" +
                                   state.strip().upper())
                        property_addr_keys.add(addr_key)
        # If still no property addresses from schedule_of_values, try matching
        # property coverage's location/address fields from the extraction
        if not property_addr_keys:
            prop_addr = prop_cov.get("address", "") or prop_cov.get("location", "") or prop_cov.get("insured_location", "")
            if prop_addr and isinstance(prop_addr, str) and len(prop_addr.strip()) > 4:
                parts = [p.strip() for p in prop_addr.split(",")]
                street = parts[0] if parts else prop_addr
                city = parts[1] if len(parts) >= 2 else ""
                state = ""
                if len(parts) >= 3:
                    st_m = re.match(r'([A-Z]{2})', parts[2].strip().upper())
                    if st_m: state = st_m.group(1)
                addr_key = (_normalize_addr(street) + "|" +
                           _normalize_city(city) + "|" +
                           state.strip().upper())
                property_addr_keys.add(addr_key)

    # FALLBACK: If property coverage exists but we still have NO property addresses
    # (no SOV uploaded, no schedule_of_values, no parseable address fields),
    # then all locations should get the property checkmark. A property quote
    # was uploaded — those locations are on it.
    _property_covers_all = ("property" in coverages and not property_addr_keys)
    if _property_covers_all:
        logger.info("Schedule of Locations - property quote exists but no addresses extracted; defaulting all locations to property-covered")

    # --- Determine which addresses are on the LIABILITY policy ---
    liability_addr_keys = set()
    gl_cov = coverages.get("general_liability", {})
    gl_classes = gl_cov.get("schedule_of_classes", []) if isinstance(gl_cov, dict) else []
    for entry in gl_classes:
        if isinstance(entry, dict):
            addr = entry.get("address", "")
            if addr:
                norm_gl_addr = _normalize_addr(addr)
                # Try to find matching city/state from SOV or locations using fuzzy matching
                matched = False
                for loc in (sov_data.get("locations", []) if sov_data else []) + locations:
                    loc_addr_norm = _normalize_addr(loc.get("address", ""))
                    if _fuzzy_addr_match(norm_gl_addr, loc_addr_norm):
                        addr_key = (loc_addr_norm + "|" +
                                   _normalize_city(loc.get("city", "")) + "|" +
                                   loc.get("state", "").strip().upper())
                        liability_addr_keys.add(addr_key)
                        matched = True
                        break
                if not matched:
                    # Also try parsing city/state from the GL address itself
                    parts = [p.strip() for p in addr.split(",")]
                    street = parts[0] if parts else addr
                    city = ""
                    state = ""
                    if len(parts) >= 3:
                        city = parts[1]
                        st_m = re.match(r'([A-Z]{2})\s*\d*', parts[2].strip().upper())
                        if st_m: state = st_m.group(1)
                    elif len(parts) == 2:
                        st_m = re.match(r'([A-Z]{2})\s*\d*', parts[1].strip().upper())
                        if st_m:
                            state = st_m.group(1)
                        else:
                            city = parts[1]
                    liability_addr_keys.add(_normalize_addr(street) + "|" + _normalize_city(city) + "|" + state.strip().upper())
    logger.info(f"Schedule of Locations - liability_addr_keys: {liability_addr_keys}")

    # Count unique PHYSICAL locations on the GL schedule_of_classes
    # (GL entries use identifiers like "Primary", "location#3" rather than addresses)
    _skip_gl_loc_classes = {"hired auto", "non-owned auto", "loss control", "package store",
                            "category vi", "sundry", "flat"}
    gl_unique_loc_ids = set()
    for entry in gl_classes:
        if isinstance(entry, dict):
            loc_id = (entry.get("location", "") or "").strip().lower()
            classification = (entry.get("classification", "") or "").lower()
            if any(skip in classification for skip in _skip_gl_loc_classes):
                continue
            # Only count entries that represent physical locations
            if loc_id and loc_id not in ("", "n/a", "all", "various"):
                gl_unique_loc_ids.add(loc_id)

    # If GL schedule has >= as many unique locations as property, GL covers ALL property locations
    # In this case, mark all property addresses as liability-covered
    if gl_unique_loc_ids and len(gl_unique_loc_ids) >= len(property_addr_keys):
        logger.info(f"Schedule of Locations - GL has {len(gl_unique_loc_ids)} unique locations "
                     f">= {len(property_addr_keys)} property locations. Marking ALL property locations as liability-covered.")
        liability_addr_keys.update(property_addr_keys)
    elif gl_unique_loc_ids and len(gl_unique_loc_ids) > len(liability_addr_keys):
        # GL has more locations than we matched — likely all property locations are GL-covered
        # but we couldn't match addresses because GL uses "location#N" references
        logger.info(f"Schedule of Locations - GL has {len(gl_unique_loc_ids)} unique locations "
                     f"but only {len(liability_addr_keys)} address matches. Adding all property addresses.")
        liability_addr_keys.update(property_addr_keys)

    # BLANKET GL RULE: In hospitality, if a GL policy exists with a carrier
    # and schedule_of_classes entries, the GL covers ALL property locations.
    # Schedule_of_classes is for rating/classification, not coverage limitation.
    _gl_carrier = gl_cov.get("carrier", "") if isinstance(gl_cov, dict) else ""
    if _gl_carrier and len(gl_classes) >= 3 and property_addr_keys:
        if not liability_addr_keys.issuperset(property_addr_keys):
            logger.info(f"GL blanket rule: carrier with {len(gl_classes)} schedule entries "
                         f"-- marking all {len(property_addr_keys)} property locations as GL-covered")
            liability_addr_keys.update(property_addr_keys)

    
    # --- Pre-scan designated premises to add to liability_addr_keys ---
    
    # PRIMARY SOURCE: designated_premises array (extracted by GPT from CG2144/NXLL110)
    _cg2144_addrs = []  # Save parsed addresses for the fourth pass later
    designated_premises = gl_cov.get("designated_premises", []) if isinstance(gl_cov, dict) else []
    for raw_addr in designated_premises:
        if not isinstance(raw_addr, str) or not raw_addr.strip() or len(raw_addr.strip()) < 5:
            continue
        raw_addr = raw_addr.strip()
        _cg2144_addrs.append(raw_addr)
        parts = [p.strip() for p in raw_addr.split(",")]
        street = parts[0] if parts else raw_addr
        city = ""
        state = ""
        if len(parts) >= 3:
            city = parts[1]
            st_m = re.match(r'([A-Z]{2})\s*\d*', parts[2].strip().upper())
            if st_m: state = st_m.group(1)
        elif len(parts) == 2:
            st_m = re.match(r'([A-Z]{2})\s*\d*', parts[1].strip().upper())
            if st_m:
                state = st_m.group(1)
            else:
                city = parts[1]
        addr_key = (_normalize_addr(street) + "|" + _normalize_city(city) + "|" + state.strip().upper())
        liability_addr_keys.add(addr_key)
        # Also try matching against SOV/locations for better key resolution using fuzzy matching
        norm_street = _normalize_addr(street)
        for loc in (sov_data.get("locations", []) if sov_data else []) + locations:
            loc_addr_norm = _normalize_addr(loc.get("address", ""))
            if _fuzzy_addr_match(norm_street, loc_addr_norm):
                resolved_key = (loc_addr_norm + "|" +
                               _normalize_city(loc.get("city", "")) + "|" +
                               loc.get("state", "").strip().upper())
                liability_addr_keys.add(resolved_key)
                break
    
    # FALLBACK: Parse CG2144/NXLL110 form descriptions for addresses
    gl_forms = gl_cov.get("forms_endorsements", []) if isinstance(gl_cov, dict) else []
    for form in gl_forms:
        if not isinstance(form, dict):
            continue
        desc = (form.get("description", "") or "").upper()
        if not any(kw in desc for kw in ["DESIGNATED PREMISES", "CG 21 44", "CG2144", "NXLL110", "NXLL 110", "LIMITATION OF COVERAGE"]):
            continue
        # Parse numbered addresses: "1) 4285 Highway 51, LaPlace, LA 70068"
        addr_pattern = re.findall(r'\d+\)\s*(.+?)(?=\s*\d+\)|$)', desc, re.DOTALL)
        if not addr_pattern:
            addr_pattern = [a.strip() for a in re.split(r'[;\n]', desc) if re.search(r'\d+\s+\w+', a.strip())]
        for raw_addr in addr_pattern:
            raw_addr = raw_addr.strip().rstrip(',')
            if not raw_addr or len(raw_addr) < 5:
                continue
            _cg2144_addrs.append(raw_addr)
            # Parse city/state from address
            parts = [p.strip() for p in raw_addr.split(",")]
            street = parts[0] if parts else raw_addr
            city = ""
            state = ""
            if len(parts) >= 3:
                city = parts[1]
                st_m = re.match(r'([A-Z]{2})\s*\d*', parts[2].strip().upper())
                if st_m: state = st_m.group(1)
            elif len(parts) == 2:
                st_m = re.match(r'([A-Z]{2})\s*\d*', parts[1].strip().upper())
                if st_m:
                    state = st_m.group(1)
                else:
                    city = parts[1]
            addr_key = (_normalize_addr(street) + "|" + _normalize_city(city) + "|" + state.strip().upper())
            liability_addr_keys.add(addr_key)
            # Also try matching against SOV/locations for better key resolution using fuzzy matching
            norm_street_cg = _normalize_addr(street)
            for loc in (sov_data.get("locations", []) if sov_data else []) + locations:
                loc_addr_norm_cg = _normalize_addr(loc.get("address", ""))
                if _fuzzy_addr_match(norm_street_cg, loc_addr_norm_cg):
                    resolved_key = (loc_addr_norm_cg + "|" +
                                   _normalize_city(loc.get("city", "")) + "|" +
                                   loc.get("state", "").strip().upper())
                    liability_addr_keys.add(resolved_key)
                    break
    
    # DESIGNATED PREMISES FALLBACK: If GL designated_premises count >= property locations,
    # all property locations are liability-covered (GL covers at least as many as property)
    if designated_premises and len(designated_premises) >= len(property_addr_keys) and property_addr_keys:
        logger.info(f"Schedule of Locations - designated_premises has {len(designated_premises)} entries "
                     f">= {len(property_addr_keys)} property locations. Marking ALL property as liability-covered.")
        liability_addr_keys.update(property_addr_keys)

    # NOTE: Blanket liability fallback removed. Multi-pass extraction (Pass 3) in
    # proposal_extractor.py now handles focused address extraction for GL when
    # designated_premises is empty. Liability checkmarks are only applied to
    # locations explicitly confirmed on the liability quote per Stefan's rule.

    # FALLBACK: If GL quote has sparse data ("See attached for Schedule of Locations"),
    # assume ALL property locations are GL-covered. This handles renewals where the
    # GL quote says "See attached for Schedule of Locations" but the attachment isn't in the PDF.
    _gl_has_sparse_data = (len(gl_classes) + len(designated_premises)) < 3
    if _gl_has_sparse_data and property_addr_keys and sov_data and sov_data.get("locations"):
        logger.info(f"GL has sparse location data ({len(gl_classes)} classes, {len(designated_premises)} designated_premises) "
                   f"— assuming all {len(property_addr_keys)} property locations are GL-covered")
        liability_addr_keys.update(property_addr_keys)

    # HEURISTIC: If GL designated_premises covers majority (>=50%) of property locations,
    # the GL likely covers ALL property locations but extraction missed some addresses.
    if not _gl_has_sparse_data and designated_premises and property_addr_keys:
        _dp_matched = 0
        for dp_addr in designated_premises:
            if isinstance(dp_addr, str) and dp_addr.strip():
                dp_norm = _normalize_addr(dp_addr)
                for pk in property_addr_keys:
                    pk_addr = pk.split("|")[0] if "|" in pk else pk
                    if _fuzzy_addr_match(dp_norm, pk_addr):
                        _dp_matched += 1
                        break
        if _dp_matched >= len(property_addr_keys) * 0.5 and _dp_matched < len(property_addr_keys):
            logger.info(f"GL designated_premises matched {_dp_matched} of {len(property_addr_keys)} property locations "
                       f"(>=50%) -- marking all property locations as GL-covered")
            liability_addr_keys.update(property_addr_keys)
    
    # --- Build master location list ---
    # _fuzzy_addr_match is now a module-level function (used by both generate_locations and generate_information_summary)
    
    # Helper: check if an addr_key matches any key in liability_addr_keys using fuzzy matching
    def _is_on_liability(addr_key):
        """Check if addr_key is in liability_addr_keys, with fuzzy address matching."""
        if addr_key in liability_addr_keys:
            return True
        # Fuzzy match: compare the street portion of the addr_key against all liability keys
        parts = addr_key.split("|")
        if len(parts) != 3:
            return False
        addr_norm = parts[0]
        state_norm = parts[2]
        for lk in liability_addr_keys:
            lk_parts = lk.split("|")
            if len(lk_parts) != 3:
                continue
            # Must match state (or one is empty)
            if state_norm and lk_parts[2] and state_norm != lk_parts[2]:
                continue
            if _fuzzy_addr_match(addr_norm, lk_parts[0]):
                return True
        return False
    
    # Helper: check if an addr_key matches any key in property_addr_keys
    # Uses STRICT matching (exact normalized address) — no fuzzy tolerance
    # Property checkmarks must be precise: only SOV/property quote locations
    def _is_on_property(addr_key):
        """Check if addr_key is in property_addr_keys. Strict match only."""
        if _property_covers_all:
            return True
        if addr_key in property_addr_keys:
            return True
        # Also try matching just the street portion (ignoring city differences)
        parts = addr_key.split("|")
        if len(parts) != 3:
            return False
        addr_norm = parts[0]
        state_norm = parts[2]
        for pk in property_addr_keys:
            pk_parts = pk.split("|")
            if len(pk_parts) != 3:
                continue
            if state_norm and pk_parts[2] and state_norm != pk_parts[2]:
                continue
            # STRICT: normalized addresses must be identical (no house number tolerance)
            if addr_norm == pk_parts[0]:
                return True
        return False
    
    # First: SOV locations (property locations)
    # Determine fallback corporate name from SOV summary or client_info
    _fallback_corp = ""
    if sov_data and sov_data.get("summary", {}).get("named_insured"):
        _ni = sov_data["summary"]["named_insured"]
        # If named_insured contains " - ", split into corp and DBA
        if " - " in _ni:
            _fallback_corp = _ni.split(" - ")[0].strip()
        else:
            _fallback_corp = _ni.strip()
    if not _fallback_corp:
        _fallback_corp = (data.get("client_info", {}).get("client_name", "") or "").strip()
    
    # Build a lookup of user-edited location names from data["locations"] array
    # The web UI stores name overrides in locations[i].name
    _name_overrides = {}  # addr_key -> user-edited name
    for ui_loc in data.get("locations", []):
        ui_addr_key = (_normalize_addr(ui_loc.get("address", "")) + "|" +
                      _normalize_city(ui_loc.get("city", "")) + "|" +
                      ui_loc.get("state", "").strip().upper())
        ui_name = (ui_loc.get("name", "") or "").strip()
        if ui_name and ui_addr_key != "||":
            _name_overrides[ui_addr_key] = ui_name
    
    # SKIP vacant land — it belongs on property SOV but NOT on the Schedule of Locations
    if sov_data and sov_data.get("locations"):
        for loc in sov_data["locations"]:
            # Filter out vacant land from Schedule of Locations
            desc = (loc.get("description", "") or loc.get("hotel_flag", "") or
                    loc.get("occupancy", "") or loc.get("dba", "") or "").lower()
            if "vacant" in desc or ("land" in desc and "hotel" not in desc and "inn" not in desc):
                # Still track the addr_key so we don't re-add it later
                addr_key = (_normalize_addr(loc.get("address", "")) + "|" +
                           _normalize_city(loc.get("city", "")) + "|" +
                           loc.get("state", "").strip().upper())
                seen_addr_keys.add(addr_key)
                continue
            
            addr_key = (_normalize_addr(loc.get("address", "")) + "|" +
                       _normalize_city(loc.get("city", "")) + "|" +
                       loc.get("state", "").strip().upper())
            
            # Check for user-edited name override first
            if addr_key in _name_overrides:
                name = _name_overrides[addr_key]
            else:
                # Build "Corporate Name - DBA" format for property name
                corporate_name = (loc.get("corporate_name", "") or "").strip()
                if not corporate_name:
                    corporate_name = _fallback_corp
                dba = (loc.get("dba", "") or loc.get("hotel_flag", "") or "").strip()
                if corporate_name and dba:
                    name = f"{corporate_name} - {dba}"
                elif dba:
                    name = dba
                elif corporate_name:
                    name = corporate_name
                else:
                    name = "Pending"
            tiv = loc.get("tiv", 0)
            on_liab = _is_on_liability(addr_key)
            logger.info(f"Schedule of Locations - SOV loc: addr_key={addr_key}, name={name}, on_liability={on_liab}")
            master_locations.append({
                "name": name,
                "address": loc.get("address", ""),
                "city": loc.get("city", ""),
                "state": loc.get("state", ""),
                "tiv": tiv,
                "on_property": _is_on_property(addr_key),
                "on_liability": on_liab,
            })
            seen_addr_keys.add(addr_key)
    
    # Second: extracted locations not already in SOV (skip vacant land)
    for loc in locations:
        desc_check = (loc.get("description", "") or loc.get("corporate_entity", "") or "").lower()
        if "vacant" in desc_check or ("land" in desc_check and "hotel" not in desc_check):
            # Track but skip
            addr_key = (_normalize_addr(loc.get("address", "")) + "|" +
                       _normalize_city(loc.get("city", "")) + "|" +
                       loc.get("state", "").strip().upper())
            seen_addr_keys.add(addr_key)
            continue
        addr_key = (_normalize_addr(loc.get("address", "")) + "|" +
                   _normalize_city(loc.get("city", "")) + "|" +
                   loc.get("state", "").strip().upper())
        if addr_key not in seen_addr_keys and loc.get("address"):
            name = loc.get("description", "") or loc.get("corporate_entity", "")
            if not name or not name.strip():
                name = "Pending"
            master_locations.append({
                "name": name,
                "address": loc.get("address", ""),
                "city": loc.get("city", ""),
                "state": loc.get("state", ""),
                "tiv": 0,
                "on_property": _is_on_property(addr_key),
                "on_liability": _is_on_liability(addr_key),
            })
            seen_addr_keys.add(addr_key)
    
    # Third: GL schedule_of_classes locations not already in master list
    # This catches liability-only locations (e.g., LaPlace, vacant land) that aren't on SOV or extracted locations
    import re
    gl_seen_addrs = set()  # Deduplicate GL entries (multiple classes per location)
    
    for entry in gl_classes:
        if not isinstance(entry, dict):
            continue
        addr = entry.get("address", "")
        if not addr:
            continue
        # Skip vacant land entries from GL schedule
        classification = (entry.get("classification", "") or "").lower()
        if "vacant" in classification or ("land" in classification and "hotel" not in classification):
            continue
        # Parse address - may contain "Street, City, ST ZIP" or just street
        addr_norm = _normalize_addr(addr)
        if addr_norm in gl_seen_addrs:
            continue
        gl_seen_addrs.add(addr_norm)
        
        # Check if this address is already in the master list using fuzzy matching
        already_in = False
        for ml in master_locations:
            ml_addr_norm = _normalize_addr(ml.get("address", ""))
            if _fuzzy_addr_match(addr_norm, ml_addr_norm):
                # Mark existing location as on_liability if not already
                ml["on_liability"] = True
                already_in = True
                break
        
        if not already_in:
            # Try to parse city/state from the address string (e.g., "4285 Highway 51, LaPlace, LA 70068")
            parts = [p.strip() for p in addr.split(",")]
            street = parts[0] if parts else addr
            city = ""
            state = ""
            if len(parts) >= 3:
                street = parts[0]
                city = parts[1]
                st_match = re.match(r'([A-Z]{2})\s*\d*', parts[2].strip())
                if st_match:
                    state = st_match.group(1)
            elif len(parts) == 2:
                street = parts[0]
                st_match = re.match(r'([A-Z]{2})\s*\d*', parts[1].strip())
                if st_match:
                    state = st_match.group(1)
                else:
                    city = parts[1]
            
            brand = entry.get("brand_dba", "") or entry.get("classification", "")
            if not brand or not brand.strip():
                brand = "Pending"
            
            # Check fuzzy match against seen_addr_keys too
            addr_key = (_normalize_addr(street) + "|" +
                       _normalize_city(city) + "|" +
                       state.strip().upper())
            key_already_seen = False
            for existing_key in seen_addr_keys:
                existing_parts = existing_key.split("|")
                new_parts = addr_key.split("|")
                if len(existing_parts) == 3 and len(new_parts) == 3:
                    # Relax city matching: if street addresses match, consider same location
                    # (city names can differ: "LA PLACE" vs "LAPLACE")
                    state_match = (not existing_parts[2] or not new_parts[2] or existing_parts[2] == new_parts[2])
                    if _fuzzy_addr_match(existing_parts[0], new_parts[0]) and state_match:
                        key_already_seen = True
                        break
            
            if not key_already_seen:
                # Cross-reference SOV for property name only (NOT TIV — TIV is property-only)
                if sov_data and sov_data.get("locations"):
                    for sov_loc in sov_data["locations"]:
                        if _fuzzy_addr_match(_normalize_addr(street), _normalize_addr(sov_loc.get("address", ""))):
                            # Get name from SOV if brand is generic
                            if brand in ("Pending", "") or brand == entry.get("classification", ""):
                                _cn = (sov_loc.get("corporate_name", "") or "").strip()
                                _db = (sov_loc.get("dba", "") or sov_loc.get("hotel_flag", "") or "").strip()
                                if _cn and _db:
                                    brand = f"{_cn} - {_db}"
                                elif _db:
                                    brand = _db
                                elif _cn:
                                    brand = _cn
                            break
                master_locations.append({
                    "name": brand,
                    "address": street,
                    "city": city,
                    "state": state,
                    "tiv": 0,  # TIV only comes from property SOV, never from GL
                    "on_property": _is_on_property(addr_key),
                    "on_liability": True,
                })
                seen_addr_keys.add(addr_key)
    
    # Fourth: Extract locations from GL forms_endorsements (e.g., CG 21 44 designated premises)
    # These forms often list ALL covered locations with full addresses
    gl_forms = gl_cov.get("forms_endorsements", []) if isinstance(gl_cov, dict) else []
    for form in gl_forms:
        if not isinstance(form, dict):
            continue
        desc = (form.get("description", "") or "").upper()
        # Look for designated premises forms that contain address lists
        if not any(kw in desc for kw in ["DESIGNATED PREMISES", "CG 21 44", "CG2144", "NXLL110", "NXLL 110", "LIMITATION OF COVERAGE"]):
            continue
        # Try to extract addresses from the description text
        # Format: "1) 4285 Highway 51, LaPlace, LA 70068" or similar numbered lists
        addr_pattern = re.findall(r'\d+\)\s*(.+?)(?=\s*\d+\)|$)', desc, re.DOTALL)
        if not addr_pattern:
            # Try semicolon or newline separated
            addr_pattern = [a.strip() for a in re.split(r'[;\n]', desc) if re.search(r'\d+\s+\w+', a.strip())]
        for raw_addr in addr_pattern:
            raw_addr = raw_addr.strip().rstrip(',')
            if not raw_addr or len(raw_addr) < 5:
                continue
            addr_norm = _normalize_addr(raw_addr)
            # Check if already in master list
            already_exists = False
            for ml in master_locations:
                if _fuzzy_addr_match(addr_norm, _normalize_addr(ml.get("address", ""))):
                    ml["on_liability"] = True
                    already_exists = True
                    break
            if not already_exists:
                parts = [p.strip() for p in raw_addr.split(",")]
                street = parts[0] if parts else raw_addr
                city = ""
                state = ""
                if len(parts) >= 3:
                    city = parts[1]
                    st_m = re.match(r'([A-Z]{2})\s*\d*', parts[2].strip().upper())
                    if st_m: state = st_m.group(1)
                elif len(parts) == 2:
                    st_m = re.match(r'([A-Z]{2})\s*\d*', parts[1].strip().upper())
                    if st_m:
                        state = st_m.group(1)
                    else:
                        city = parts[1]
                # Try to find name from SOV data
                loc_name = "Pending"
                if sov_data and sov_data.get("locations"):
                    for sov_loc in sov_data["locations"]:
                        if _fuzzy_addr_match(_normalize_addr(street), _normalize_addr(sov_loc.get("address", ""))):
                            # Use Corporate Name - DBA format
                            _cn = (sov_loc.get("corporate_name", "") or "").strip()
                            _db = (sov_loc.get("dba", "") or sov_loc.get("hotel_flag", "") or "").strip()
                            if _cn and _db:
                                loc_name = f"{_cn} - {_db}"
                            elif _db:
                                loc_name = _db
                            elif _cn:
                                loc_name = _cn
                            if not city and sov_loc.get("city"): city = sov_loc["city"]
                            if not state and sov_loc.get("state"): state = sov_loc["state"]
                            break
                addr_key = (_normalize_addr(street) + "|" + _normalize_city(city) + "|" + state.strip().upper())
                key_already_seen = False
                for existing_key in seen_addr_keys:
                    ep = existing_key.split("|")
                    np = addr_key.split("|")
                    if len(ep) == 3 and len(np) == 3:
                        if _fuzzy_addr_match(ep[0], np[0]) and ep[2] == np[2]:
                            key_already_seen = True
                            break
                if not key_already_seen:
                    master_locations.append({
                        "name": loc_name,
                        "address": street,
                        "city": city,
                        "state": state,
                        "tiv": 0,
                        "on_property": _is_on_property(addr_key),
                        "on_liability": True,
                    })
                    seen_addr_keys.add(addr_key)
    
    # Fifth: Add GL designated_premises entries not yet in master list
    # These are the primary GL location list — may have addresses the other passes missed
    for dp_addr in designated_premises:
        if not isinstance(dp_addr, str) or not dp_addr.strip() or len(dp_addr.strip()) < 5:
            continue
        dp_addr = dp_addr.strip()
        dp_addr_norm = _normalize_addr(dp_addr)
        # Check if already in master list
        already_exists = False
        for ml in master_locations:
            if _fuzzy_addr_match(dp_addr_norm, _normalize_addr(ml.get("address", ""))):
                ml["on_liability"] = True  # Ensure liability checkmark
                already_exists = True
                break
        if not already_exists:
            parts = [p.strip() for p in dp_addr.split(",")]
            street = parts[0] if parts else dp_addr
            city = ""
            state = ""
            if len(parts) >= 3:
                city = parts[1]
                st_m = re.match(r'([A-Z]{2})\s*\d*', parts[2].strip().upper())
                if st_m: state = st_m.group(1)
            elif len(parts) == 2:
                st_m = re.match(r'([A-Z]{2})\s*\d*', parts[1].strip().upper())
                if st_m:
                    state = st_m.group(1)
                else:
                    city = parts[1]
            # Cross-reference SOV for name
            loc_name = "Pending"
            if sov_data and sov_data.get("locations"):
                for sov_loc in sov_data["locations"]:
                    if _fuzzy_addr_match(_normalize_addr(street), _normalize_addr(sov_loc.get("address", ""))):
                        _cn = (sov_loc.get("corporate_name", "") or "").strip()
                        _db = (sov_loc.get("dba", "") or sov_loc.get("hotel_flag", "") or "").strip()
                        if _cn and _db:
                            loc_name = f"{_cn} - {_db}"
                        elif _db:
                            loc_name = _db
                        elif _cn:
                            loc_name = _cn
                        if not city and sov_loc.get("city"): city = sov_loc["city"]
                        if not state and sov_loc.get("state"): state = sov_loc["state"]
                        break
            addr_key = (_normalize_addr(street) + "|" + _normalize_city(city) + "|" + state.strip().upper())
            key_already = False
            for ek in seen_addr_keys:
                ep = ek.split("|")
                np = addr_key.split("|")
                if len(ep) == 3 and len(np) == 3 and _fuzzy_addr_match(ep[0], np[0]):
                    state_ok = (not ep[2] or not np[2] or ep[2] == np[2])
                    if state_ok:
                        key_already = True
                        break
            if not key_already:
                master_locations.append({
                    "name": loc_name,
                    "address": street,
                    "city": city,
                    "state": state,
                    "tiv": 0,
                    "on_property": _is_on_property(addr_key),
                    "on_liability": True,
                })
                seen_addr_keys.add(addr_key)

    # Normalize ALL CAPS text in location data to title case
    for ml in master_locations:
        if ml["name"] and ml["name"] == ml["name"].upper() and len(ml["name"]) > 2:
            ml["name"] = _proper_case(ml["name"])
        if ml["address"] and ml["address"] == ml["address"].upper():
            ml["address"] = _proper_case(ml["address"])
        if ml["city"] and ml["city"] == ml["city"].upper() and len(ml["city"]) > 2:
            ml["city"] = _proper_case(ml["city"])
    
    if master_locations:
        CHECK = "\u2713"  # Unicode checkmark
        DASH = "\u2014"   # Em-dash for missing coverage (rendered in RED)
        headers = ["#", "Property Name", "Address", "City", "ST", "TIV", "Property", "Liability"]
        rows = []
        total_tiv = 0
        # Track which rows have missing coverages for RED formatting
        missing_property_rows = []  # row indices (0-based in rows list)
        missing_liability_rows = []
        
        for i, loc in enumerate(master_locations, 1):
            tiv_val = loc["tiv"] or 0
            total_tiv += tiv_val
            prop_cell = CHECK if loc["on_property"] else DASH
            liab_cell = CHECK if loc["on_liability"] else DASH
            if not loc["on_property"]:
                missing_property_rows.append(len(rows))  # current row index
            if not loc["on_liability"]:
                missing_liability_rows.append(len(rows))
            rows.append([
                str(i),
                loc["name"],
                loc["address"],
                loc["city"],
                loc["state"],
                fmt_currency(tiv_val) if tiv_val else "",
                prop_cell,
                liab_cell,
            ])
        
        # Add totals row
        rows.append([
            "", "TOTAL", "", "", "",
            fmt_currency(total_tiv) if total_tiv else "",
            "", ""
        ])
        
        L = WD_ALIGN_PARAGRAPH.LEFT
        C = WD_ALIGN_PARAGRAPH.CENTER
        R = WD_ALIGN_PARAGRAPH.RIGHT
        table = create_styled_table(doc, headers, rows,
                          col_widths=[0.3, 1.5, 1.6, 0.9, 0.3, 1.0, 0.7, 0.7],
                          header_size=8, body_size=8,
                          header_alignments={0: L, 1: L, 2: L, 3: L, 4: L, 5: R, 6: C, 7: C},
                          col_alignments={5: R, 6: C, 7: C})
        
        # Apply RED color to missing coverage cells (em-dashes)
        RED = RGBColor(0xCC, 0x00, 0x00)
        for row_idx in missing_property_rows:
            cell = table.rows[row_idx + 1].cells[6]  # +1 for header row
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RED
                    run.font.bold = True
        for row_idx in missing_liability_rows:
            cell = table.rows[row_idx + 1].cells[7]  # +1 for header row
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RED
                    run.font.bold = True
        
        # Legend
        add_formatted_paragraph(doc, "", size=4)
        legend_p = doc.add_paragraph()
        legend_p.paragraph_format.space_before = Pt(2)
        legend_p.paragraph_format.space_after = Pt(2)
        run_check = legend_p.add_run("\u2713")
        run_check.font.size = Pt(8)
        run_check.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Green
        run_text = legend_p.add_run(" = Covered     ")
        run_text.font.size = Pt(8)
        run_text.font.color.rgb = CHARCOAL
        run_dash = legend_p.add_run(DASH)
        run_dash.font.size = Pt(8)
        run_dash.font.color.rgb = RED
        run_dash.font.bold = True
        run_text2 = legend_p.add_run(" = Not Currently Quoted")
        run_text2.font.size = Pt(8)
        run_text2.font.color.rgb = CHARCOAL
        
        # Add note about SOV
        if sov_data and sov_data.get("locations"):
            add_formatted_paragraph(doc, "", size=6)
            add_formatted_paragraph(doc, "See attached Statement of Values for complete property details.",
                                  size=9, italic=True, color=CHARCOAL)
    else:
        add_formatted_paragraph(doc, "Location schedule to be confirmed.", size=11)


def generate_coverage_section(doc, data, coverage_key, display_name):
    """Generate a coverage section (Property, GL, Umbrella, WC, Auto)."""
    # Standard crime insuring clause names (fallback when GPT extraction is incomplete)
    STANDARD_CRIME_CLAUSES = [
        "Employee Theft",
        "Forgery or Alteration",
        "Inside the Premises - Theft of Money and Securities",
        "Inside the Premises - Robbery or Safe Burglary of Other Property",
        "Outside the Premises",
        "Computer and Funds Transfer Fraud",
        "Money Orders and Counterfeit Money",
        "Client's Property",
        "Identity Fraud Expense",
        "Social Engineering Fraud",
        "Impersonation Fraud",
        "Erroneous Payments",
        "Investigation Expenses",
    ]

    coverages = data.get("coverages", {})
    cov = coverages.get(coverage_key)
    if not cov:
        return
    
    # Skip phantom coverage sections with no meaningful data
    carrier = cov.get("carrier", "") or ""
    premium = cov.get("premium", 0) or 0
    total_premium = cov.get("total_premium", 0) or 0
    limits = cov.get("coverage_limits", []) or []
    if not carrier.strip() and not premium and not total_premium and not limits:
        logger.info(f"Skipping phantom coverage section: {coverage_key} (no carrier, premium, or limits)")
        return
    
    add_page_break(doc)
    add_section_header(doc, display_name)
    
    # Coverage Summary table
    carrier = _clean_carrier_name(cov.get("carrier", "N/A"))
    admitted = "Admitted" if cov.get("carrier_admitted", True) else "Non-Admitted"
    am_best = cov.get("am_best_rating", "N/A")
    # Fallback to lookup table if not provided in quote
    if not am_best or am_best == "N/A":
        looked_up = lookup_am_best(carrier)
        if looked_up:
            am_best = looked_up
    
    add_subsection_header(doc, "Coverage Summary")
    
    carrier_rows = [
        ["Carrier", carrier],
        ["Admitted Status", admitted],
        ["AM Best Rating", am_best],
    ]
    
    # Add wholesaler if present
    if cov.get("wholesaler"):
        carrier_rows.append(["Wholesaler", cov["wholesaler"]])
    
    # Add policy form if present
    if cov.get("policy_form"):
        carrier_rows.append(["Policy Form", cov["policy_form"]])
    
    # Add policy period if present
    if cov.get("policy_period"):
        carrier_rows.append(["Policy Period", cov["policy_period"]])
    
    # Add layer description for excess property
    layer_desc = cov.get("layer_description", "")
    if layer_desc and coverage_key in ("excess_property", "excess_property_2"):
        carrier_rows.append(["Layer", layer_desc])
    
    # Add TIV if present (for property coverages)
    tiv = cov.get("tiv", "")
    if tiv and coverage_key in ("property", "excess_property", "excess_property_2"):
        carrier_rows.append(["Total Insured Value", tiv])
    
    # Add GL deductible if present
    gl_ded = cov.get("gl_deductible", "")
    if gl_ded and gl_ded not in ("$0", "None", "N/A", "", "0"):
        carrier_rows.append(["Deductible", gl_ded])
    
    # Add defense basis if present
    defense = cov.get("defense_basis", "")
    if defense and defense not in ("N/A", ""):
        carrier_rows.append(["Defense Basis", defense])
    
    L = WD_ALIGN_PARAGRAPH.LEFT
    create_styled_table(doc, ["Item", "Details"], carrier_rows, col_widths=[2.5, 5.0],
                       header_size=10, body_size=10,
                       header_alignments={0: L, 1: L})
    
    # Schedule of Values (Property) - prefer SOV data if available
    sov_data = data.get("sov_data")
    sov_from_quote = cov.get("schedule_of_values", [])
    
    if coverage_key == "property" and sov_data and sov_data.get("locations"):
        # Use SOV spreadsheet data for detailed Schedule of Values
        add_subsection_header(doc, "Schedule of Values")
        # Prefer property SOV building-level data if available (from merged SOV)
        if sov_data.get("_property_sov") and sov_data["_property_sov"].get("locations"):
            _prop_sov = sov_data["_property_sov"]
            sov_locs = _prop_sov["locations"]
            totals = _prop_sov.get("totals", sov_data.get("totals", {}))
        else:
            sov_locs = sov_data["locations"]
            # Filter to only locations with TIV > 0 for the schedule of values
            sov_locs_with_tiv = [loc for loc in sov_locs if (loc.get("tiv", 0) or 0) > 0]
            if sov_locs_with_tiv:
                sov_locs = sov_locs_with_tiv
            totals = sov_data.get("totals", {})
        # Check if any location has "other_value" data (Sign, Pool, Other)
        has_other = any(loc.get("other_value", 0) for loc in sov_locs)
        if has_other:
            headers = ["#", "Location", "Building", "Contents", "Other", "BI/Rents", "TIV"]
        else:
            headers = ["#", "Location", "Building", "Contents", "BI/Rents", "TIV"]
        rows = []
        total_other = 0
        for i, loc in enumerate(sov_locs, 1):
            name = loc.get("dba") or loc.get("hotel_flag") or loc.get("corporate_name", "")
            if name and name == name.upper() and len(name) > 2:
                name = _proper_case(name)
            _addr = loc.get('address', '')
            _city = loc.get('city', '')
            if _addr and _addr == _addr.upper(): _addr = _proper_case(_addr)
            if _city and _city == _city.upper() and len(_city) > 2: _city = _proper_case(_city)
            addr = f"{_addr}, {_city}, {loc.get('state', '')}"
            loc_label = f"{name}\n{addr}" if name else addr
            other_val = loc.get("other_value", 0) or 0
            total_other += other_val
            if has_other:
                rows.append([
                    str(i),
                    loc_label,
                    fmt_currency(loc.get("building_value", 0)),
                    fmt_currency(loc.get("contents_value", 0)),
                    fmt_currency(other_val) if other_val else "\u2014",
                    fmt_currency(loc.get("bi_value", 0)),
                    fmt_currency(loc.get("tiv", 0))
                ])
            else:
                rows.append([
                    str(i),
                    loc_label,
                    fmt_currency(loc.get("building_value", 0)),
                    fmt_currency(loc.get("contents_value", 0)),
                    fmt_currency(loc.get("bi_value", 0)),
                    fmt_currency(loc.get("tiv", 0))
                ])
        # Add totals row
        if has_other:
            rows.append([
                "", "TOTAL",
                fmt_currency(totals.get("building_value", 0)),
                fmt_currency(totals.get("contents_value", 0)),
                fmt_currency(total_other) if total_other else "",
                fmt_currency(totals.get("bi_value", 0)),
                fmt_currency(totals.get("tiv", 0))
            ])
            create_styled_table(doc, headers, rows,
                              col_widths=[0.3, 2.0, 1.1, 0.9, 0.8, 0.9, 1.1],
                              header_size=9, body_size=8,
                              col_alignments={2: WD_ALIGN_PARAGRAPH.CENTER, 3: WD_ALIGN_PARAGRAPH.CENTER,
                                             4: WD_ALIGN_PARAGRAPH.CENTER, 5: WD_ALIGN_PARAGRAPH.CENTER,
                                             6: WD_ALIGN_PARAGRAPH.CENTER})
        else:
            rows.append([
                "", "TOTAL",
                fmt_currency(totals.get("building_value", 0)),
                fmt_currency(totals.get("contents_value", 0)),
                fmt_currency(totals.get("bi_value", 0)),
                fmt_currency(totals.get("tiv", 0))
            ])
            create_styled_table(doc, headers, rows,
                              col_widths=[0.3, 2.2, 1.2, 1.0, 1.0, 1.3],
                              header_size=9, body_size=8,
                              col_alignments={2: WD_ALIGN_PARAGRAPH.CENTER, 3: WD_ALIGN_PARAGRAPH.CENTER,
                                             4: WD_ALIGN_PARAGRAPH.CENTER, 5: WD_ALIGN_PARAGRAPH.CENTER})
    elif sov_from_quote:
        add_subsection_header(doc, "Schedule of Values")
        headers = ["Location", "Building", "Contents", "BI/Rents", "TIV"]
        rows = [[
            s.get("location", ""),
            fmt_currency(s.get("building", 0)),
            fmt_currency(s.get("contents", 0)),
            fmt_currency(s.get("business_income", 0)),
            fmt_currency(s.get("tiv", 0))
        ] for s in sov_from_quote]
        create_styled_table(doc, headers, rows,
                          col_widths=[2.0, 1.4, 1.2, 1.2, 1.2],
                          header_size=9, body_size=9)
    
    # Crime Insuring Clauses (3-column: Clause, Limit, Retention)
    insuring_clauses = cov.get("insuring_clauses", []) or cov.get("insuring_agreements", [])
    # For Crime: if no insuring_clauses but limits contain descriptions+retentions, use as clauses
    if coverage_key == "crime" and not insuring_clauses:
        _crime_limits = cov.get("limits", [])
        if _crime_limits and any(
            isinstance(lim, dict) and (lim.get("description") or lim.get("type"))
            and (lim.get("retention") or lim.get("deductible"))
            for lim in _crime_limits
        ):
            insuring_clauses = _crime_limits
    if coverage_key == "crime" and insuring_clauses:
        add_subsection_header(doc, "Insuring Clauses")
        headers = ["Insuring Clause", "Limit", "Retention"]
        rows = []
        clause_idx = 0
        for clause in insuring_clauses:
            if isinstance(clause, dict):
                # Use extracted description if available, otherwise use standard clause name
                clause_name = clause.get("description", "").strip()
                if not clause_name and clause_idx < len(STANDARD_CRIME_CLAUSES):
                    clause_name = STANDARD_CRIME_CLAUSES[clause_idx]
                rows.append([
                    clause_name,
                    clause.get("limit", ""),
                    clause.get("retention", clause.get("deductible", ""))
                ])
            else:
                # Use standard clause name as fallback
                clause_name = str(clause).strip() if str(clause).strip() else (STANDARD_CRIME_CLAUSES[clause_idx] if clause_idx < len(STANDARD_CRIME_CLAUSES) else "")
                rows.append([clause_name, "", ""])
            clause_idx += 1
        L = WD_ALIGN_PARAGRAPH.LEFT
        R = WD_ALIGN_PARAGRAPH.RIGHT
        create_styled_table(doc, headers, rows, col_widths=[4.0, 1.5, 1.5],
                           header_size=10, body_size=9,
                           header_alignments={0: L, 1: R, 2: R},
                           col_alignments={1: R, 2: R})
    
    # Limits (non-crime coverages, or crime without insuring_clauses)
    limits = cov.get("limits", [])
    if limits and not (coverage_key == "crime" and insuring_clauses):
        # For Crime without insuring_clauses, show limits with retentions if available
        has_retention = coverage_key == "crime" and any(
            isinstance(lim, dict) and (lim.get("retention") or lim.get("deductible"))
            for lim in limits
        )
        if has_retention:
            add_subsection_header(doc, "Insuring Agreements / Limits of Insurance")
            headers = ["Insuring Agreement", "Limit", "Retention"]
            rows = []
            for lim in limits:
                if isinstance(lim, dict):
                    rows.append([
                        lim.get("description", "") or lim.get("type", ""),
                        lim.get("limit", "") or lim.get("amount", ""),
                        lim.get("retention", "") or lim.get("deductible", "")
                    ])
                else:
                    rows.append([str(lim), "", ""])
            L = WD_ALIGN_PARAGRAPH.LEFT
            R = WD_ALIGN_PARAGRAPH.RIGHT
            create_styled_table(doc, headers, rows, col_widths=[4.0, 1.5, 1.5],
                               header_size=10, body_size=9,
                               header_alignments={0: L, 1: R, 2: R},
                               col_alignments={1: R, 2: R})
        else:
            add_subsection_header(doc, "Coverage Limits")
            headers = ["Description", "Limit"]
            rows = [[lim.get("description", "") or lim.get("type", ""), lim.get("limit", "") or lim.get("amount", "")] if isinstance(lim, dict) else [str(lim), ""] for lim in limits]
            # Left-align headers; center Limit values for umbrella/excess
            L = WD_ALIGN_PARAGRAPH.LEFT
            limit_body_align = {}
            if coverage_key in ("umbrella", "cyber", "epli", "flood", "terrorism"):
                limit_body_align = {1: WD_ALIGN_PARAGRAPH.CENTER}
            create_styled_table(doc, headers, rows, col_widths=[4.5, 3.0],
                               header_size=10, body_size=10,
                               header_alignments={0: L, 1: L},
                               col_alignments=limit_body_align)
    
    # Deductibles (Property)
    deductibles = cov.get("deductibles", [])
    if deductibles:
        add_subsection_header(doc, "Deductibles")
        headers = ["Peril", "Deductible"]
        rows = [[ded.get("description", "") or ded.get("type", ""), ded.get("amount", "")] if isinstance(ded, dict) else [str(ded), ""] for ded in deductibles]
        L = WD_ALIGN_PARAGRAPH.LEFT
        create_styled_table(doc, headers, rows, col_widths=[4.5, 3.0],
                           header_size=10, body_size=10,
                           header_alignments={0: L, 1: L})
    
    # Coinsurance & Valuation (Property)
    coinsurance = cov.get("coinsurance", [])
    valuation = cov.get("valuation", "")
    if coinsurance or valuation:
        add_subsection_header(doc, "Coinsurance & Valuation")
        headers = ["Coverage", "Coinsurance / Limitation"]
        rows = []
        for ci in coinsurance:
            if isinstance(ci, dict):
                cov_name = ci.get("coverage", "")
                pct = ci.get("percentage", "")
                limitation = ci.get("limitation", "")
                val = limitation if limitation else pct
                if val:
                    rows.append([cov_name, val])
        if valuation:
            rows.append(["Valuation", valuation])
        if rows:
            L = WD_ALIGN_PARAGRAPH.LEFT
            create_styled_table(doc, headers, rows, col_widths=[4.5, 3.0],
                               header_size=10, body_size=10,
                               header_alignments={0: L, 1: L})
    
    # Layer Description (Excess Property)
    layer_desc = cov.get("layer_description", "")
    if layer_desc and coverage_key in ("excess_property", "excess_property_2"):
        add_formatted_paragraph(doc, f"Layer: {layer_desc}", size=11, bold=True, space_after=6)
    
    # Schedule of Hazards (GL)
    hazards = cov.get("schedule_of_hazards", [])
    if hazards:
        add_subsection_header(doc, "Schedule of Hazards")
        headers = ["Location", "Classification", "Code", "Basis", "Exposure"]
        rows = [[
            h.get("location", ""),
            h.get("classification", ""),
            h.get("code", ""),
            h.get("basis", ""),
            h.get("exposure", "")
        ] if isinstance(h, dict) else [str(h), "", "", "", ""] for h in hazards]
        create_styled_table(doc, headers, rows,
                          col_widths=[1.5, 2.5, 0.8, 1.0, 1.2],
                          header_size=9, body_size=9)
    
    # Schedule of Classes (GL - location exposures)
    classes = cov.get("schedule_of_classes", [])
    if classes:
        add_subsection_header(doc, "Exposures")
        from docx.enum.text import WD_ALIGN_PARAGRAPH as WD_ALIGN
        # Check if we have address/brand data for the enhanced format
        has_address = any(c.get("address") or c.get("brand_dba") for c in classes if isinstance(c, dict))
        # Check if we have class codes and rates (class-code-based format like AmTrust)
        has_class_code = any(c.get("class_code") for c in classes if isinstance(c, dict))
        has_rate = any(c.get("rate") for c in classes if isinstance(c, dict))
        
        if has_address:
            # Enhanced format with address/brand: Address | Brand/DBA | Code | Classification | Rate | Exposure
            if has_class_code or has_rate:
                headers = ["Address", "Brand / DBA", "Code", "Classification", "Rate", "Exposure"]
                rows = []
                for c in classes:
                    if not isinstance(c, dict):
                        continue
                    addr = c.get("address", "")
                    brand = c.get("brand_dba", "")
                    if not addr and c.get("location"):
                        addr = c.get("location", "")
                    classification = c.get("classification", "")
                    class_code = c.get("class_code", "")
                    rate = c.get("rate", "")
                    exposure_basis = c.get("exposure_basis", "")
                    exposure = c.get("exposure", "")
                    # Format exposure with basis
                    if exposure_basis and exposure:
                        exposure_str = f"{exposure} ({exposure_basis})"
                    else:
                        exposure_str = str(exposure) if exposure else ""
                    rows.append([addr, brand, str(class_code), classification, str(rate), exposure_str])
                create_styled_table(doc, headers, rows,
                                  col_widths=[1.8, 1.2, 0.6, 1.5, 0.6, 1.3],
                                  header_size=8, body_size=7,
                                  col_alignments={4: WD_ALIGN.RIGHT, 5: WD_ALIGN.RIGHT})
            else:
                # Address-based without class codes: Address | Brand/DBA | Classification | Exposure | Premium
                headers = ["Address", "Brand / DBA", "Classification", "Exposure", "Premium"]
                rows = []
                for c in classes:
                    if not isinstance(c, dict):
                        continue
                    addr = c.get("address", "")
                    brand = c.get("brand_dba", "")
                    if not addr and c.get("location"):
                        addr = c.get("location", "")
                    classification = c.get("classification", "")
                    exposure_basis = c.get("exposure_basis", "")
                    exposure = c.get("exposure", "")
                    premium = c.get("premium", "")
                    if exposure_basis and exposure:
                        exposure_str = f"{exposure} ({exposure_basis})"
                    else:
                        exposure_str = str(exposure) if exposure else ""
                    rows.append([addr, brand, classification, exposure_str, str(premium) if premium else ""])
                create_styled_table(doc, headers, rows,
                                  col_widths=[2.0, 1.5, 1.5, 1.3, 1.0],
                                  header_size=8, body_size=8,
                                  col_alignments={3: WD_ALIGN.RIGHT, 4: WD_ALIGN.RIGHT})
        else:
            # Class-code-based format (no addresses): Code | Classification | Rate | Exposure | Basis
            headers = ["Code", "Classification", "Rate", "Exposure", "Exposure Basis"]
            rows = []
            for c in classes:
                if not isinstance(c, dict):
                    continue
                class_code = c.get("class_code", c.get("location", ""))
                classification = c.get("classification", "")
                rate = c.get("rate", "")
                exposure = c.get("exposure", "")
                exposure_basis = c.get("exposure_basis", "")
                rows.append([str(class_code), classification, str(rate), str(exposure), exposure_basis])
            create_styled_table(doc, headers, rows,
                              col_widths=[0.8, 2.5, 0.8, 1.5, 1.4],
                              header_size=9, body_size=8,
                              col_alignments={2: WD_ALIGN.RIGHT, 3: WD_ALIGN.RIGHT})
    
    # Rating Basis (WC)
    rating = cov.get("rating_basis", [])
    if rating:
        add_subsection_header(doc, "Rating Basis")
        headers = ["State", "Location", "Class Code", "Classification", "Payroll", "Rate"]
        rows = [[
            r.get("state", ""),
            r.get("location", ""),
            r.get("class_code", ""),
            r.get("classification", ""),
            r.get("payroll", ""),
            r.get("rate", "")
        ] for r in rating]
        create_styled_table(doc, headers, rows,
                          col_widths=[0.6, 1.5, 0.8, 2.0, 1.2, 0.9],
                          header_size=9, body_size=9)
    
    # Vehicle Schedule (Auto)
    vehicles = cov.get("vehicle_schedule", [])
    if vehicles:
        add_subsection_header(doc, "Vehicle Schedule")
        headers = ["Year", "Make", "Model", "VIN", "Garage Location"]
        rows = [[v.get("year", ""), v.get("make", ""), v.get("model", ""),
                 v.get("vin", ""), v.get("garage_location", "")] for v in vehicles]
        create_styled_table(doc, headers, rows,
                          col_widths=[0.6, 1.2, 1.2, 2.5, 2.0],
                          header_size=9, body_size=9)
    
    # Additional Coverages
    addl = cov.get("additional_coverages", [])
    if addl:
        # Use "Sublimits of Liability / Extensions" for property, "Additional Coverages" for others
        addl_title = "Sublimits of Liability / Extensions" if coverage_key == "property" else "Additional Coverages"
        add_subsection_header(doc, addl_title)
        has_ded = any(ac.get("deductible") for ac in addl if isinstance(ac, dict))
        L = WD_ALIGN_PARAGRAPH.LEFT
        if has_ded:
            headers = ["Description", "Limit", "Deductible"]
            rows = [[ac.get("description", "") or ac.get("coverage", "") or ac.get("name", ""), ac.get("limit", ""), ac.get("deductible", "")] if isinstance(ac, dict) else [str(ac), "", ""] for ac in addl]
            create_styled_table(doc, headers, rows, col_widths=[3.5, 2.0, 2.0],
                              header_alignments={0: L, 1: L, 2: L})
        else:
            headers = ["Description", "Limit"]
            rows = [[ac.get("description", "") or ac.get("coverage", "") or ac.get("name", ""), ac.get("limit", "")] if isinstance(ac, dict) else [str(ac), ""] for ac in addl]
            create_styled_table(doc, headers, rows, col_widths=[4.5, 3.0],
                              header_alignments={0: L, 1: L})
    
    # Underlying Insurance (Umbrella)
    underlying = cov.get("underlying_insurance", [])
    if underlying:
        add_subsection_header(doc, "Underlying Insurance")
        headers = ["Carrier", "Coverage", "Limits"]
        rows = [[u.get("carrier", ""), u.get("coverage", ""), u.get("limits", "")] for u in underlying]
        create_styled_table(doc, headers, rows, col_widths=[2.5, 2.5, 2.5],
                          col_alignments={2: WD_ALIGN_PARAGRAPH.CENTER})
    
    # Tower Structure (Umbrella)
    tower = cov.get("tower_structure", [])
    if tower:
        add_subsection_header(doc, "Umbrella Tower Structure")
        headers = ["Layer", "Carrier", "Limits", "Premium", "Total (incl. taxes/fees)"]
        rows = [[
            t.get("layer", ""),
            t.get("carrier", ""),
            t.get("limits", ""),
            fmt_currency(t.get("premium", 0)),
            fmt_currency(t.get("total_cost", 0))
        ] for t in tower]
        create_styled_table(doc, headers, rows,
                          col_widths=[0.8, 2.0, 1.5, 1.2, 1.5],
                          header_size=9, body_size=9)
    
    # Forms & Endorsements
    forms = cov.get("forms_endorsements", [])
    # Critical coverages MUST show forms — add placeholder if empty
    _critical_form_coverages = {"general_liability", "crime", "umbrella", "umbrella_layer_2",
                                 "umbrella_layer_3", "umbrella_layer_4", "workers_comp",
                                 "workers_compensation", "commercial_auto"}
    if forms:
        add_subsection_header(doc, "Forms & Endorsements")
        headers = ["Form Number", "Description"]
        rows = [[f.get("form_number", ""), f.get("description", "")] if isinstance(f, dict) else ["", str(f)] for f in forms]
        L = WD_ALIGN_PARAGRAPH.LEFT
        create_styled_table(doc, headers, rows, col_widths=[2.0, 5.5],
                           header_size=9, body_size=9,
                           header_alignments={0: L, 1: L})
    elif coverage_key in _critical_form_coverages:
        add_subsection_header(doc, "Forms & Endorsements")
        add_formatted_paragraph(doc,
            "Forms and endorsements schedule to be confirmed upon policy issuance. "
            "Please refer to the carrier quote documents for the complete list of applicable forms.",
            size=9, italic=True, color=CHARCOAL, space_after=6)
    
    # Additional Named Insureds (per-coverage)
    # Pull from the coverage data and cross-reference with SOV for hotel matching
    ani_list = cov.get("additional_named_insureds", []) or []
    # Also check additional_insureds (some extractors use this key)
    if not ani_list:
        ani_list = cov.get("additional_insureds", []) or []
    if ani_list:
        add_subsection_header(doc, "Additional Named Insureds")
        # Cross-reference with SOV data to match insureds to hotels
        sov_locs = (data.get("sov_data") or {}).get("locations", []) if data.get("sov_data") else []
        headers_ani = ["#", "Additional Named Insured", "Location / Hotel"]
        rows_ani = []
        for idx_ani, ani in enumerate(ani_list, 1):
            if isinstance(ani, dict):
                ani_name = ani.get("name", "") or ani.get("insured", "") or ""
                ani_loc = ani.get("location", "") or ani.get("hotel", "") or ani.get("address", "") or ""
            else:
                ani_name = str(ani)
                ani_loc = ""
            # Try to match to SOV location if no location specified
            if not ani_loc and sov_locs:
                ani_name_lower = ani_name.lower()
                for sov_loc in sov_locs:
                    dba = (sov_loc.get("dba", "") or sov_loc.get("hotel_flag", "") or "").lower()
                    corp = (sov_loc.get("corporate_name", "") or "").lower()
                    if dba and dba in ani_name_lower:
                        ani_loc = f"{sov_loc.get('address', '')}, {sov_loc.get('city', '')}, {sov_loc.get('state', '')}"
                        break
                    elif corp and corp in ani_name_lower:
                        ani_loc = f"{sov_loc.get('address', '')}, {sov_loc.get('city', '')}, {sov_loc.get('state', '')}"
                        break
            rows_ani.append([str(idx_ani), ani_name, ani_loc])
        L = WD_ALIGN_PARAGRAPH.LEFT
        create_styled_table(doc, headers_ani, rows_ani, col_widths=[0.4, 3.5, 3.6],
                           header_size=9, body_size=8,
                           header_alignments={0: L, 1: L, 2: L})

    # Schedule of Locations reference (GL and Crime — critical E&O requirement)
    if coverage_key in ("general_liability", "crime"):
        # Count locations from GL schedule_of_classes or designated_premises
        _loc_count = 0
        if coverage_key == "general_liability":
            _gl_classes = cov.get("schedule_of_classes", []) or []
            _gl_premises = cov.get("designated_premises", []) or []
            _skip_classes = {"hired auto", "non-owned auto", "loss control", "package store",
                            "category vi", "sundry", "flat"}
            _gl_loc_ids = set()
            for _entry in _gl_classes:
                if isinstance(_entry, dict):
                    _loc_id = (_entry.get("location", "") or "").strip().lower()
                    _classif = (_entry.get("classification", "") or "").lower()
                    if any(_sk in _classif for _sk in _skip_classes):
                        continue
                    if _loc_id and _loc_id not in ("", "n/a", "all", "various"):
                        _gl_loc_ids.add(_loc_id)
            _loc_count = max(len(_gl_loc_ids), len(_gl_premises))
        if _loc_count > 0:
            add_formatted_paragraph(doc,
                f"Schedule of Locations: {_loc_count} locations are covered under this policy. "
                "See the Schedule of Locations section for the complete list with addresses.",
                size=9, italic=True, color=CHARCOAL, space_before=6, space_after=6)

    # Covered Locations (GL only) - backup list of liability locations from GL quote
    if coverage_key == "general_liability":
        import re as _re_gl
        gl_loc_list = []
        gl_loc_seen = set()
        def _gl_loc_already_seen(addr_norm):
            """Check if addr is already in gl_loc_seen using fuzzy matching.
            Handles city/state format differences between GL and SOV addresses."""
            if addr_norm in gl_loc_seen:
                return True
            for existing in gl_loc_seen:
                if _fuzzy_addr_match(addr_norm, existing):
                    return True
            # Compare house number + first 2 street words to handle format differences
            # e.g. "5370 CLEARWATER CT BEAUMONT TX" vs "5370 CLEARWATER CT BEAUMONT TEXAS 77705"
            import re as _re_dup
            m_new = _re_dup.match(r'^(\d+)\s+(.+)', addr_norm)
            if m_new:
                new_num = m_new.group(1)
                new_words = m_new.group(2).split()[:2]
                for existing in gl_loc_seen:
                    m_ex = _re_dup.match(r'^(\d+)\s+(.+)', existing)
                    if m_ex and m_ex.group(1) == new_num:
                        ex_words = m_ex.group(2).split()[:2]
                        if new_words == ex_words:
                            return True
            return False
        # Source 1: designated_premises array (PRIMARY - extracted by GPT)
        for dp_addr in cov.get("designated_premises", []):
            if not isinstance(dp_addr, str) or not dp_addr.strip() or len(dp_addr.strip()) < 5:
                continue
            addr_norm = _normalize_addr(dp_addr)
            if not _gl_loc_already_seen(addr_norm):
                gl_loc_seen.add(addr_norm)
                gl_loc_list.append({
                    "address": dp_addr.strip(),
                    "brand": "",
                    "classification": "",
                })
        # Source 2: schedule_of_classes addresses
        for entry in cov.get("schedule_of_classes", []):
            if isinstance(entry, dict):
                addr = entry.get("address", "") or entry.get("location", "")
                if addr and addr.strip():
                    addr_norm = _normalize_addr(addr)
                    if not _gl_loc_already_seen(addr_norm):
                        gl_loc_seen.add(addr_norm)
                        brand = entry.get("brand_dba", "") or ""
                        classification = entry.get("classification", "") or ""
                        gl_loc_list.append({
                            "address": addr.strip(),
                            "brand": brand.strip(),
                            "classification": classification.strip(),
                        })
        # Source 3: CG2144/NXLL110 designated premises forms (fallback)
        for form in cov.get("forms_endorsements", []):
            if not isinstance(form, dict):
                continue
            desc = (form.get("description", "") or "").upper()
            if not any(kw in desc for kw in ["DESIGNATED PREMISES", "CG 21 44", "CG2144", "NXLL110", "NXLL 110", "LIMITATION OF COVERAGE"]):
                continue
            addr_pattern = _re_gl.findall(r'\d+\)\s*(.+?)(?=\s*\d+\)|$)', desc, _re_gl.DOTALL)
            if not addr_pattern:
                addr_pattern = [a.strip() for a in _re_gl.split(r'[;\n]', desc) if _re_gl.search(r'\d+\s+\w+', a.strip())]
            for raw_addr in addr_pattern:
                raw_addr = raw_addr.strip().rstrip(',')
                if not raw_addr or len(raw_addr) < 5:
                    continue
                addr_norm = _normalize_addr(raw_addr)
                if not _gl_loc_already_seen(addr_norm):
                    gl_loc_seen.add(addr_norm)
                    gl_loc_list.append({
                        "address": raw_addr.strip(),
                        "brand": "",
                        "classification": "",
                    })
        
        # Source 4: If GL covered locations list is incomplete compared to SOV,
        # supplement with SOV locations (which have corporate names and addresses)
        sov_data = data.get("sov_data")
        _sov_loc_count = len(sov_data.get("locations", [])) if sov_data else 0
        if len(gl_loc_list) < max(3, _sov_loc_count):
            all_locs = data.get("locations", [])
            if sov_data and sov_data.get("locations"):
                for sov_loc in sov_data["locations"]:
                    addr = sov_loc.get("address", "")
                    if not addr or len(addr.strip()) < 5:
                        continue
                    addr_norm = _normalize_addr(addr)
                    if not _gl_loc_already_seen(addr_norm):
                        gl_loc_seen.add(addr_norm)
                        corp = (sov_loc.get("corporate_name") or sov_loc.get("client_name") or "").strip()
                        dba = (sov_loc.get("dba") or sov_loc.get("hotel_flag") or "").strip()
                        # Handle "Corp LLC dba Brand" in dba field
                        if not corp and dba and " dba " in dba.lower():
                            import re as _re_dba
                            parts = _re_dba.split(r'\s+dba\s+', dba, flags=_re_dba.IGNORECASE)
                            if len(parts) == 2:
                                corp = parts[0].strip()
                                dba = parts[1].strip()
                        brand = f"{corp} / {dba}" if corp and dba else (corp or dba or "")
                        gl_loc_list.append({
                            "address": addr.strip(),
                            "brand": brand,
                            "classification": "",
                        })
                logger.info(f"GL covered locations: added {len(gl_loc_list)} from SOV fallback")
            elif all_locs:
                for loc in all_locs:
                    addr = loc.get("address", "")
                    if not addr or len(addr.strip()) < 5:
                        continue
                    addr_norm = _normalize_addr(addr)
                    if not _gl_loc_already_seen(addr_norm):
                        gl_loc_seen.add(addr_norm)
                        name = loc.get("name", "") or loc.get("corporate_entity", "") or ""
                        gl_loc_list.append({
                            "address": addr.strip(),
                            "brand": name.strip(),
                            "classification": "",
                        })

        # Cross-reference GL locations with SOV to pull Corporate Name - DBA
        sov_data = data.get("sov_data")
        if sov_data and sov_data.get("locations"):
            for gl_loc in gl_loc_list:
                if gl_loc["brand"]:  # Already has a brand, skip
                    continue
                gl_addr_norm = _normalize_addr(gl_loc["address"])
                if not gl_addr_norm:
                    continue
                for sov_loc in sov_data["locations"]:
                    sov_addr_norm = _normalize_addr(sov_loc.get("address", ""))
                    if not sov_addr_norm:
                        continue
                    # Use fuzzy address matching (handles typos, abbreviations, house# tolerance)
                    if _fuzzy_addr_match(gl_addr_norm, sov_addr_norm):
                        corp = (sov_loc.get("corporate_name", "") or "").strip()
                        dba = (sov_loc.get("dba", "") or sov_loc.get("hotel_flag", "") or "").strip()
                        if corp and dba:
                            gl_loc["brand"] = f"{corp} - {dba}"
                        elif dba:
                            gl_loc["brand"] = dba
                        elif corp:
                            gl_loc["brand"] = corp
                        break
        
        if gl_loc_list:
            add_subsection_header(doc, "Covered Locations")
            add_formatted_paragraph(doc, 
                "The following locations are covered under this General Liability policy "
                "as identified on the carrier quote:",
                size=9, italic=True, color=CHARCOAL)
            headers = ["#", "Address", "Corporate Name / DBA"]
            rows = []
            for i, loc in enumerate(gl_loc_list, 1):
                rows.append([
                    str(i),
                    loc["address"],
                    loc["brand"],
                ])
            L = WD_ALIGN_PARAGRAPH.LEFT
            create_styled_table(doc, headers, rows,
                              col_widths=[0.4, 3.5, 3.6],
                              header_size=9, body_size=8,
                              header_alignments={0: L, 1: L, 2: L})


def generate_confirmation_to_bind(doc, data):
    """Section 14: Confirmation to Bind Agreement"""
    add_page_break(doc)
    add_section_header(doc, "Confirmation to Bind Agreement")
    
    # Show effective date prominently
    effective_date = data.get("client_info", {}).get("effective_date", "")
    if effective_date:
        add_formatted_paragraph(doc, f"Effective Date: {effective_date}", size=12,
                               color=ELECTRIC_BLUE, bold=True, space_after=8)
    
    add_formatted_paragraph(doc,
        "By signing below, the undersigned authorized representative of the Applicant confirms "
        "the following statements and authorizes HUB International to bind the coverages as outlined "
        "in this proposal, subject to the terms and conditions of the respective policies.",
        size=10, space_after=6)
    
    # Application Statements
    add_subsection_header(doc, "Application Statements")
    
    statements = [
        "The information provided in the applications and supplemental materials is true, accurate, and complete to the best of my knowledge.",
        "I understand that any material misrepresentation or omission may void coverage or result in denial of claims.",
        "I have reviewed the proposed coverages, limits, deductibles, and premiums outlined in this proposal.",
        "I understand that the coverages described herein are subject to the terms, conditions, and exclusions of the actual policies issued.",
        "I authorize HUB International to bind the coverages as outlined in this proposal on behalf of the named insured(s).",
        "I understand that subjectivities, if any, must be satisfied within the timeframes specified or coverage may be subject to cancellation.",
        "I acknowledge that surplus lines placements, if any, are not covered by state guaranty funds.",
        "I have been offered Terrorism Risk Insurance Act (TRIA) coverage and have made my election as indicated in this proposal.",
        "I understand that additional policies are available and recommended which include Equipment Breakdown (power surges, electrical arcing, mechanical failure), Employment Practices Liability (excluded by the liability carrier), Pollution (claims such as mold and legionella), Cyber, Flood, Earthquake, Deductible Buy Downs, Sexual Abuse & Molestation. If you would like to get these options quoted please request in writing to the producer or account executive."
    ]
    
    for i, stmt in enumerate(statements, 1):
        add_formatted_paragraph(doc, f"{i}. {stmt}", size=9, space_after=2)
    
    # Underwriting Confirmations (True/False checkmarks)
    add_subsection_header(doc, "Underwriting Confirmations")
    add_formatted_paragraph(doc,
        "The insured confirms the following by checking True or False:",
        size=10, space_after=6)
    
    confirmations = [
        "No prior losses for abuse & molestation, assault & battery, or human trafficking.",
        "Pest control service contract including bed bug prevention / detection.",
        "No homeless or government shelters.",
        "Human trafficking awareness program (annual training all employees).",
        "Background checks on all employees.",
    ]
    
    # Create a table with columns: #, Confirmation, True, False
    conf_table = doc.add_table(rows=len(confirmations) + 1, cols=4)
    conf_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Header row
    header_cells = conf_table.rows[0].cells
    for ci, (hdr_text, hdr_width) in enumerate([
        ("#", Inches(0.4)), ("Confirmation", Inches(5.0)),
        ("True", Inches(0.8)), ("False", Inches(0.8))
    ]):
        header_cells[ci].width = hdr_width
        p = header_cells[ci].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(hdr_text)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci >= 2 else WD_ALIGN_PARAGRAPH.LEFT
        # Blue background for header
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CLASSIC_BLUE_HEX}" w:val="clear"/>')
        header_cells[ci]._tc.get_or_add_tcPr().append(shading_elm)
    
    # Data rows
    for ri, conf_text in enumerate(confirmations):
        row_cells = conf_table.rows[ri + 1].cells
        # Row number
        p = row_cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(str(ri + 1))
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Confirmation text
        p = row_cells[1].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(conf_text)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        
        # True checkbox
        p = row_cells[2].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("\u2610")  # ☐ empty checkbox
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # False checkbox
        p = row_cells[3].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("\u2610")  # ☐ empty checkbox
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Alternate row shading
        if ri % 2 == 0:
            for cell in row_cells:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F6FA" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
    
    add_formatted_paragraph(doc, "", space_after=6)  # spacer
    
    # Earned premium / cancellation disclaimer - small font, bold, red
    _add_earned_premium_disclaimer(doc)
    
    # Signature block
    add_formatted_paragraph(doc, "", space_before=6)
    sig_table = doc.add_table(rows=5, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    sig_fields = [
        ("Authorized Signature:", ""),
        ("Printed Name:", ""),
        ("Title:", ""),
        ("Date:", ""),
        ("Company:", "")
    ]
    
    for i, (label, val) in enumerate(sig_fields):
        cell_label = sig_table.rows[i].cells[0]
        cell_label.width = Inches(2.0)
        p = cell_label.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(label)
        run.font.size = Pt(10)
        run.font.color.rgb = CLASSIC_BLUE
        run.font.bold = True
        run.font.name = "Calibri"
        
        cell_val = sig_table.rows[i].cells[1]
        cell_val.width = Inches(5.0)
        p = cell_val.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{CLASSIC_BLUE_HEX}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)


# --- Earned Premium Disclaimer (shared between Payment Options and Confirmation to Bind) ---
_EARNED_PREMIUM_DISCLAIMER = (
    "All insurance policies, including but not limited to property, general liability, "
    "umbrella/excess liability, and ancillary coverages, may be subject to minimum earned premiums "
    "as determined by the issuing carrier. Property policies frequently carry hurricane or named storm "
    "minimum earned premiums, which may require a significant portion of the annual premium to be fully "
    "earned regardless of the policy\u2019s cancellation or replacement date. Liability and umbrella/excess "
    "policies may also carry minimum earned premium provisions that limit or eliminate premium refunds "
    "upon cancellation.\n\n"
    "Additionally, most policies \u2014 both admitted and non-admitted (surplus lines) \u2014 are subject to "
    "short rate cancellation penalties if cancelled mid-term at the insured\u2019s request. Policy fees, "
    "inspection fees, and membership or association fees are typically fully earned at inception and "
    "non-refundable regardless of cancellation.\n\n"
    "These provisions vary by carrier, program, and policy form. Clients should carefully consider the "
    "financial implications of any mid-term policy changes, cancellations, or carrier transitions, as "
    "premium refunds may be limited or unavailable. HUB recommends reviewing all earned premium and "
    "cancellation provisions with your service team prior to binding or making any policy changes."
)

def _add_earned_premium_disclaimer(doc):
    """Add the earned premium disclaimer in small bold red font."""
    RED = RGBColor(0xCC, 0x00, 0x00)
    for para_text in _EARNED_PREMIUM_DISCLAIMER.split("\n\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(para_text.strip())
        run.font.size = Pt(6.5)
        run.font.bold = True
        run.font.color.rgb = RED
        run.font.name = "Calibri"


def generate_electronic_consent(doc):
    """Section 15: Electronic Documents Consent"""
    add_page_break(doc)
    add_section_header(doc, "Electronic Documents Consent")
    
    add_formatted_paragraph(doc,
        "By signing below, I consent to receive insurance policy documents, endorsements, "
        "certificates of insurance, notices of cancellation, renewal notices, and other "
        "insurance-related documents electronically from HUB International and/or the "
        "insurance carriers providing coverage.",
        size=11, space_after=10)
    
    add_formatted_paragraph(doc,
        "I understand that I may withdraw this consent at any time by providing written "
        "notice to my HUB International representative, and that I may request paper "
        "copies of any documents at no additional charge.",
        size=11, space_after=15)
    
    # Signature line
    sig_table = doc.add_table(rows=3, cols=2)
    for i, label in enumerate(["Authorized Signature:", "Printed Name:", "Date:"]):
        cell = sig_table.rows[i].cells[0]
        cell.width = Inches(2.0)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(label)
        run.font.size = Pt(11)
        run.font.color.rgb = CLASSIC_BLUE
        run.font.bold = True
        run.font.name = "Calibri"
        
        cell_val = sig_table.rows[i].cells[1]
        cell_val.width = Inches(5.0)
        p = cell_val.paragraphs[0]
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{CLASSIC_BLUE_HEX}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)


def generate_carrier_rating(doc, data):
    """Section 16: Carrier Rating"""
    add_page_break(doc)
    add_section_header(doc, "Carrier Ratings Summary")
    
    add_formatted_paragraph(doc,
        "AM Best is a full-service credit rating organization dedicated to serving the insurance "
        "industry. AM Best ratings provide an independent third-party evaluation of an insurer's "
        "financial strength and ability to meet its ongoing insurance policy and contract obligations.",
        size=11, space_after=10)
    
    # Build carrier rating table from all coverages
    coverages = data.get("coverages", {})
    carriers_seen = {}
    coverage_names = {
        "property": "Property",
        "property_alt_1": "Property (Option 2)",
        "property_alt_2": "Property (Option 3)",
        "excess_property": "Excess Property (Layer 1)",
        "excess_property_2": "Excess Property (Layer 2)",
        "general_liability": "General Liability",
        "general_liability_alt_1": "General Liability (Option 2)",
        "general_liability_alt_2": "General Liability (Option 3)",
        "umbrella": "Umbrella / Excess 1",
        "umbrella_alt_1": "Umbrella / Excess 2",
        "umbrella_alt_2": "Umbrella / Excess 3",
        "umbrella_alt_3": "Umbrella / Excess 4",
        "umbrella_layer_2": "2nd Excess Layer",
        "umbrella_layer_3": "3rd Excess Layer",
        "umbrella_layer_4": "4th Excess Layer",
        "excess_liability": "Excess Liability",
        "excess": "Excess Liability",
        "workers_comp": "Workers Compensation",
        "workers_compensation": "Workers Compensation",
        "workers_compensation_alt_1": "Workers Comp (Option 2)",
        "commercial_auto": "Commercial Auto",
        "terrorism": "Terrorism / TRIA",
        "cyber": "Cyber Liability",
        "cyber_alt_1": "Cyber (Option 2)",
        "epli": "Employment Practices Liability",
        "crime": "Crime",
        "flood": "Flood",
        "inland_marine": "Inland Marine",
        "equipment_breakdown": "Equipment Breakdown",
        "liquor_liability": "Liquor Liability",
        "innkeepers_liability": "Innkeepers Liability",
        "environmental": "Environmental / Pollution",
        "workplace_violence": "Workplace Violence",
        "garage_keepers": "Garage Keepers",
        "enviro_pack": "Enviro Pack",
        "wind_deductible_buydown": "Wind Deductible Buy Down",
        "earthquake": "Earthquake",
        "pollution": "Pollution Liability",
        "abuse_molestation": "Sexual Abuse & Molestation",
        "active_assailant": "Active Assailant",
        "deductible_buydown": "Deductible Buy Down",
    }
    
    for key, display_name in coverage_names.items():
        cov = coverages.get(key)
        if cov:
            carrier = _clean_carrier_name(cov.get("carrier", ""))
            if carrier and carrier not in carriers_seen:
                rating = cov.get("am_best_rating", "N/A")
                if not rating or rating == "N/A":
                    looked_up = lookup_am_best(carrier)
                    if looked_up:
                        rating = looked_up
                carriers_seen[carrier] = {
                    "rating": rating,
                    "admitted": "Admitted" if cov.get("carrier_admitted", True) else "Non-Admitted",
                    "coverages": [display_name]
                }
            elif carrier in carriers_seen:
                carriers_seen[carrier]["coverages"].append(display_name)
    
    if carriers_seen:
        headers = ["Carrier", "AM Best Rating", "Admitted Status", "Coverages"]
        rows = []
        for carrier, info in carriers_seen.items():
            rows.append([carrier, info["rating"], info["admitted"], ", ".join(info["coverages"])])
        create_styled_table(doc, headers, rows,
                          col_widths=[2.5, 1.2, 1.3, 2.5],
                          header_size=10, body_size=10)


def generate_general_statement(doc):
    """Section 17: General Statement"""
    add_page_break(doc)
    add_section_header(doc, "General Statement")
    
    sections = [
        ("Important Notice", "This proposal of insurance is provided for informational purposes only and does not constitute a binder of insurance. Coverage is not effective until confirmed in writing by the insurance carrier. The actual terms, conditions, and exclusions of coverage will be governed by the policies as issued. Please review your policies carefully upon receipt and report any discrepancies to your HUB International representative immediately."),
        ("Proposal Limitations", "This proposal is based on the information provided to us and the coverages available at the time of preparation. Insurance markets, rates, and terms are subject to change. HUB International makes no guarantee that the proposed coverages, limits, or premiums will remain available at the time of binding."),
        ("Claims Reporting", "All claims and potential claims should be reported immediately to your HUB International representative and directly to the applicable insurance carrier. Failure to report claims promptly may jeopardize coverage."),
        ("Policy Review", "We strongly recommend that you review all insurance policies upon receipt to ensure they accurately reflect the coverages intended. Any errors or omissions should be reported to your HUB International representative within 30 days of policy receipt."),
        ("Surplus Lines Notice", "Certain coverages in this proposal may be placed with surplus lines carriers. Surplus lines carriers are not members of state guaranty funds, and in the event of insolvency, claims may not be covered by state guaranty fund protection. Surplus lines placements are subject to surplus lines taxes and fees as required by applicable state law."),
    ]
    
    for title, text in sections:
        add_subsection_header(doc, title)
        add_formatted_paragraph(doc, text, size=10, space_after=6)


def generate_property_definitions(doc):
    """Section 18: Property Coverage Definitions"""
    add_page_break(doc)
    add_section_header(doc, "Property Coverage Definitions")
    
    definitions = [
        ("Actual Cash Value (ACV)", "The cost to repair or replace damaged property with material of like kind and quality, less depreciation."),
        ("Replacement Cost", "The cost to repair or replace damaged property with material of like kind and quality, without deduction for depreciation."),
        ("Agreed Value", "A predetermined value agreed upon by the insurer and insured, eliminating the coinsurance penalty."),
        ("Coinsurance", "A provision requiring the insured to carry insurance equal to a specified percentage of the property's value. Failure to maintain adequate limits may result in a penalty at the time of loss."),
        ("Business Income", "Coverage for loss of income sustained due to the necessary suspension of operations during the period of restoration following a covered loss."),
        ("Extra Expense", "Coverage for expenses incurred to avoid or minimize the suspension of business operations following a covered loss."),
        ("Ordinance or Law", "Coverage for the increased cost of construction due to enforcement of building codes or ordinances following a covered loss."),
        ("Equipment Breakdown", "Coverage for loss or damage to covered equipment caused by mechanical breakdown, electrical arcing, or other covered causes."),
        ("Flood", "Coverage for direct physical loss or damage caused by flood, as defined in the policy. Flood coverage may be subject to separate limits, deductibles, and waiting periods."),
        ("Earthquake", "Coverage for direct physical loss or damage caused by earthquake or earth movement. Subject to separate limits and deductibles."),
        ("Named Storm", "A storm system that has been designated a tropical storm or hurricane by the National Weather Service. Named storm deductibles typically apply as a percentage of insured values."),
        ("Mold", "Mold damage is typically excluded or subject to limited coverage under standard property policies."),
        ("Terrorism", "Coverage for acts of terrorism as defined by the Terrorism Risk Insurance Act (TRIA). See TRIA Disclosure section for details."),
    ]
    
    for term, definition in definitions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(13)
        run_term = p.add_run(f"{term}: ")
        run_term.font.size = Pt(10)
        run_term.font.color.rgb = ELECTRIC_BLUE
        run_term.font.bold = True
        run_term.font.name = "Calibri"
        run_def = p.add_run(definition)
        run_def.font.size = Pt(10)
        run_def.font.color.rgb = CLASSIC_BLUE
        run_def.font.name = "Calibri"


def generate_how_we_get_paid(doc):
    """Section 19: How We Get Paid"""
    add_page_break(doc)
    add_section_header(doc, "How We Get Paid")
    
    add_formatted_paragraph(doc,
        "HUB International takes pride in the services our brokerages provide to you, our client, "
        "for insurance and risk management programs. For our efforts we are compensated in a variety "
        "of ways, primarily in the form of commissions and contingency amounts paid by insurance "
        "companies and, in some cases, fees paid by clients or third parties. The means by which we "
        "are compensated are described below.",
        size=10, space_after=10)
    
    add_subsection_header(doc, "Commission income")
    add_formatted_paragraph(doc,
        "Commission, normally calculated as a percentage of the premium paid to the insurer for the "
        "specific policy, is paid to us by the insurer to distribute and service your insurance policy. "
        "Our commission is included in the premium paid by you. The individuals at HUB International "
        "who place and service your insurance may be paid compensation that varies directly with the "
        "commissions we receive.",
        size=10, space_after=6)
    
    add_subsection_header(doc, "Contingency income")
    add_formatted_paragraph(doc,
        "We also receive income through contingency arrangements with most insurers. They are called "
        "\"contingent\" because to qualify for payment we normally need to meet certain criteria, usually "
        "measured on an annual basis. Contingency arrangements vary, but payment under these agreements "
        "is normally the result of growing the business by attracting new customers, helping the insurance "
        "company gather and assess underwriting information and/or working to renew the policies of "
        "existing insureds. There is currently no meaningful method to determine the exact impact that "
        "any particular insurance policy has on contingency arrangements. However, brokers tend to receive "
        "higher contingency payments when they grow their business and retain clients through better service. "
        "In other words, the amount of earned contingency income depends on the overall size and/or "
        "profitability of all of a group of accounts, as opposed to the placement or profitability of any "
        "particular insurance policy. For this reason, the individuals involved in placing or servicing "
        "insurance are rarely, if ever, compensated directly for the contingent income that we receive.",
        size=10, space_after=6)
    
    add_formatted_paragraph(doc,
        "Please also feel free to ask any questions about our compensation generally, or as to your "
        "specific insurance proposal or placement, by contacting your HUB broker or customer service "
        "representative directly, or by calling our client hotline at 1-866-857-4073.",
        size=10, space_after=10)
    
    add_subsection_header(doc, "Privacy Policy")
    add_formatted_paragraph(doc,
        "To view our privacy policy, please visit: www.hubinternational.com/about-us/privacy-policy/",
        size=10)


def generate_hub_advantage(doc):
    """Section 20: The HUB Advantage"""
    add_page_break(doc)
    add_section_header(doc, "Our Commitment — The HUB Advantage")
    
    add_formatted_paragraph(doc,
        "HUB International is dedicated to maintaining and upholding the highest standards of ethical "
        "conduct and integrity in all of our dealings with you, our client. We want to be your trusted "
        "risk advisor, and as such, we need to earn your confidence. So we are making a promise. We call "
        "it The HUB Advantage. Our mission is to make the advantage yours — and this is our commitment.",
        size=10, space_after=10)
    
    commitments = [
        "We strive to secure the most favorable terms from insurers, taking into account all of the circumstances — the risk you need to insure, the cost of insurance, the financial condition of the insurer, the insurer's reputation for service, and any other factors that are specific to your situation.",
        "We are open and honest as to how we are paid for placing your insurance. Our answers to your questions will be forthright and understandable. When we intend to seek a fixed fee for our efforts, we will disclose it to you in writing and obtain your approval prior to coverage being bound.",
        "We comply with the laws of every jurisdiction in which we operate, including those that apply to how insurance brokerages and agencies are paid. If the laws change, we will respond in a timely and appropriate manner.",
    ]
    
    for commitment in commitments:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(13)
        run = p.add_run(f"• {commitment}")
        run.font.size = Pt(10)
        run.font.color.rgb = CLASSIC_BLUE
        run.font.name = "Calibri"
    
    add_formatted_paragraph(doc,
        "We take our responsibility to our customers very seriously. If at any time you feel that we are "
        "not fulfilling your expectations — that we are not meeting our Client Commitment — please contact "
        "your account executive or call our toll free client hotline at 1-866-857-4073, and your concerns "
        "will be addressed as soon as possible.",
        size=10, space_before=10, space_after=10)
    
    add_formatted_paragraph(doc, "The HUB Advantage", size=14, color=ELECTRIC_BLUE, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "The privilege is ours, but the advantage is yours.", size=12,
                           color=CLASSIC_BLUE, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def generate_tria_disclosure(doc):
    """Section 21: TRIA Disclosure"""
    add_page_break(doc)
    add_section_header(doc, "Terrorism Risk Insurance Act (TRIA) Disclosure")
    
    add_formatted_paragraph(doc,
        'You are hereby notified that under the Terrorism Risk Insurance Act, as amended in 2015, '
        'the definition of act of terrorism has changed. As defined in Section 102 (1) of the Act: '
        'The term "act of terrorism" means any act or acts that is certified by the Secretary of the '
        'Treasury — in consultation with the Secretary of Homeland Security, and the Attorney General '
        'of the United States — to be an act of terrorism; to be a violent act or an act that is '
        'dangerous to human life, property, or infrastructure; to have resulted in damage within the '
        'United States, or outside the United States in the case of certain air carriers or vessels or '
        'the premises of a United States mission; and to have been committed by an individual or '
        'individuals as a part of an effort to coerce the civilian population of the United States or '
        'to influence the policy or affect the conduct of the United States Government by coercion.',
        size=10, space_after=10)
    
    add_formatted_paragraph(doc,
        'Under the coverage, any losses resulting from certified acts of terrorism may be partially '
        'reimbursed by the United States Government under a formula established by the Terrorism Risk '
        'Insurance Act, as amended. However, your policy may contain other exclusions which might affect '
        'the coverage, such as an exclusion for nuclear events. Under the formula, the United States '
        'Government generally reimburses 80% beginning on January 1, 2020 of covered terrorism losses '
        'exceeding the statutorily established deductible paid by the insurance company providing the '
        'coverage. The Terrorism Risk Insurance Act, as amended, contains a $100 billion cap that limits '
        'United States government reimbursement as well as insurers\' liability for losses resulting from '
        'certified acts of terrorism when the amount of such losses exceed $100 billion in any one calendar '
        'year. If the aggregate insured losses for all insured exceed $100 billion, your coverage may be reduced.',
        size=10)


def generate_california_licenses(doc):
    """Section 22: California Licenses"""
    add_page_break(doc)
    add_section_header(doc, "California Licenses")
    
    add_formatted_paragraph(doc,
        "The following HUB International entities are licensed in the State of California:",
        size=11, space_after=8)
    
    headers = ["Entity Name", "License Number"]
    rows = [[name, lic] for name, lic in CA_LICENSES]
    create_styled_table(doc, headers, rows, col_widths=[5.5, 2.0],
                       header_size=9, body_size=8)


def generate_coverage_recommendations(doc):
    """Section 23: Coverage Recommendations"""
    add_page_break(doc)
    add_section_header(doc, "Coverage Recommendations")
    
    add_formatted_paragraph(doc,
        "HUB International recommends that you consider the following coverages and risk management "
        "strategies to protect your hospitality business. These recommendations are based on our "
        "extensive experience in the hotel and hospitality insurance industry.",
        size=11, space_after=10)
    
    recommendations = [
        ("Umbrella/Excess Liability", "We recommend maintaining umbrella or excess liability coverage with limits adequate to protect your assets. Hotels face unique liability exposures including guest injuries, swimming pool incidents, and food service operations."),
        ("Cyber Liability", "Hotels collect and store sensitive guest information including credit card numbers, personal identification, and travel details. Cyber liability coverage protects against data breaches, ransomware attacks, and regulatory fines."),
        ("Employment Practices Liability (EPLI)", "Hotels employ large numbers of workers in various roles, creating exposure to employment-related claims including wrongful termination, discrimination, harassment, and wage disputes."),
        ("Crime/Employee Dishonesty", "Hotels handle significant amounts of cash and guest valuables. Crime coverage protects against employee theft, forgery, computer fraud, and funds transfer fraud."),
        ("Flood Insurance", "Standard property policies exclude flood damage. If your properties are located in flood-prone areas, we strongly recommend purchasing separate flood coverage."),
        ("Equipment Breakdown", "Hotels rely heavily on mechanical and electrical equipment including HVAC systems, elevators, kitchen equipment, and laundry facilities."),
        ("Liquor Liability", "If your hotel serves alcohol, liquor liability coverage is essential. This coverage protects against claims arising from the sale, service, or furnishing of alcoholic beverages."),
    ]
    
    for title, text in recommendations:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(13)
        run_title = p.add_run(f"{title}: ")
        run_title.font.size = Pt(10)
        run_title.font.color.rgb = ELECTRIC_BLUE
        run_title.font.bold = True
        run_title.font.name = "Calibri"
        run_text = p.add_run(text)
        run_text.font.size = Pt(10)
        run_text.font.color.rgb = CLASSIC_BLUE
        run_text.font.name = "Calibri"
    
    add_formatted_paragraph(doc, "", space_before=10)
    add_callout_box(doc,
        "Please discuss these recommendations with your HUB International representative to determine "
        "which coverages are appropriate for your specific operations and risk profile.")


# ─── Main Generator ───────────────────────────────────────────

def generate_proposal(data: dict, output_path: str) -> str:
    """
    Generate a complete branded DOCX proposal.
    
    Args:
        data: Structured insurance data from extraction
        output_path: Path to save the DOCX file
        
    Returns:
        Path to the generated DOCX file
    """
    logger.info(f"Generating proposal for: {data.get('client_info', {}).get('named_insured', 'Unknown')}")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = CLASSIC_BLUE
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Part 1: Front Matter
    generate_cover_page(doc, data)
    generate_service_team(doc, data)
    generate_premium_summary(doc, data)
    generate_payment_options(doc, data)
    generate_subjectivities(doc, data)
    generate_named_insureds(doc, data)
    generate_information_summary(doc, data)
    generate_locations(doc, data)
    
    # Part 2: Coverage Sections (only if quoted)
    coverages = data.get("coverages", {})
    if "property" in coverages:
        generate_coverage_section(doc, data, "property", "Property Coverage")
    if "property_alt_1" in coverages:
        generate_coverage_section(doc, data, "property_alt_1", "Property Coverage — Option 2")
    if "property_alt_2" in coverages:
        generate_coverage_section(doc, data, "property_alt_2", "Property Coverage — Option 3")
    if "excess_property" in coverages:
        generate_coverage_section(doc, data, "excess_property", "Excess Property Coverage — Layer 1")
    if "excess_property_2" in coverages:
        generate_coverage_section(doc, data, "excess_property_2", "Excess Property Coverage — Layer 2")
    if "general_liability" in coverages:
        generate_coverage_section(doc, data, "general_liability", "General Liability Coverage")
    if "general_liability_alt_1" in coverages:
        generate_coverage_section(doc, data, "general_liability_alt_1", "General Liability Coverage — Option 2")
    if "general_liability_alt_2" in coverages:
        generate_coverage_section(doc, data, "general_liability_alt_2", "General Liability Coverage — Option 2")
    # Support both workers_comp and workers_compensation keys
    _wc_key = "workers_comp" if "workers_comp" in coverages else ("workers_compensation" if "workers_compensation" in coverages else None)
    if _wc_key:
        generate_coverage_section(doc, data, _wc_key, "Workers Compensation Coverage")
    if "commercial_auto" in coverages:
        generate_coverage_section(doc, data, "commercial_auto", "Commercial Auto Coverage")
    # Normalize any excess_liability / excess variants into umbrella layer slots
    _excess_aliases = ["excess_liability", "excess", "excess_layer_2", "excess_layer_3",
                       "2nd_excess", "second_excess", "3rd_excess", "third_excess"]
    for _alias in _excess_aliases:
        if _alias in coverages:
            if "umbrella" not in coverages:
                coverages["umbrella"] = coverages.pop(_alias)
            elif "umbrella_layer_2" not in coverages:
                coverages["umbrella_layer_2"] = coverages.pop(_alias)
            elif "umbrella_layer_3" not in coverages:
                coverages["umbrella_layer_3"] = coverages.pop(_alias)
            elif "umbrella_layer_4" not in coverages:
                coverages["umbrella_layer_4"] = coverages.pop(_alias)
    
    # Sort umbrella layers by attachment point before rendering
    # The extractor may assign layers in PDF order, not by actual layer structure
    _umb_keys = [k for k in ["umbrella", "umbrella_layer_2", "umbrella_layer_3", "umbrella_layer_4"] if k in coverages]
    if len(_umb_keys) > 1:
        import re as _re_umb
        def _parse_attachment_point(cov_data):
            """Determine attachment point from underlying insurance or limits.
            Lower attachment = lower layer (1st excess)."""
            underlying = cov_data.get("underlying_insurance", []) if isinstance(cov_data, dict) else []
            # If underlying includes primary policies (GL, Auto, WC), it's the 1st layer
            has_primary = False
            has_umbrella_underlying = False
            max_underlying_limit = 0
            for u in underlying:
                if isinstance(u, dict):
                    cov_type = (u.get("coverage", "") or "").lower()
                    limits_str = (u.get("limits", "") or "")
                    # Parse limit amount
                    limit_nums = _re_umb.findall(r'[\$]?([\d,]+)', str(limits_str))
                    for n in limit_nums:
                        try:
                            val = int(n.replace(',', ''))
                            if val > max_underlying_limit:
                                max_underlying_limit = val
                        except ValueError:
                            pass
                    if any(kw in cov_type for kw in ["general liability", "auto", "workers", "employer"]):
                        has_primary = True
                    if any(kw in cov_type for kw in ["umbrella", "excess"]):
                        has_umbrella_underlying = True
            # If it has primary underlying, it's the 1st layer (attachment ~$1M)
            if has_primary and not has_umbrella_underlying:
                return 1000000  # $1M attachment
            # If it has umbrella/excess underlying, it's a higher layer
            # Use the max underlying limit as a proxy for attachment point
            if has_umbrella_underlying and max_underlying_limit > 0:
                return max_underlying_limit
            # Fallback: try to parse from limits description
            limits = cov_data.get("limits", []) if isinstance(cov_data, dict) else []
            for lim in limits:
                if isinstance(lim, dict):
                    desc = (lim.get("description", "") or "").lower()
                    val = (lim.get("limit", "") or "")
                    if "retention" in desc or "attachment" in desc or "underlying" in desc:
                        nums = _re_umb.findall(r'[\$]?([\d,]+)', str(val))
                        for n in nums:
                            try:
                                return int(n.replace(',', ''))
                            except ValueError:
                                pass
            # Check if underlying_insurance has any entries with $1,000,000 limits
            # This is the most reliable indicator: primary excess sits on top of $1M primary policies
            if underlying:
                for u in underlying:
                    if isinstance(u, dict):
                        limits_str = str(u.get("limits", "") or "")
                        if "1,000,000" in limits_str or "1000000" in limits_str:
                            return 1000000  # Primary excess layer
            
            # If no underlying info, use premium as proxy (higher premium = lower/primary layer usually)
            prem = cov_data.get("total_premium", 0) or cov_data.get("premium", 0) or 0
            if isinstance(cov_data, dict) and prem > 0:
                # Invert so higher premium sorts first (lower attachment point)
                return 10000000 - prem
            return 50000000  # Unknown layers sort last
        
        # Sort umbrella keys by attachment point (ascending = 1st layer first)
        _umb_sorted = sorted(_umb_keys, key=lambda k: _parse_attachment_point(coverages.get(k, {})))
        # Reassign to canonical keys: umbrella, umbrella_layer_2, umbrella_layer_3
        _canonical_keys = ["umbrella", "umbrella_layer_2", "umbrella_layer_3", "umbrella_layer_4"]
        _umb_data_backup = {k: coverages[k] for k in _umb_keys}
        for i, sorted_key in enumerate(_umb_sorted):
            target_key = _canonical_keys[i]
            coverages[target_key] = _umb_data_backup[sorted_key]
        logger.info(f"Umbrella layer order after sorting: {[coverages[k].get('carrier', 'unknown') for k in _canonical_keys[:len(_umb_keys)]]}")
    
    _umb_titles = {
        "umbrella": "Umbrella / Excess Liability Coverage",
        "umbrella_layer_2": "2nd Excess Liability Layer",
        "umbrella_layer_3": "3rd Excess Liability Layer",
        "umbrella_layer_4": "4th Excess Liability Layer",
    }
    for _uk in ["umbrella", "umbrella_layer_2", "umbrella_layer_3", "umbrella_layer_4"]:
        if _uk in coverages:
            generate_coverage_section(doc, data, _uk, _umb_titles[_uk])
    if "cyber" in coverages:
        generate_coverage_section(doc, data, "cyber", "Cyber Liability Coverage")
    if "epli" in coverages:
        generate_coverage_section(doc, data, "epli", "Employment Practices Liability (EPLI) Coverage")
    if "flood" in coverages:
        generate_coverage_section(doc, data, "flood", "Flood Coverage")
    if "terrorism" in coverages:
        generate_coverage_section(doc, data, "terrorism", "Terrorism / TRIA Coverage")
    if "crime" in coverages:
        generate_coverage_section(doc, data, "crime", "Crime Coverage")
    if "inland_marine" in coverages:
        generate_coverage_section(doc, data, "inland_marine", "Inland Marine Coverage")
    if "equipment_breakdown" in coverages:
        generate_coverage_section(doc, data, "equipment_breakdown", "Equipment Breakdown Coverage")
    if "liquor_liability" in coverages:
        generate_coverage_section(doc, data, "liquor_liability", "Liquor Liability Coverage")
    if "innkeepers_liability" in coverages:
        generate_coverage_section(doc, data, "innkeepers_liability", "Innkeepers Liability Coverage")
    if "environmental" in coverages:
        generate_coverage_section(doc, data, "environmental", "Environmental / Pollution Coverage")
    if "workplace_violence" in coverages:
        generate_coverage_section(doc, data, "workplace_violence", "Workplace Violence Coverage")
    if "garage_keepers" in coverages:
        generate_coverage_section(doc, data, "garage_keepers", "Garage Keepers Coverage")
    if "wind_deductible_buydown" in coverages:
        generate_coverage_section(doc, data, "wind_deductible_buydown", "Wind Deductible Buy Down Coverage")
    if "enviro_pack" in coverages:
        generate_coverage_section(doc, data, "enviro_pack", "Enviro Pack Coverage")
    if "earthquake" in coverages:
        generate_coverage_section(doc, data, "earthquake", "Earthquake Coverage")
    if "pollution" in coverages:
        generate_coverage_section(doc, data, "pollution", "Pollution Liability Coverage")
    if "abuse_molestation" in coverages:
        generate_coverage_section(doc, data, "abuse_molestation", "Sexual Abuse & Molestation Coverage")
    if "active_assailant" in coverages:
        generate_coverage_section(doc, data, "active_assailant", "Active Assailant Coverage")
    if "deductible_buydown" in coverages:
        generate_coverage_section(doc, data, "deductible_buydown", "Deductible Buy Down Coverage")

    # Catch-all: generate sections for any remaining _alt_ keys not explicitly handled above
    for cov_key in sorted(coverages.keys()):
        if "_alt_" in cov_key and cov_key not in ("property_alt_1", "property_alt_2",
                "general_liability_alt_1", "general_liability_alt_2"):
            base_name = cov_key.split("_alt_")[0].replace("_", " ").title()
            alt_num = cov_key.split("_alt_")[-1]
            display = f"{base_name} Coverage — Option {int(alt_num) + 1}"
            generate_coverage_section(doc, data, cov_key, display)
    
    # Part 3: Coverage Recommendations (before signature pages)
    generate_coverage_recommendations(doc)
    
    # Part 4: Signature Pages
    generate_confirmation_to_bind(doc, data)
    
    # Part 5: Compliance Pages (ALWAYS REQUIRED)
    generate_electronic_consent(doc)
    generate_carrier_rating(doc, data)
    generate_general_statement(doc)
    generate_property_definitions(doc)
    generate_how_we_get_paid(doc)
    generate_hub_advantage(doc)
    generate_tria_disclosure(doc)
    generate_california_licenses(doc)
    
    # Save
    doc.save(output_path)
    logger.info(f"Proposal saved to: {output_path}")
    return output_path
