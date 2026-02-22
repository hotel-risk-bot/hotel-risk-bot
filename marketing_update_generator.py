#!/usr/bin/env python3
"""
Marketing Update DOCX Generator.
Pulls policies from the Airtable Sales System for a given client/opportunity
and generates a professional Marketing Update document in DOCX format.

Two versions:
  - Internal (/marketingsummary): includes commission, revenue, rates, broker info
  - Client-facing (/marketingsummaryclient): hides sensitive financial data

Uses HUB International branding with color-coded carrier comparison tables.
"""

import os
import logging
import tempfile
from collections import defaultdict
from datetime import datetime

import requests as http_requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

# ── Airtable Config ──
AIRTABLE_PAT = os.environ.get("AIRTABLE_PAT", "")
SALES_BASE_ID = "appnFKEzmdLbR4CHY"
POLICIES_TABLE_ID = "tbl8vZP2oHrinwVfd"
OPPORTUNITIES_TABLE_ID = "tblMKuUsG1cosdQPN"

# ── HUB Colors ──
ELECTRIC_BLUE = RGBColor(0x16, 0x7B, 0xD4)
CLASSIC_BLUE = RGBColor(0x26, 0x38, 0x45)
ARCTIC_GRAY = RGBColor(0xB8, 0xC4, 0xCE)
EGGSHELL = RGBColor(0xF3, 0xF5, 0xF1)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CHARCOAL = RGBColor(0x4A, 0x4A, 0x4A)
DARK_GREEN = RGBColor(0x1B, 0x7A, 0x3D)
DARK_RED = RGBColor(0xC0, 0x39, 0x2B)

ELECTRIC_BLUE_HEX = "167BD4"
CLASSIC_BLUE_HEX = "263845"
EGGSHELL_HEX = "F3F5F1"
LIGHT_GREEN_HEX = "E8F5E9"
LIGHT_RED_HEX = "FDEDEC"
LIGHT_BLUE_HEX = "E3F2FD"
LIGHT_GRAY_HEX = "F0F0F0"
PROPOSED_GREEN_HEX = "E8F5E9"
EXPIRING_GRAY_HEX = "EDEFF1"
PENDING_YELLOW_HEX = "FFF8E1"

LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "hub_logo.png")

# ── Coverage ordering ──
COVERAGE_ORDER = [
    "Property", "Liability", "Umbrella", "Workers Compensation",
    "Auto", "Employment Practices Liability", "Cyber", "Flood",
    "Equipment Breakdown", "Terrorism", "Crime", "Inland Marine", "Package",
]

COVERAGE_DISPLAY_NAMES = {
    "Property": "Property",
    "Liability": "General Liability",
    "Umbrella": "Umbrella / Excess Liability",
    "Workers Compensation": "Workers Compensation",
    "Auto": "Commercial Auto",
    "Employment Practices Liability": "Employment Practices Liability (EPLI)",
    "Cyber": "Cyber Liability",
    "Flood": "Flood",
    "Equipment Breakdown": "Equipment Breakdown",
    "Terrorism": "Terrorism / TRIA",
    "Crime": "Crime",
    "Inland Marine": "Inland Marine",
    "Package": "Package",
}

COVERAGE_SHORT = {
    "Property": "Property",
    "Liability": "General Liability",
    "Umbrella": "Umbrella",
    "Workers Compensation": "Workers Comp",
    "Auto": "Auto",
    "Employment Practices Liability": "EPLI",
    "Cyber": "Cyber",
    "Flood": "Flood",
    "Equipment Breakdown": "Equip Breakdown",
    "Terrorism": "Terrorism",
    "Crime": "Crime",
    "Inland Marine": "Inland Marine",
    "Package": "Package",
}

STATUS_PRIORITY = {
    "Incumbent": 0, "Bound": 1, "Proposed": 2, "Quoted": 3,
    "Market": 4, "Submit": 5, "Pending": 6, "Declined": 7,
    "Blocked": 8, "Lost": 9,
}

# ── Carrier Abbreviation Map ──
CARRIER_ABBR_MAP = {
    'SLAKE': 'Southlake Specialty Insurance Company',
    'ORGIN': 'Origin Specialty',
    'ARCH': 'Arch Insurance',
    'CONCERT': 'Concert Group',
    'TRAV': 'Travelers',
    'ZENITH': 'Zenith National Insurance',
    'HARTFORD': 'The Hartford',
    'PHILA': 'Philadelphia Insurance',
    'BERKS': 'Berkshire Hathaway Specialty Insurance',
    'WESTFIELD': 'Westfield Specialty',
    'CORE': 'Core Specialty',
    'EVEREST': 'Everest National Insurance Company',
    'RIVINGTON': 'Rivington Insurance',
    'KINSALE': 'Kinsale Capital Group',
    'IRONSHORE': 'Ironshore Insurance Services',
    'HDI': 'HDI Global Specialty',
    'GRAY': 'Gray Surplus Lines',
    'EVANSTON': 'Evanston Insurance Company',
    'ADMIRAL': 'Admiral Insurance Company',
    'AMWINS': 'AmWins',
    'RT': 'RT Specialty',
    'STARR': 'Starr Surplus Lines',
    'ZURICH': 'Zurich Insurance',
    'CHUBB': 'Chubb',
    'AIG': 'AIG',
    'CNA': 'CNA Insurance',
    'LIBERTY': 'Liberty Mutual',
    'MARKEL': 'Markel Insurance',
    'COLONY': 'Colony Specialty',
    'SCOTTSDALE': 'Scottsdale Insurance',
    'NAUTILUS': 'Nautilus Insurance',
    'EMPLOYERS': 'Employers Insurance',
    'PINNACOL': 'Pinnacol Assurance',
    'GUARD': 'Guard Insurance',
    'AMTRUST': 'AmTrust Financial',
}

# ── Broker Abbreviation Map ──
BROKER_ABBR_MAP = {
    'AMWINS': 'AmWins',
    'RT': 'RT Specialty',
    'CRC': 'CRC Group',
    'BURNS': 'Burns & Wilcox',
    'RSG': 'Ryan Specialty Group',
    'USI': 'USI Insurance Services',
    'WHOLESALE': 'Wholesale Trading',
}


# ══════════════════════════════════════════════════════════════════════════
# UTILITY FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════

def airtable_headers():
    return {
        "Authorization": f"Bearer {AIRTABLE_PAT}",
        "Content-Type": "application/json",
    }


def _safe_str(val, default="—"):
    if val is None:
        return default
    if isinstance(val, list):
        return ", ".join(str(v) for v in val if v is not None) or default
    s = str(val).strip()
    return s if s else default


def _safe_currency(val, default="—"):
    """Format as currency without cents (e.g., $1,788,571)."""
    if val is None:
        return default
    try:
        n = float(str(val).replace("$", "").replace(",", ""))
        if n == 0:
            return default
        return f"${n:,.0f}"
    except (ValueError, TypeError):
        return default


def _safe_currency_int(val, default="—"):
    if val is None:
        return default
    try:
        n = float(str(val).replace("$", "").replace(",", ""))
        if n == 0:
            return default
        return f"${n:,.0f}"
    except (ValueError, TypeError):
        return default


def _safe_number(val, default="—"):
    if val is None:
        return default
    try:
        n = float(str(val).replace(",", ""))
        if n == 0:
            return default
        if n == int(n):
            return f"{int(n):,}"
        return f"{n:,.2f}"
    except (ValueError, TypeError):
        return default


def _safe_percent(val, default="—"):
    if val is None:
        return default
    try:
        n = float(val)
        if n == 0:
            return default
        return f"{n:.2%}" if n < 1 else f"{n:.2f}%"
    except (ValueError, TypeError):
        return default


def _resolve_carrier_name(raw_name):
    """Resolve carrier abbreviation to full name using the carrier map."""
    if not raw_name or raw_name == "N/A":
        return raw_name or "N/A"
    name = str(raw_name).strip()
    # Check if the raw name is an abbreviation
    upper = name.upper()
    if upper in CARRIER_ABBR_MAP:
        return CARRIER_ABBR_MAP[upper]
    # Check if it contains a known abbreviation as a substring
    for abbr, full in CARRIER_ABBR_MAP.items():
        if name.upper() == abbr:
            return full
    return name


def _fmt_limit(val):
    """Format limits as $20M, $10M, $5M, $500k, etc."""
    if not val or val == "\u2014" or val == "—":
        return "\u2014"
    try:
        clean = str(val).replace("$", "").replace(",", "").strip()
        num = float(clean)
        if num >= 1_000_000:
            m = num / 1_000_000
            if m == int(m):
                return f"${int(m)}M"
            else:
                return f"${m:.1f}M"
        elif num >= 1_000:
            k = num / 1_000
            if k == int(k):
                return f"${int(k)}k"
            else:
                return f"${k:.1f}k"
        elif num > 0:
            return f"${int(num):,}"
        else:
            return "\u2014"
    except (ValueError, TypeError):
        return str(val) if val else "\u2014"


def _get_float(val, default=0):
    if val is None:
        return default
    try:
        return float(str(val).replace("$", "").replace(",", "").replace("%", ""))
    except (ValueError, TypeError):
        return default


def _resolve_broker_names(broker_record_ids: list) -> str:
    """Resolve broker linked record IDs to company names via Airtable API."""
    if not broker_record_ids or not AIRTABLE_PAT:
        return "—"
    names = []
    companies_table = "tblMPEvjv6mcSwdSd"  # Companies table
    for rid in broker_record_ids:
        if not isinstance(rid, str) or not rid.startswith("rec"):
            continue
        url = f"https://api.airtable.com/v0/{SALES_BASE_ID}/{companies_table}/{rid}"
        try:
            resp = http_requests.get(url, headers=airtable_headers(), timeout=10)
            resp.raise_for_status()
            data = resp.json()
            flds = data.get("fields", {})
            # Try ABBR first, then Name
            name = flds.get("ABBR") or flds.get("Name") or flds.get("Company Name") or ""
            if name:
                names.append(str(name).strip())
        except Exception as e:
            logger.warning(f"Could not resolve broker record {rid}: {e}")
    return ", ".join(names) if names else "—"


def _resolve_broker_from_fields(flds: dict) -> str:
    """Resolve broker name from policy fields. Checks multiple paths."""
    # Path 1: Broker ABBR rollup (from Related Broker)
    broker_abbr = flds.get("Broker ABBR")
    if broker_abbr and str(broker_abbr).strip() and str(broker_abbr).strip() != "\u2014":
        raw = str(broker_abbr).strip()
        # Map known abbreviations to full names
        return BROKER_ABBR_MAP.get(raw.upper(), raw)

    # Path 2: Check Direct field - if True, it's a direct placement
    direct = flds.get("Direct")
    is_direct = False
    if isinstance(direct, list):
        is_direct = any(v is True for v in direct)
    elif isinstance(direct, bool):
        is_direct = direct

    # Path 3: Resolve Brokers linked record IDs
    brokers = flds.get("Brokers")
    if brokers and isinstance(brokers, list) and any(isinstance(b, str) and b.startswith("rec") for b in brokers):
        resolved = _resolve_broker_names(brokers)
        if resolved != "\u2014":
            return resolved

    # If direct placement with no broker
    if is_direct:
        return "Direct"

    return "\u2014"


def _normalize_coverage_type(policy_type_raw):
    if isinstance(policy_type_raw, list):
        pt = policy_type_raw[0] if policy_type_raw else "Other"
    else:
        pt = str(policy_type_raw) if policy_type_raw else "Other"
    pt_lower = pt.lower().strip()
    if "property" in pt_lower:
        return "Property"
    elif "terror" in pt_lower or "tria" in pt_lower:
        return "Terrorism"
    elif "liability" in pt_lower and "employment" not in pt_lower:
        return "Liability"
    elif "umbrella" in pt_lower or "excess" in pt_lower:
        return "Umbrella"
    elif "worker" in pt_lower or "comp" in pt_lower:
        return "Workers Compensation"
    elif "auto" in pt_lower:
        return "Auto"
    elif "employment" in pt_lower or "epli" in pt_lower:
        return "Employment Practices Liability"
    elif "cyber" in pt_lower:
        return "Cyber"
    elif "flood" in pt_lower:
        return "Flood"
    elif "equipment" in pt_lower or "breakdown" in pt_lower:
        return "Equipment Breakdown"
    elif "crime" in pt_lower:
        return "Crime"
    elif "inland" in pt_lower or "marine" in pt_lower:
        return "Inland Marine"
    elif "package" in pt_lower or "pkg" in pt_lower:
        return "Package"
    else:
        return pt if pt else "Other"


def _coverage_sort_key(coverage_type):
    try:
        return COVERAGE_ORDER.index(coverage_type)
    except ValueError:
        return len(COVERAGE_ORDER)


def _status_sort_key(status):
    return STATUS_PRIORITY.get(status, 99)


def _map_status_for_display(status):
    """Map Airtable status to display status for the document."""
    s = (status or "").strip()
    if s == "Incumbent":
        return "Expiring"
    elif s in ("Bound", "Proposed"):
        return s
    elif s in ("Quoted",):
        return "Quoted"
    elif s in ("Market", "Submit"):
        return "Pending"
    elif s in ("Declined", "Blocked", "Lost"):
        return s
    else:
        return s or "Unknown"


def _status_color_hex(display_status):
    """Return background color hex for carrier column based on display status."""
    if display_status == "Expiring":
        return EXPIRING_GRAY_HEX
    elif display_status in ("Bound", "Proposed"):
        return PROPOSED_GREEN_HEX
    elif display_status == "Quoted":
        return None  # White / alternating
    elif display_status == "Pending":
        return PENDING_YELLOW_HEX
    elif display_status in ("Declined", "Blocked", "Lost"):
        return "FDE8E8"  # Light red for declined/blocked
    else:
        return None


# ══════════════════════════════════════════════════════════════════════════
# AIRTABLE DATA FETCHING (reuse from marketing_summary.py)
# ══════════════════════════════════════════════════════════════════════════

def _sanitize_for_formula(text: str) -> str:
    return text.replace('"', '\\"')


def search_opportunity(client_name: str, upcoming_only: bool = True) -> list:
    safe_name = _sanitize_for_formula(client_name)
    dq = '"'
    search_part = (
        f"OR("
        f"SEARCH(LOWER({dq}{safe_name}{dq}), LOWER({{Opportunity Name}})),"
        f"SEARCH(LOWER({dq}{safe_name}{dq}), LOWER(ARRAYJOIN({{Corporate Name}},{dq}{dq})))"
        f")"
    )
    if upcoming_only:
        formula = f"AND({search_part}, {{Days to Expiration}} >= 0)"
    else:
        formula = search_part

    url = f"https://api.airtable.com/v0/{SALES_BASE_ID}/{OPPORTUNITIES_TABLE_ID}"
    params = {
        "filterByFormula": formula,
        "sort[0][field]": "Effective Date",
        "sort[0][direction]": "asc",
        "pageSize": 20,
    }
    try:
        resp = http_requests.get(url, headers=airtable_headers(), params=params, timeout=30)
        resp.raise_for_status()
        return resp.json().get("records", [])
    except Exception as e:
        logger.error(f"Error searching opportunities: {e}")
        return []


def fetch_policies_by_record_ids(policy_record_ids: list) -> list:
    if not policy_record_ids:
        return []
    all_records = []
    batch_size = 20
    for i in range(0, len(policy_record_ids), batch_size):
        batch = policy_record_ids[i:i + batch_size]
        conditions = [f'RECORD_ID()="{rid}"' for rid in batch]
        formula = f"OR({','.join(conditions)})" if len(conditions) > 1 else conditions[0]
        url = f"https://api.airtable.com/v0/{SALES_BASE_ID}/{POLICIES_TABLE_ID}"
        offset = None
        while True:
            params = {"filterByFormula": formula, "pageSize": 100}
            if offset:
                params["offset"] = offset
            try:
                resp = http_requests.get(url, headers=airtable_headers(), params=params, timeout=30)
                resp.raise_for_status()
                data = resp.json()
                all_records.extend(data.get("records", []))
                offset = data.get("offset")
                if not offset:
                    break
            except Exception as e:
                logger.error(f"Error fetching policies by record IDs: {e}")
                break
    return all_records


def fetch_policies_for_client(client_name: str) -> list:
    safe_name = _sanitize_for_formula(client_name)
    dq = '"'
    formula = (
        f"OR("
        f"SEARCH(LOWER({dq}{safe_name}{dq}), LOWER({{Name}})),"
        f"SEARCH(LOWER({dq}{safe_name}{dq}), LOWER(ARRAYJOIN({{Companies}},{dq}{dq})))"
        f")"
    )
    url = f"https://api.airtable.com/v0/{SALES_BASE_ID}/{POLICIES_TABLE_ID}"
    all_records = []
    offset = None
    while True:
        params = {"filterByFormula": formula, "pageSize": 100}
        if offset:
            params["offset"] = offset
        try:
            resp = http_requests.get(url, headers=airtable_headers(), params=params, timeout=30)
            resp.raise_for_status()
            data = resp.json()
            all_records.extend(data.get("records", []))
            offset = data.get("offset")
            if not offset:
                break
        except Exception as e:
            logger.error(f"Error fetching policies: {e}")
            break
    return all_records


async def resolve_client_data(client_name: str):
    """Resolve client name to opportunity + policies. Returns (opp_fields, policies) or (None, [])."""
    opportunities = search_opportunity(client_name)
    if opportunities:
        opp = opportunities[0]
        opp_fields = opp.get("fields", {})
        policy_ids = opp_fields.get("Policies", [])
        if policy_ids:
            policies = fetch_policies_by_record_ids(policy_ids)
        else:
            policies = fetch_policies_for_client(client_name)
        if policies:
            return opp_fields, policies

    # Direct policy search
    policies = fetch_policies_for_client(client_name)
    if policies:
        return None, policies

    # Try individual words
    words = client_name.strip().split()
    if len(words) > 1:
        for word in words:
            if len(word) >= 3:
                opportunities = search_opportunity(word)
                if opportunities:
                    opp = opportunities[0]
                    opp_fields = opp.get("fields", {})
                    policy_ids = opp_fields.get("Policies", [])
                    if policy_ids:
                        policies = fetch_policies_by_record_ids(policy_ids)
                        if policies:
                            return opp_fields, policies
                policies = fetch_policies_for_client(word)
                if policies:
                    return None, policies

    return None, []


# ══════════════════════════════════════════════════════════════════════════
# POLICY DATA PARSING
# ══════════════════════════════════════════════════════════════════════════

def parse_policies(policies: list):
    """Parse raw Airtable policy records into structured dicts grouped by coverage type."""
    parsed = []
    for rec in policies:
        flds = rec.get("fields", {})
        coverage_type = _normalize_coverage_type(flds.get("Policy Type"))

        insurance_co_raw = flds.get("Insurance Company", "")
        if isinstance(insurance_co_raw, list):
            insurance_co = ", ".join(str(ic) for ic in insurance_co_raw if ic) or "N/A"
        else:
            insurance_co = str(insurance_co_raw).strip() if insurance_co_raw else "N/A"
        # Resolve carrier abbreviation to full name
        insurance_co = _resolve_carrier_name(insurance_co)

        status = flds.get("Status", "Unknown")
        display_status = _map_status_for_display(status)

        base_premium = _get_float(flds.get("Base Premium"))
        premium_tx = _get_float(flds.get("Premium Tx"))
        # Use Premium Tx if available, else Base Premium
        premium_with_tax = premium_tx if premium_tx > 0 else base_premium

        commission = _get_float(flds.get("Commission"))
        revenue = _get_float(flds.get("Revenue"))

        parsed.append({
            "coverage_type": coverage_type,
            "carrier": insurance_co,
            "status": status,
            "display_status": display_status,
            "base_premium": base_premium,
            "premium_tx": premium_with_tax,
            "commission": commission,
            "revenue": revenue,
            "broker": _resolve_broker_from_fields(flds),
            "units": flds.get("Units"),
            "num_locs": flds.get("# of Locs"),
            "tiv": flds.get("TIV"),
            "property_rate": flds.get("Property Rate"),
            "property_limit": _fmt_limit(flds.get("Property Limit")),
            "aop": _safe_str(flds.get("AOP")),
            "wind_type": _safe_str(flds.get("Wind Type")),
            "wind": _safe_str(flds.get("Wind")),
            "aow": _safe_str(flds.get("AOW")),
            "water_damage": _safe_str(flds.get("Water Damage")),
            "flood_limit": _fmt_limit(flds.get("Flood Limit")),
            "flood_deductible": _safe_str(flds.get("Flood Deductible")),
            "eq_limit": _fmt_limit(flds.get("EQ Limit")),
            "eq_deductible": _safe_str(flds.get("Earthquake Deductible")),
            "gross_sales": flds.get("Gross Sales"),
            "gl_rate": flds.get("GL Rate $"),
            "gl_rate_unit": flds.get("GL Rate (u)"),
            "gl_deductible": _safe_str(flds.get("GL Deductible")),
            "umb_limit": _fmt_limit(flds.get("UMB Limit")),
            "total_payroll": flds.get("Total Payroll"),
            "exp_mod": flds.get("Exp Mod"),
            "safety_credit": flds.get("Safety"),
            "drug_free_credit": flds.get("Drug Free"),
            "comments": flds.get("Comments", ""),
            "fields": flds,
        })

    return parsed


def group_by_coverage(parsed_policies):
    """Group parsed policies by coverage type, sorted by coverage order."""
    by_coverage = defaultdict(list)
    for p in parsed_policies:
        by_coverage[p["coverage_type"]].append(p)

    # Sort carriers within each coverage: Expiring first, then Bound/Proposed, then Quoted, then Pending
    for ct in by_coverage:
        by_coverage[ct].sort(key=lambda p: _status_sort_key(p["status"]))

    return by_coverage


# ══════════════════════════════════════════════════════════════════════════
# DOCX BUILDING UTILITIES
# ══════════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_width(cell, width_inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width_inches * 1440)}" w:type="dxa"/>')
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)


def set_cell_vertical_alignment(cell, val="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val}"/>')
    existing = tcPr.find(qn('w:vAlign'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(vAlign)


def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tcBorders>'
    )
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(borders)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def set_thin_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def add_page_header(doc):
    section = doc.sections[-1]
    header = section.header
    header.is_linked_to_previous = False
    htable = header.add_table(1, 2, width=Inches(7))
    htable.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo_cell = htable.rows[0].cells[0]
    logo_cell.width = Inches(2.5)
    if os.path.exists(LOGO_PATH):
        p = logo_cell.paragraphs[0]
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.8))
    text_cell = htable.rows[0].cells[1]
    text_cell.width = Inches(4.5)
    p = text_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Franchise Division")
    run.font.size = Pt(12)
    run.font.color.rgb = ELECTRIC_BLUE
    run.font.bold = True
    run.font.name = "Calibri"
    p2 = text_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(0)
    run2 = p2.add_run("Hotel Insurance Programs")
    run2.font.size = Pt(12)
    run2.font.color.rgb = ELECTRIC_BLUE
    run2.font.bold = True
    run2.font.name = "Calibri"
    for row in htable.rows:
        for cell in row.cells:
            remove_cell_borders(cell)


def add_formatted_paragraph(doc, text, size=11, color=CLASSIC_BLUE, bold=False,
                            italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(size + 3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Calibri"
    return p


def add_section_header(doc, text):
    return add_formatted_paragraph(doc, text, size=22, color=CLASSIC_BLUE, bold=True,
                                   space_before=24, space_after=12)


def add_subsection_header(doc, text):
    return add_formatted_paragraph(doc, text, size=14, color=ELECTRIC_BLUE, bold=True,
                                   space_before=14, space_after=6)


def add_callout_box(doc, text, size=10, shading_hex=EGGSHELL_HEX):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, shading_hex)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = CLASSIC_BLUE
    run.font.name = "Calibri"
    run.font.italic = True
    return table


def add_rich_callout_box(doc, lines, size=10, shading_hex=EGGSHELL_HEX):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, shading_hex)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)
    for i, line_data in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.1)
        text, bold, color = line_data
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = "Calibri"
        run.font.bold = bold
    return table


def add_divider(doc):
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(4)
    p_div.paragraph_format.space_after = Pt(4)
    pPr = p_div._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="{ELECTRIC_BLUE_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_page_break(doc):
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# TABLE BUILDERS
# ══════════════════════════════════════════════════════════════════════════

def create_premium_summary_table(doc, headers, rows, col_widths=None, highlight_last_row=False):
    """Create the premium summary comparison table."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    existing = tblPr.find(qn('w:tblLayout'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblLayout)
    total_width = sum(col_widths) if col_widths else 7.5
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{int(total_width * 1440)}" w:type="dxa"/>')
    existing_tblW = tblPr.find(qn('w:tblW'))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblPr.append(tblW)

    if not col_widths:
        col_widths = [total_width / num_cols] * num_cols

    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(header_text)
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.bold = True
        run.font.name = "Calibri"
        set_cell_shading(cell, ELECTRIC_BLUE_HEX)
        set_cell_width(cell, col_widths[i])
        set_cell_vertical_alignment(cell, "center")

    # Data rows
    for row_idx, row_data in enumerate(rows):
        is_total = highlight_last_row and row_idx == len(rows) - 1
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            if col_idx < 2:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Calibri"

            if is_total:
                run.font.color.rgb = WHITE
                run.font.bold = True
                set_cell_shading(cell, CLASSIC_BLUE_HEX)
            else:
                run.font.color.rgb = CLASSIC_BLUE
                # Color code $ Change column
                val_str = str(val)
                if "Change" in headers[col_idx] if col_idx < len(headers) else False:
                    pass
                if col_idx >= 2 and val_str.startswith("-$"):
                    run.font.color.rgb = DARK_GREEN
                elif col_idx >= 2 and val_str.startswith("+$"):
                    run.font.color.rgb = DARK_RED
                if row_idx % 2 == 1:
                    set_cell_shading(cell, EGGSHELL_HEX)

            set_cell_width(cell, col_widths[col_idx])
            set_cell_vertical_alignment(cell, "center")

    return table


def create_carrier_comparison_table(doc, coverage_title, metrics, carriers):
    """
    Create a carrier comparison table for a coverage line.
    carriers: list of dicts with keys: name, status, values (dict), notes
    metrics: list of metric labels (rows)
    """
    num_cols = 1 + len(carriers)
    num_rows = 1 + len(metrics)

    has_notes = any(c.get("notes") for c in carriers)
    if has_notes:
        num_rows += 1

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_thin_borders(table)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    existing = tblPr.find(qn('w:tblLayout'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblLayout)

    total_width = 7.5
    label_width = 1.6
    carrier_width = (total_width - label_width) / max(len(carriers), 1)

    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{int(total_width * 1440)}" w:type="dxa"/>')
    existing_tblW = tblPr.find(qn('w:tblW'))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblPr.append(tblW)

    # Header row
    header_cell = table.rows[0].cells[0]
    header_cell.text = ""
    p = header_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(coverage_title)
    run.font.size = Pt(10)
    run.font.color.rgb = WHITE
    run.font.bold = True
    run.font.name = "Calibri"
    set_cell_shading(header_cell, ELECTRIC_BLUE_HEX)
    set_cell_width(header_cell, label_width)
    set_cell_vertical_alignment(header_cell, "center")

    for c_idx, carrier in enumerate(carriers):
        cell = table.rows[0].cells[1 + c_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(12)

        run = p.add_run(carrier["name"])
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.bold = True
        run.font.name = "Calibri"

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(3)
        p2.paragraph_format.line_spacing = Pt(10)
        run2 = p2.add_run(f"({carrier['status']})")
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0xCC, 0xE5, 0xFF)
        run2.font.italic = True
        run2.font.name = "Calibri"

        set_cell_shading(cell, ELECTRIC_BLUE_HEX)
        set_cell_width(cell, carrier_width)
        set_cell_vertical_alignment(cell, "center")

    # Metric rows
    for m_idx, metric in enumerate(metrics):
        label_cell = table.rows[1 + m_idx].cells[0]
        label_cell.text = ""
        p = label_cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(12)
        run = p.add_run(metric)
        run.font.size = Pt(9)
        run.font.color.rgb = CLASSIC_BLUE
        run.font.bold = True
        run.font.name = "Calibri"
        set_cell_width(label_cell, label_width)
        set_cell_vertical_alignment(label_cell, "center")

        is_premium_row = (metric == "Premium")

        for c_idx, carrier in enumerate(carriers):
            cell = table.rows[1 + m_idx].cells[1 + c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = Pt(12)

            val = carrier["values"].get(metric, "—")
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Calibri"

            if is_premium_row:
                run.font.bold = True
                run.font.size = Pt(10)

            # Color coding by status
            status = carrier["status"]
            bg_hex = _status_color_hex(status)
            if bg_hex:
                set_cell_shading(cell, bg_hex)
            elif m_idx % 2 == 1:
                set_cell_shading(cell, "FAFAFA")

            run.font.color.rgb = CLASSIC_BLUE if status != "Pending" else CHARCOAL

            set_cell_width(cell, carrier_width)
            set_cell_vertical_alignment(cell, "center")

    # Notes row
    if has_notes:
        notes_row_idx = 1 + len(metrics)
        label_cell = table.rows[notes_row_idx].cells[0]
        label_cell.text = ""
        p = label_cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run("Notes")
        run.font.size = Pt(9)
        run.font.color.rgb = CLASSIC_BLUE
        run.font.bold = True
        run.font.italic = True
        run.font.name = "Calibri"
        set_cell_width(label_cell, label_width)

        for c_idx, carrier in enumerate(carriers):
            cell = table.rows[notes_row_idx].cells[1 + c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = Pt(11)
            note = carrier.get("notes", "—") or "—"
            run = p.add_run(note)
            run.font.size = Pt(8)
            run.font.color.rgb = CHARCOAL
            run.font.italic = True
            run.font.name = "Calibri"
            set_cell_width(cell, carrier_width)

            bg_hex = _status_color_hex(carrier["status"])
            if bg_hex:
                set_cell_shading(cell, bg_hex)

    return table


# ══════════════════════════════════════════════════════════════════════════
# COVERAGE-SPECIFIC METRIC BUILDERS
# ══════════════════════════════════════════════════════════════════════════

def _build_property_carrier(p, is_internal=True):
    """Build carrier dict for property comparison table."""
    is_declined = p["display_status"] in ("Declined", "Blocked", "Lost")
    if is_declined:
        values = {"Premium": "Declined"}
        if p["broker"] != "\u2014":
            values["Broker"] = p["broker"]
        comments = p.get("comments", "")
        if isinstance(comments, str) and len(comments) > 150:
            comments = comments[:147] + "..."
        return {
            "name": p["carrier"],
            "status": p["display_status"],
            "values": values,
            "notes": comments if comments else "Declined to quote",
        }
    values = {
        "Premium": _safe_currency(p["premium_tx"]) if p["premium_tx"] else "Pending",
        "TIV": _safe_currency_int(p["tiv"]),
        "# of Locations": _safe_number(p["num_locs"]),
        "AOP Deductible": p["aop"],
        "Wind": f"{p['wind']} ({p['wind_type']})" if p["wind"] != "\u2014" and p["wind_type"] != "\u2014" else p["wind"],
        "AOW (All Other Wind)": p["aow"],
        "Water Damage": p["water_damage"],
    }
    if is_internal:
        # Calculate rate from base premium / TIV if property_rate not available
        if p["property_rate"]:
            values["Property Rate"] = _safe_currency(p["property_rate"])
        elif p["base_premium"] and p["tiv"]:
            try:
                tiv_val = float(str(p["tiv"]).replace("$", "").replace(",", ""))
                if tiv_val > 0:
                    calc_rate = (p["base_premium"] / tiv_val) * 100
                    values["Property Rate"] = f"${calc_rate:.2f}"
                else:
                    values["Property Rate"] = "\u2014"
            except (ValueError, TypeError):
                values["Property Rate"] = "\u2014"
        else:
            values["Property Rate"] = "\u2014"
    if p["property_limit"] != "\u2014":
        values["Property Limit"] = p["property_limit"]
    if p["flood_limit"] != "\u2014":
        values["Flood Limit"] = p["flood_limit"]
    if p["eq_limit"] != "\u2014":
        values["EQ Limit"] = p["eq_limit"]
    # Broker on all policies
    if p["broker"] != "\u2014":
        values["Broker"] = p["broker"]
    if is_internal:
        if p["commission"]:
            values["Commission"] = _safe_percent(p["commission"])
        if p["revenue"]:
            values["Revenue"] = _safe_currency(p["revenue"])

    # Build notes from comments
    comments = p.get("comments", "")
    if isinstance(comments, str) and len(comments) > 100:
        comments = comments[:97] + "..."

    return {
        "name": p["carrier"],
        "status": p["display_status"],
        "values": values,
        "notes": comments if comments else "",
    }


def _get_property_metrics(carriers_data, is_internal=True):
    """Determine which property metrics to show based on available data."""
    metrics = ["Premium", "TIV", "# of Locations"]
    if is_internal:
        metrics.append("Property Rate")
    metrics.extend(["AOP Deductible", "Wind", "AOW (All Other Wind)", "Water Damage"])

    # Only add optional metrics if any carrier has data
    all_values = {}
    for c in carriers_data:
        for k, v in c["values"].items():
            if k not in all_values:
                all_values[k] = []
            all_values[k].append(v)

    optional_metrics = ["Property Limit", "Flood Limit", "EQ Limit", "Broker"]
    if is_internal:
        optional_metrics.extend(["Commission", "Revenue"])

    for m in optional_metrics:
        if m in all_values and any(v != "\u2014" for v in all_values[m]):
            metrics.append(m)

    return metrics


def _build_gl_carrier(p, is_internal=True):
    is_declined = p["display_status"] in ("Declined", "Blocked", "Lost")
    if is_declined:
        # For declined carriers, show minimal info with reason
        values = {
            "Premium": "Declined",
        }
        if p["broker"] != "\u2014":
            values["Broker"] = p["broker"]
        comments = p.get("comments", "")
        if isinstance(comments, str) and len(comments) > 150:
            comments = comments[:147] + "..."
        return {
            "name": p["carrier"],
            "status": p["display_status"],
            "values": values,
            "notes": comments if comments else "Declined to quote",
        }

    values = {
        "Premium": _safe_currency(p["premium_tx"]) if p["premium_tx"] else "Pending",
        "# of Units": _safe_number(p["units"]),
        "Total Sales": _safe_currency_int(p["gross_sales"]),
        "GL Deductible": p["gl_deductible"],
        "# of Locations": _safe_number(p["num_locs"]),
    }
    if is_internal:
        values["GL Rate"] = _safe_currency(p["gl_rate"]) if p["gl_rate"] else "\u2014"
        values["GL Rate/Unit"] = _safe_currency(p["gl_rate_unit"]) if p["gl_rate_unit"] else "\u2014"
    # Broker on all policies
    if p["broker"] != "\u2014":
        values["Broker"] = p["broker"]
    if is_internal:
        if p["commission"]:
            values["Commission"] = _safe_percent(p["commission"])
        if p["revenue"]:
            values["Revenue"] = _safe_currency(p["revenue"])

    comments = p.get("comments", "")
    if isinstance(comments, str) and len(comments) > 100:
        comments = comments[:97] + "..."

    return {
        "name": p["carrier"],
        "status": p["display_status"],
        "values": values,
        "notes": comments if comments else "",
    }


def _get_gl_metrics(carriers_data, is_internal=True):
    metrics = ["Premium", "# of Units", "Total Sales"]
    if is_internal:
        metrics.extend(["GL Rate", "GL Rate/Unit"])
    metrics.extend(["GL Deductible", "# of Locations"])
    all_values = {}
    for c in carriers_data:
        for k, v in c["values"].items():
            if k not in all_values:
                all_values[k] = []
            all_values[k].append(v)
    # Broker on all versions
    if "Broker" in all_values and any(v != "\u2014" for v in all_values["Broker"]):
        metrics.append("Broker")
    if is_internal:
        for m in ["Commission", "Revenue"]:
            if m in all_values and any(v != "\u2014" for v in all_values[m]):
                metrics.append(m)
    return metrics


def _build_umbrella_carrier(p, is_internal=True):
    is_declined = p["display_status"] in ("Declined", "Blocked", "Lost")
    if is_declined:
        values = {"Premium": "Declined"}
        if p["broker"] != "\u2014":
            values["Broker"] = p["broker"]
        comments = p.get("comments", "")
        if isinstance(comments, str) and len(comments) > 150:
            comments = comments[:147] + "..."
        return {
            "name": p["carrier"],
            "status": p["display_status"],
            "values": values,
            "notes": comments if comments else "Declined to quote",
        }
    values = {
        "Premium": _safe_currency(p["premium_tx"]) if p["premium_tx"] else "Pending",
        "# of Units": _safe_number(p["units"]),
        "# of Locations": _safe_number(p["num_locs"]),
        "Umbrella Limit": p["umb_limit"],
        "Total Sales": _safe_currency_int(p["gross_sales"]),
    }
    # Calculate umbrella rate per unit if data available
    if p["premium_tx"] and p["units"]:
        try:
            rate_per_unit = p["premium_tx"] / float(p["units"])
            values["Rate/Unit"] = f"${rate_per_unit:,.2f}"
        except (ValueError, TypeError, ZeroDivisionError):
            pass
    # Calculate umbrella rate per sales if data available
    if p["premium_tx"] and p["gross_sales"]:
        try:
            rate_per_sales = (p["premium_tx"] / float(p["gross_sales"])) * 1000
            values["Rate/$1K Sales"] = f"${rate_per_sales:,.2f}"
        except (ValueError, TypeError, ZeroDivisionError):
            pass
    # Broker on all policies
    if p["broker"] != "\u2014":
        values["Broker"] = p["broker"]
    if is_internal:
        if p["commission"]:
            values["Commission"] = _safe_percent(p["commission"])
        if p["revenue"]:
            values["Revenue"] = _safe_currency(p["revenue"])

    comments = p.get("comments", "")
    if isinstance(comments, str) and len(comments) > 100:
        comments = comments[:97] + "..."

    return {
        "name": p["carrier"],
        "status": p["display_status"],
        "values": values,
        "notes": comments if comments else "",
    }


def _get_umbrella_metrics(carriers_data, is_internal=True):
    metrics = ["Premium", "# of Units", "# of Locations", "Umbrella Limit", "Total Sales"]
    all_values = {}
    for c in carriers_data:
        for k, v in c["values"].items():
            if k not in all_values:
                all_values[k] = []
            all_values[k].append(v)
    # Rate metrics
    for m in ["Rate/Unit", "Rate/$1K Sales"]:
        if m in all_values and any(v != "\u2014" for v in all_values.get(m, [])):
            metrics.append(m)
    # Broker on all versions
    if "Broker" in all_values and any(v != "\u2014" for v in all_values["Broker"]):
        metrics.append("Broker")
    if is_internal:
        for m in ["Commission", "Revenue"]:
            if m in all_values and any(v != "\u2014" for v in all_values[m]):
                metrics.append(m)
    return metrics


def _build_wc_carrier(p, is_internal=True):
    is_declined = p["display_status"] in ("Declined", "Blocked", "Lost")
    if is_declined:
        values = {"Premium": "Declined"}
        if p["broker"] != "\u2014":
            values["Broker"] = p["broker"]
        comments = p.get("comments", "")
        if isinstance(comments, str) and len(comments) > 150:
            comments = comments[:147] + "..."
        return {
            "name": p["carrier"],
            "status": p["display_status"],
            "values": values,
            "notes": comments if comments else "Declined to quote",
        }
    values = {
        "Premium": _safe_currency(p["premium_tx"]) if p["premium_tx"] else "Pending",
        "Total Payroll": _safe_currency_int(p["total_payroll"]),
    }
    # Broker on all policies
    if p["broker"] != "\u2014":
        values["Broker"] = p["broker"]
    if p["exp_mod"]:
        try:
            em = float(p["exp_mod"])
            if em > 0:
                values["Exp Mod"] = f"{em:.2f}"
        except (ValueError, TypeError):
            pass
    if p["safety_credit"]:
        values["Safety Credit"] = "Yes" if p["safety_credit"] else "No"
    if p["drug_free_credit"]:
        values["Drug Free Credit"] = "Yes" if p["drug_free_credit"] else "No"
    if is_internal:
        if p["commission"]:
            values["Commission"] = _safe_percent(p["commission"])
        if p["revenue"]:
            values["Revenue"] = _safe_currency(p["revenue"])

    comments = p.get("comments", "")
    if isinstance(comments, str) and len(comments) > 100:
        comments = comments[:97] + "..."

    return {
        "name": p["carrier"],
        "status": p["display_status"],
        "values": values,
        "notes": comments if comments else "",
    }


def _get_wc_metrics(carriers_data, is_internal=True):
    metrics = ["Premium", "Total Payroll"]
    all_values = {}
    for c in carriers_data:
        for k, v in c["values"].items():
            if k not in all_values:
                all_values[k] = []
            all_values[k].append(v)
    for m in ["Exp Mod", "Safety Credit", "Drug Free Credit"]:
        if m in all_values and any(v not in ("\u2014", "No") for v in all_values[m]):
            metrics.append(m)
    # Broker on all versions
    if "Broker" in all_values and any(v != "\u2014" for v in all_values["Broker"]):
        metrics.append("Broker")
    if is_internal:
        for m in ["Commission", "Revenue"]:
            if m in all_values and any(v != "\u2014" for v in all_values[m]):
                metrics.append(m)
    return metrics


def _build_generic_carrier(p, is_internal=True):
    """Generic carrier builder for EPLI, Cyber, Flood, Auto, etc."""
    is_declined = p["display_status"] in ("Declined", "Blocked", "Lost")
    if is_declined:
        values = {"Premium": "Declined"}
        if p["broker"] != "\u2014":
            values["Broker"] = p["broker"]
        comments = p.get("comments", "")
        if isinstance(comments, str) and len(comments) > 150:
            comments = comments[:147] + "..."
        return {
            "name": p["carrier"],
            "status": p["display_status"],
            "values": values,
            "notes": comments if comments else "Declined to quote",
        }
    values = {
        "Premium": _safe_currency(p["premium_tx"]) if p["premium_tx"] else "Pending",
    }
    if p["units"]:
        values["# of Units"] = _safe_number(p["units"])
    if p["num_locs"]:
        values["# of Locations"] = _safe_number(p["num_locs"])

    # Coverage-specific fields
    ct = p["coverage_type"]
    if ct == "Flood":
        if p["flood_limit"] != "\u2014":
            values["Flood Limit"] = p["flood_limit"]
        if p["flood_deductible"] != "\u2014":
            values["Flood Deductible"] = p["flood_deductible"]
    elif ct == "Employment Practices Liability":
        if p["umb_limit"] != "\u2014":
            values["Limit"] = p["umb_limit"]
    elif ct == "Cyber":
        # Cyber Liability: include Gross Sales
        if p["gross_sales"]:
            values["Gross Sales"] = _safe_currency_int(p["gross_sales"])

    # Broker on all policies (not just internal)
    if p["broker"] != "\u2014":
        values["Broker"] = p["broker"]

    if is_internal:
        if p["commission"]:
            values["Commission"] = _safe_percent(p["commission"])
        if p["revenue"]:
            values["Revenue"] = _safe_currency(p["revenue"])

    comments = p.get("comments", "")
    if isinstance(comments, str) and len(comments) > 100:
        comments = comments[:97] + "..."

    return {
        "name": p["carrier"],
        "status": p["display_status"],
        "values": values,
        "notes": comments if comments else "",
    }


def _get_generic_metrics(carriers_data, is_internal=True):
    metrics = ["Premium"]
    all_values = {}
    for c in carriers_data:
        for k, v in c["values"].items():
            if k not in all_values:
                all_values[k] = []
            all_values[k].append(v)
    # Add metrics that have data (available on all versions)
    optional = ["# of Units", "# of Locations", "Gross Sales", "Limit", "Flood Limit", "Flood Deductible",
                "Building Limit", "BPP Limit", "Retention", "Broker"]
    for m in optional:
        if m in all_values and any(v != "\u2014" for v in all_values[m]):
            metrics.append(m)
    # Internal-only metrics
    if is_internal:
        for m in ["Commission", "Revenue"]:
            if m in all_values and any(v != "\u2014" for v in all_values[m]):
                metrics.append(m)
    return metrics


# ══════════════════════════════════════════════════════════════════════════
# TOWER BUILDERS (combine multi-carrier placements into single column)
# ══════════════════════════════════════════════════════════════════════════

def _build_property_tower(policies, is_internal=True):
    """For multi-carrier property placements, combine incumbents into a single 'Property Tower' column."""
    incumbents = sorted([p for p in policies if p["status"] == "Incumbent"],
                        key=lambda x: x["premium_tx"], reverse=True)
    if len(incumbents) <= 1:
        return None  # No tower needed

    tower_prem = sum(p["premium_tx"] for p in incumbents)
    # TIV is the same across all carriers in the tower (shared exposure)
    tower_tiv = 0
    for p in incumbents:
        if p["tiv"]:
            try:
                tiv_val = float(str(p["tiv"]).replace("$", "").replace(",", ""))
                if tiv_val > tower_tiv:
                    tower_tiv = tiv_val
            except (ValueError, TypeError):
                pass
    # Blended rate = combined base premium / TIV * 100 (rate per $100)
    tower_base_prem = sum(p["base_premium"] for p in incumbents)
    blended_rate = (tower_base_prem / tower_tiv * 100) if tower_tiv > 0 and tower_base_prem > 0 else 0

    # Get max limits
    def _max_limit(pols, key):
        vals = []
        for p in pols:
            v = p.get(key, "\u2014")
            if v and v != "\u2014":
                try:
                    clean = str(v).replace("$", "").replace(",", "").replace("M", "000000").replace("k", "000")
                    vals.append(float(clean))
                except (ValueError, TypeError):
                    pass
        return _fmt_limit(str(int(max(vals)))) if vals else "\u2014"

    values = {
        "Premium": _safe_currency(tower_prem),
        "TIV": _safe_currency_int(tower_tiv),
        "# of Locations": _safe_number(incumbents[0]["num_locs"]),
        "AOP Deductible": incumbents[0]["aop"],
        "Wind": f"{incumbents[0]['wind']} ({incumbents[0]['wind_type']})" if incumbents[0]["wind"] != "\u2014" else "\u2014",
        "AOW (All Other Wind)": incumbents[0]["aow"],
        "Water Damage": incumbents[0]["water_damage"],
    }
    if is_internal:
        values["Property Rate"] = f"${blended_rate:.2f}" if blended_rate else "\u2014"
    prop_limit = _max_limit(incumbents, "property_limit")
    if prop_limit != "\u2014":
        values["Property Limit"] = prop_limit
    flood_limit = _max_limit(incumbents, "flood_limit")
    if flood_limit != "\u2014":
        values["Flood Limit"] = flood_limit
    eq_limit = _max_limit(incumbents, "eq_limit")
    if eq_limit != "\u2014":
        values["EQ Limit"] = eq_limit
    if incumbents[0]["broker"] != "\u2014":
        values["Broker"] = incumbents[0]["broker"]
    if is_internal:
        total_rev = sum(p["revenue"] for p in incumbents if p["revenue"])
        comm = incumbents[0]["commission"]
        all_same = all(p["commission"] == comm for p in incumbents)
        if comm:
            values["Commission"] = _safe_percent(comm) + ("" if all_same else " (blended)")
        if total_rev:
            values["Revenue"] = _safe_currency(total_rev)

    return {
        "name": f"Property Tower ({len(incumbents)} carriers)",
        "status": "Expiring",
        "values": values,
        "notes": f"Expiring tower: {', '.join(p['carrier'] for p in incumbents)}",
    }


def _build_umbrella_tower(policies, is_internal=True):
    """For multi-carrier umbrella placements, combine incumbents into a single 'Umbrella Tower' column."""
    incumbents = sorted([p for p in policies if p["status"] == "Incumbent"],
                        key=lambda x: x["premium_tx"], reverse=True)
    if len(incumbents) <= 1:
        return None  # No tower needed

    tower_prem = sum(p["premium_tx"] for p in incumbents)
    total_limit = 0
    layer_names = ["Primary", "1st Excess", "2nd Excess", "3rd Excess", "4th Excess"]
    tower_desc = []
    for i, p in enumerate(incumbents):
        layer = layer_names[i] if i < len(layer_names) else f"{i+1}th Excess"
        tower_desc.append(f"{p['carrier']} ({p['umb_limit']} {layer})")
        try:
            clean = str(p["umb_limit"]).replace("$", "").replace(",", "").replace("M", "000000").replace("k", "000")
            total_limit += float(clean)
        except (ValueError, TypeError):
            pass

    values = {
        "Premium": _safe_currency(tower_prem),
        "# of Units": _safe_number(incumbents[0]["units"]),
        "# of Locations": _safe_number(incumbents[0]["num_locs"]),
        "Umbrella Limit": _fmt_limit(str(int(total_limit))) + " (combined)" if total_limit else "\u2014",
        "Total Sales": _safe_currency_int(incumbents[0]["gross_sales"]),
    }
    if incumbents[0]["broker"] != "\u2014":
        values["Broker"] = incumbents[0]["broker"]
    if is_internal:
        total_rev = sum(p["revenue"] for p in incumbents if p["revenue"])
        if incumbents[0]["commission"]:
            values["Commission"] = _safe_percent(incumbents[0]["commission"])
        if total_rev:
            values["Revenue"] = _safe_currency(total_rev)

    return {
        "name": f"Umbrella Tower ({len(incumbents)} layers)",
        "status": "Expiring",
        "values": values,
        "notes": f"Tower: {', '.join(tower_desc)}",
    }


def _add_internal_detail_pages(doc, by_coverage, parsed_policies):
    """Add internal-only detail pages for Property Tower, Umbrella Tower, and GL Detail."""
    has_detail = False

    # Property Carrier Detail
    prop_policies = by_coverage.get("Property", [])
    incumbents = sorted([p for p in prop_policies if p["status"] == "Incumbent"],
                        key=lambda x: x["premium_tx"], reverse=True)
    if len(incumbents) > 1:
        if not has_detail:
            add_page_break(doc)
            add_section_header(doc, "Internal Detail Pages")
            add_callout_box(doc, "The following pages are for internal use only and should not be shared with the client.",
                           size=10, shading_hex="FDE8E8")
            has_detail = True

        add_formatted_paragraph(doc, "", size=6, space_before=0, space_after=0)
        add_subsection_header(doc, "Property Carrier Detail (Expiring Tower)")
        add_formatted_paragraph(doc,
            "The following table breaks out each individual carrier within the expiring property placement tower.",
            size=10, color=CLASSIC_BLUE, space_before=2, space_after=8)

        # Build individual carrier columns for the tower
        tower_carriers = []
        for p in incumbents:
            carrier_dict = _build_property_carrier(p, is_internal=True)
            # Add base premium to values
            carrier_dict["values"]["Base Premium"] = _safe_currency(p["base_premium"])
            tower_carriers.append(carrier_dict)

        metrics = ["Base Premium", "Premium"] + _get_property_metrics(tower_carriers, is_internal=True)
        # Remove duplicates while preserving order
        seen = set()
        unique_metrics = []
        for m in metrics:
            if m not in seen:
                seen.add(m)
                unique_metrics.append(m)
        create_carrier_comparison_table(doc, "Property Tower Detail", unique_metrics, tower_carriers)

        # Add totals note
        total_base = sum(p["base_premium"] for p in incumbents)
        total_prem = sum(p["premium_tx"] for p in incumbents)
        total_rev = sum(p["revenue"] for p in incumbents if p["revenue"])
        add_callout_box(doc,
            f"Tower Totals: Base Premium {_safe_currency(total_base)} | "
            f"Premium w/ Tax {_safe_currency(total_prem)} | "
            f"Total Revenue {_safe_currency(total_rev)}",
            size=9)

    # Umbrella Tower Detail
    umb_policies = by_coverage.get("Umbrella", [])
    umb_incumbents = sorted([p for p in umb_policies if p["status"] == "Incumbent"],
                            key=lambda x: x["premium_tx"], reverse=True)
    if len(umb_incumbents) > 1:
        if not has_detail:
            add_page_break(doc)
            add_section_header(doc, "Internal Detail Pages")
            add_callout_box(doc, "The following pages are for internal use only and should not be shared with the client.",
                           size=10, shading_hex="FDE8E8")
            has_detail = True

        add_page_break(doc)
        add_subsection_header(doc, "Umbrella Tower Detail (Expiring)")

        layer_names = ["Primary", "1st Excess", "2nd Excess", "3rd Excess", "4th Excess"]
        tower_carriers = []
        for i, p in enumerate(umb_incumbents):
            layer = layer_names[i] if i < len(layer_names) else f"{i+1}th Excess"
            carrier_dict = _build_umbrella_carrier(p, is_internal=True)
            carrier_dict["name"] = f"{p['carrier']} ({layer})"
            carrier_dict["values"]["Base Premium"] = _safe_currency(p["base_premium"])
            tower_carriers.append(carrier_dict)

        metrics = ["Base Premium", "Premium", "Umbrella Limit", "# of Units", "Rate/Unit", "Broker", "Commission", "Revenue"]
        # Filter to only metrics that exist
        all_vals = {}
        for c in tower_carriers:
            for k, v in c["values"].items():
                if k not in all_vals:
                    all_vals[k] = []
                all_vals[k].append(v)
        metrics = [m for m in metrics if m in all_vals]
        create_carrier_comparison_table(doc, "Umbrella Tower Detail", metrics, tower_carriers)

    # GL Quoted Detail (if more than 2 GL policies)
    gl_policies = by_coverage.get("Liability", [])
    if len(gl_policies) > 2:
        if not has_detail:
            add_page_break(doc)
            add_section_header(doc, "Internal Detail Pages")
            add_callout_box(doc, "The following pages are for internal use only and should not be shared with the client.",
                           size=10, shading_hex="FDE8E8")
            has_detail = True

        add_page_break(doc)
        add_subsection_header(doc, "GL Quoted Detail")

        gl_carriers = []
        for p in gl_policies:
            carrier_dict = _build_gl_carrier(p, is_internal=True)
            carrier_dict["values"]["Base Premium"] = _safe_currency(p["base_premium"])
            carrier_dict["values"]["Status"] = p["display_status"]
            gl_carriers.append(carrier_dict)

        metrics = ["Status", "Base Premium", "Premium", "Total Sales", "GL Rate", "GL Rate/Unit",
                   "# of Units", "Broker", "Commission", "Revenue"]
        all_vals = {}
        for c in gl_carriers:
            for k, v in c["values"].items():
                if k not in all_vals:
                    all_vals[k] = []
                all_vals[k].append(v)
        metrics = [m for m in metrics if m in all_vals]
        create_carrier_comparison_table(doc, "GL Detail", metrics, gl_carriers)

    return has_detail


# Coverage type -> (carrier_builder, metrics_builder) mapping
COVERAGE_BUILDERS = {
    "Property": (_build_property_carrier, _get_property_metrics),
    "Liability": (_build_gl_carrier, _get_gl_metrics),
    "Umbrella": (_build_umbrella_carrier, _get_umbrella_metrics),
    "Workers Compensation": (_build_wc_carrier, _get_wc_metrics),
}


# ══════════════════════════════════════════════════════════════════════════
# PREMIUM COMPARISON DATA BUILDER
# ══════════════════════════════════════════════════════════════════════════

def build_premium_comparison(by_coverage, parsed_policies):
    """Build premium comparison rows: Coverage | Carrier | Expiring | Proposed | $ Change | % Change."""
    rows = []
    total_expiring = 0
    total_proposed = 0
    pending_coverages = []

    sorted_coverages = sorted(by_coverage.keys(), key=_coverage_sort_key)

    for ct in sorted_coverages:
        policies = by_coverage[ct]
        display_name = COVERAGE_SHORT.get(ct, ct)

        # Find expiring (Incumbent) and proposed/bound
        expiring = [p for p in policies if p["status"] == "Incumbent"]
        bound = [p for p in policies if p["status"] in ("Bound", "Proposed")]
        quoted = [p for p in policies if p["status"] == "Quoted"]

        expiring_premium = sum(p["premium_tx"] for p in expiring) if expiring else 0
        proposed_premium = sum(p["premium_tx"] for p in bound) if bound else 0

        # Determine carrier name for the row
        if bound:
            carrier_name = bound[0]["carrier"]
        elif quoted:
            carrier_name = quoted[0]["carrier"]
        elif expiring:
            carrier_name = expiring[0]["carrier"]
        else:
            carrier_name = policies[0]["carrier"] if policies else "—"

        # Check if this is a "bundled" coverage (e.g., EPLI included in GL)
        is_included = False
        for p in policies:
            comments = str(p.get("comments", "")).lower()
            if "included" in comments or "inc " in comments[:10]:
                is_included = True
                break

        if proposed_premium > 0:
            exp_str = _safe_currency(expiring_premium) if expiring_premium else "—"
            prop_str = _safe_currency(proposed_premium)
            if expiring_premium > 0:
                change = proposed_premium - expiring_premium
                pct_change = (change / expiring_premium) * 100
                change_str = f"+${change:,.0f}" if change > 0 else f"-${abs(change):,.0f}"
                pct_str = f"+{pct_change:.1f}%" if change > 0 else f"{pct_change:.1f}%"
                total_expiring += expiring_premium
                total_proposed += proposed_premium
            else:
                change_str = "\u2014"
                pct_str = "\u2014"
                total_proposed += proposed_premium
            rows.append([display_name, carrier_name, exp_str, prop_str, change_str, pct_str])
        elif is_included:
            exp_str = "\u2014" if not expiring_premium else _safe_currency(expiring_premium)
            rows.append([display_name, "Included in GL", exp_str, "Included", "\u2014", "\u2014"])
        else:
            exp_str = _safe_currency(expiring_premium) if expiring_premium else "\u2014"
            if expiring_premium:
                total_expiring += expiring_premium
            pending_coverages.append(display_name)
            rows.append([display_name, carrier_name, exp_str, "Pending", "\u2014", "\u2014"])

    # Total row
    total_change = total_proposed - total_expiring if total_expiring > 0 else 0
    total_pct = (total_change / total_expiring * 100) if total_expiring > 0 else 0
    total_change_str = f"+${total_change:,.0f}" if total_change > 0 else f"-${abs(total_change):,.0f}" if total_change != 0 else "\u2014"
    total_pct_str = f"+{total_pct:.1f}%" if total_change > 0 else f"{total_pct:.1f}%" if total_change != 0 else "\u2014"

    rows.append([
        "TOTAL", "",
        _safe_currency(total_expiring) if total_expiring else "\u2014",
        _safe_currency(total_proposed) if total_proposed else "\u2014",
        total_change_str, total_pct_str,
    ])

    return rows, total_change, total_pct, pending_coverages


def build_premium_comparison_internal(by_coverage, parsed_policies):
    """Build internal premium comparison with Commission and Revenue columns."""
    rows = []
    total_expiring = 0
    total_proposed = 0
    total_commission_revenue = 0
    pending_coverages = []

    sorted_coverages = sorted(by_coverage.keys(), key=_coverage_sort_key)

    for ct in sorted_coverages:
        policies = by_coverage[ct]
        display_name = COVERAGE_SHORT.get(ct, ct)

        expiring = [p for p in policies if p["status"] == "Incumbent"]
        bound = [p for p in policies if p["status"] in ("Bound", "Proposed")]
        quoted = [p for p in policies if p["status"] == "Quoted"]

        expiring_premium = sum(p["premium_tx"] for p in expiring) if expiring else 0
        proposed_premium = sum(p["premium_tx"] for p in bound) if bound else 0

        # Get commission and revenue from bound/proposed carrier
        comm_str = "—"
        rev_str = "—"
        broker_str = "—"
        if bound:
            if len(bound) > 1:
                carrier_name = f"{len(bound)}-Carrier Placement"
            else:
                carrier_name = bound[0]["carrier"]
            if bound[0]["commission"]:
                comm_str = _safe_percent(bound[0]["commission"])
            total_rev = sum(p["revenue"] for p in bound if p["revenue"])
            if total_rev:
                rev_str = _safe_currency(total_rev)
                total_commission_revenue += total_rev
            broker_str = bound[0]["broker"]
        elif quoted:
            carrier_name = quoted[0]["carrier"]
            if quoted[0]["commission"]:
                comm_str = _safe_percent(quoted[0]["commission"])
            broker_str = quoted[0]["broker"]
        elif expiring:
            # For multi-carrier expiring (e.g., property tower), show combined name
            if len(expiring) > 1:
                carrier_name = f"{len(expiring)}-Carrier Tower"
                # Use broker from first carrier for tower
                broker_str = expiring[0]["broker"]
            else:
                carrier_name = expiring[0]["carrier"]
                broker_str = expiring[0]["broker"]
        else:
            carrier_name = policies[0]["carrier"] if policies else "—"

        is_included = False
        for p in policies:
            comments = str(p.get("comments", "")).lower()
            if "included" in comments or "inc " in comments[:10]:
                is_included = True
                break

        if proposed_premium > 0:
            exp_str = _safe_currency(expiring_premium) if expiring_premium else "—"
            prop_str = _safe_currency(proposed_premium)
            if expiring_premium > 0:
                change = proposed_premium - expiring_premium
                pct_change = (change / expiring_premium) * 100
                change_str = f"+${change:,.0f}" if change > 0 else f"-${abs(change):,.0f}"
                pct_str = f"+{pct_change:.1f}%" if change > 0 else f"{pct_change:.1f}%"
                total_expiring += expiring_premium
                total_proposed += proposed_premium
            else:
                change_str = "\u2014"
                pct_str = "\u2014"
                total_proposed += proposed_premium
            rows.append([display_name, carrier_name, exp_str, prop_str, change_str, pct_str, comm_str, rev_str, broker_str])
        elif is_included:
            exp_str = "\u2014" if not expiring_premium else _safe_currency(expiring_premium)
            rows.append([display_name, "Included in GL", exp_str, "Included", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014"])
        else:
            exp_str = _safe_currency(expiring_premium) if expiring_premium else "\u2014"
            if expiring_premium:
                total_expiring += expiring_premium
            pending_coverages.append(display_name)
            rows.append([display_name, carrier_name, exp_str, "Pending", "\u2014", "\u2014", comm_str, rev_str, broker_str])

    # Total row
    total_change = total_proposed - total_expiring if total_expiring > 0 else 0
    total_pct = (total_change / total_expiring * 100) if total_expiring > 0 else 0
    total_change_str = f"+${total_change:,.0f}" if total_change > 0 else f"-${abs(total_change):,.0f}" if total_change != 0 else "\u2014"
    total_pct_str = f"+{total_pct:.1f}%" if total_change > 0 else f"{total_pct:.1f}%" if total_change != 0 else "\u2014"

    rows.append([
        "TOTAL", "",
        _safe_currency(total_expiring) if total_expiring else "\u2014",
        _safe_currency(total_proposed) if total_proposed else "\u2014",
        total_change_str, total_pct_str, "\u2014",
        _safe_currency(total_commission_revenue) if total_commission_revenue else "\u2014", "\u2014",
    ])

    return rows, total_change, total_pct, pending_coverages


# ══════════════════════════════════════════════════════════════════════════
# MARKET ACTIVITY BUILDER
# ══════════════════════════════════════════════════════════════════════════

def build_market_activity(by_coverage):
    """Build market activity overview rows."""
    rows = []
    sorted_coverages = sorted(by_coverage.keys(), key=_coverage_sort_key)

    for ct in sorted_coverages:
        policies = by_coverage[ct]
        display_name = COVERAGE_SHORT.get(ct, ct)

        total_marketed = len(policies)
        quotes_received = len([p for p in policies if p["status"] in ("Quoted", "Bound", "Proposed")])
        proposed_count = len([p for p in policies if p["status"] in ("Bound", "Proposed")])

        # Determine overall status
        if any(p["status"] == "Bound" for p in policies):
            overall_status = "Bound"
        elif any(p["status"] == "Proposed" for p in policies):
            overall_status = "Proposed"
        elif any(p["status"] == "Quoted" for p in policies):
            overall_status = "Quoted"
        else:
            overall_status = "Pending"

        # Check if included in another coverage
        is_included = False
        for p in policies:
            comments = str(p.get("comments", "")).lower()
            if "included" in comments:
                is_included = True
                break

        if is_included:
            rows.append([display_name, "—", "—", f"Included in GL", overall_status])
        else:
            proposed_str = str(proposed_count) if proposed_count > 0 else "0"
            rows.append([display_name, str(total_marketed), str(quotes_received), proposed_str, overall_status])

    return rows


# ══════════════════════════════════════════════════════════════════════════
# COVERAGE SUMMARY AT A GLANCE
# ══════════════════════════════════════════════════════════════════════════

def build_coverage_summary(by_coverage):
    """Build coverage summary at a glance rows."""
    rows = []
    sorted_coverages = sorted(by_coverage.keys(), key=_coverage_sort_key)

    for ct in sorted_coverages:
        policies = by_coverage[ct]
        display_name = COVERAGE_SHORT.get(ct, ct)

        expiring_carriers = [p["carrier"] for p in policies if p["status"] == "Incumbent"]
        bound_carriers = [p["carrier"] for p in policies if p["status"] in ("Bound", "Proposed")]
        quoted_carriers = [p["carrier"] for p in policies if p["status"] == "Quoted"]
        pending_carriers = [p["carrier"] for p in policies if p["status"] in ("Market", "Submit", "Pending")]

        # Check if included
        is_included = any("included" in str(p.get("comments", "")).lower() for p in policies)

        exp_str = ", ".join(set(expiring_carriers)) if expiring_carriers else "—"
        if is_included:
            bound_str = "Included in GL"
        else:
            bound_str = ", ".join(set(bound_carriers)) if bound_carriers else "—"
        quoted_str = ", ".join(set(quoted_carriers)) if quoted_carriers else "—"
        pending_str = ", ".join(set(pending_carriers)) if pending_carriers else "—"

        rows.append([display_name, exp_str, bound_str, quoted_str, pending_str])

    return rows


# ══════════════════════════════════════════════════════════════════════════
# MAIN DOCUMENT GENERATOR
# ══════════════════════════════════════════════════════════════════════════

def generate_marketing_update_docx(
    opp_fields: dict,
    parsed_policies: list,
    by_coverage: dict,
    client_name: str,
    is_internal: bool = True,
) -> str:
    """
    Generate the Marketing Update DOCX document.

    Args:
        opp_fields: Opportunity record fields (or empty dict)
        parsed_policies: List of parsed policy dicts
        by_coverage: Policies grouped by coverage type
        client_name: Client/opportunity name
        is_internal: If True, include commission, revenue, rates, broker info

    Returns:
        Path to the generated DOCX file.
    """
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(1.2)
        section.bottom_margin = Inches(0.6)

    add_page_header(doc)

    # ── Determine client info ──
    opp_name = ""
    effective_date = ""
    corporate_name = ""
    if opp_fields:
        opp_name = opp_fields.get("Opportunity Name", "")
        effective_date_raw = opp_fields.get("Effective Date", "")
        if effective_date_raw:
            try:
                dt = datetime.strptime(str(effective_date_raw), "%Y-%m-%d")
                effective_date = dt.strftime("%m/%d/%Y")
            except (ValueError, TypeError):
                effective_date = str(effective_date_raw)
        corporate_name = opp_fields.get("Opportunity Corporate Name", "")
        if not corporate_name:
            cn = opp_fields.get("Corporate Name", "")
            if isinstance(cn, list):
                corporate_name = cn[0] if cn else ""
            else:
                corporate_name = str(cn) if cn else ""

    display_name = corporate_name or client_name
    version_label = "Internal" if is_internal else "Client"

    # ════════════════════════════════════════════════════════════════
    # PAGE 1: TITLE & PREMIUM COMPARISON
    # ════════════════════════════════════════════════════════════════

    add_section_header(doc, "Insurance Marketing Update")

    add_formatted_paragraph(doc, f"Prepared for: {display_name}",
                           size=13, color=CLASSIC_BLUE, bold=True, space_before=4, space_after=2)
    if effective_date:
        add_formatted_paragraph(doc, f"Effective Date: {effective_date}  |  Package Program",
                               size=11, color=ELECTRIC_BLUE, bold=False, space_before=0, space_after=2)
    add_formatted_paragraph(doc, f"Report Date: {datetime.now().strftime('%m/%d/%Y')}",
                           size=10, color=CHARCOAL, bold=False, space_before=0, space_after=4)

    add_divider(doc)

    # ── Premium Comparison ──
    add_subsection_header(doc, "Premium Comparison — Expiring vs. Proposed")

    add_callout_box(doc, "Premiums shown include applicable taxes and fees. "
                        "TRIA/Terrorism premiums are not included in totals.")

    add_formatted_paragraph(doc, "", size=4, space_before=0, space_after=0)

    if is_internal:
        premium_headers = ["Coverage", "Carrier", "Expiring", "Proposed", "$ Change", "% Change", "Comm", "Revenue", "Broker"]
        premium_rows, total_change, total_pct, pending_coverages = build_premium_comparison_internal(by_coverage, parsed_policies)
        col_widths = [0.9, 1.1, 0.85, 0.85, 0.8, 0.65, 0.55, 0.7, 0.6]
    else:
        premium_headers = ["Coverage", "Carrier", "Expiring", "Proposed", "$ Change", "% Change"]
        premium_rows, total_change, total_pct, pending_coverages = build_premium_comparison(by_coverage, parsed_policies)
        col_widths = [1.2, 1.5, 1.1, 1.1, 1.1, 1.0]

    create_premium_summary_table(
        doc, premium_headers, premium_rows,
        col_widths=col_widths,
        highlight_last_row=True,
    )

    add_formatted_paragraph(doc, "", size=4, space_before=0, space_after=0)

    # Summary notes
    notes = []
    if total_change != 0:
        change_word = "decrease" if total_change < 0 else "increase"
        notes.append((
            f"Estimated total premium change: {'-' if total_change < 0 else '+'}${abs(total_change):,.2f} "
            f"({total_pct:+.1f}%) based on coverages quoted to date.",
            False, CLASSIC_BLUE
        ))
    if pending_coverages:
        notes.append((
            f"Note: {', '.join(pending_coverages)} premiums are pending and not included in proposed total.",
            False, CHARCOAL
        ))

    if notes:
        add_rich_callout_box(doc, notes, size=10, shading_hex=LIGHT_BLUE_HEX)

    # ════════════════════════════════════════════════════════════════
    # PAGE 2: MARKET ACTIVITY & KEY HIGHLIGHTS
    # ════════════════════════════════════════════════════════════════

    add_page_break(doc)
    add_subsection_header(doc, "Market Activity Overview")

    add_formatted_paragraph(doc,
        "The following summarizes our marketing efforts across all coverage lines. "
        "Our team has engaged multiple carriers to ensure competitive pricing and comprehensive coverage options.",
        size=10, color=CLASSIC_BLUE, space_before=2, space_after=8)

    activity_headers = ["Coverage", "Carriers Marketed", "Quotes Received", "Proposed", "Status"]
    activity_rows = build_market_activity(by_coverage)

    create_premium_summary_table(
        doc, activity_headers, activity_rows,
        col_widths=[1.5, 1.5, 1.3, 1.2, 1.5],
    )

    # Key Highlights section - auto-generated from data
    add_formatted_paragraph(doc, "", size=6, space_before=0, space_after=0)
    add_subsection_header(doc, "Key Highlights & Recommendations")

    highlights = _generate_highlights(by_coverage, parsed_policies)
    for title, description in highlights:
        hl_table = doc.add_table(rows=1, cols=1)
        hl_cell = hl_table.rows[0].cells[0]
        set_cell_shading(hl_cell, EGGSHELL_HEX)
        p_title = hl_cell.paragraphs[0]
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after = Pt(2)
        p_title.paragraph_format.left_indent = Inches(0.1)
        run_t = p_title.add_run(title)
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = ELECTRIC_BLUE
        run_t.font.bold = True
        run_t.font.name = "Calibri"
        p_desc = hl_cell.add_paragraph()
        p_desc.paragraph_format.space_before = Pt(2)
        p_desc.paragraph_format.space_after = Pt(6)
        p_desc.paragraph_format.left_indent = Inches(0.1)
        run_d = p_desc.add_run(description)
        run_d.font.size = Pt(9)
        run_d.font.color.rgb = CLASSIC_BLUE
        run_d.font.name = "Calibri"
        add_formatted_paragraph(doc, "", size=3, space_before=0, space_after=0)

    # ════════════════════════════════════════════════════════════════
    # CARRIER COMPARISON TABLES
    # ════════════════════════════════════════════════════════════════

    add_page_break(doc)
    add_section_header(doc, "Carrier Comparisons")

    color_legend = "Side-by-side carrier comparisons for each coverage line. "
    color_legend += "Expiring carriers shown in gray, proposed/bound in green, quoted in white, and pending in yellow."
    add_callout_box(doc, color_legend)

    add_formatted_paragraph(doc, "", size=6, space_before=0, space_after=0)

    sorted_coverages = sorted(by_coverage.keys(), key=_coverage_sort_key)
    first_on_page = True

    for ct in sorted_coverages:
        policies = by_coverage[ct]
        if not policies:
            continue

        display_title = COVERAGE_DISPLAY_NAMES.get(ct, ct)
        short_title = COVERAGE_SHORT.get(ct, ct)

        # Build carrier data for comparison table
        builder_func, metrics_func = COVERAGE_BUILDERS.get(ct, (_build_generic_carrier, _get_generic_metrics))

        carriers_data = []

        # For Property: use tower view if multiple incumbents
        if ct == "Property":
            tower = _build_property_tower(policies, is_internal=is_internal)
            if tower:
                carriers_data.append(tower)
                # Add non-incumbent carriers individually
                for p in policies:
                    if p["status"] != "Incumbent":
                        carrier_dict = builder_func(p, is_internal=is_internal)
                        carriers_data.append(carrier_dict)
            else:
                for p in policies:
                    carrier_dict = builder_func(p, is_internal=is_internal)
                    carriers_data.append(carrier_dict)
        # For Umbrella: use tower view if multiple incumbents (client version)
        elif ct == "Umbrella":
            tower = _build_umbrella_tower(policies, is_internal=is_internal)
            if tower and not is_internal:
                # Client version: show combined tower
                carriers_data.append(tower)
                for p in policies:
                    if p["status"] != "Incumbent":
                        carrier_dict = builder_func(p, is_internal=is_internal)
                        carriers_data.append(carrier_dict)
            else:
                # Internal version: show individual carriers for full detail
                for p in policies:
                    carrier_dict = builder_func(p, is_internal=is_internal)
                    carriers_data.append(carrier_dict)
        else:
            for p in policies:
                carrier_dict = builder_func(p, is_internal=is_internal)
                carriers_data.append(carrier_dict)

        if not carriers_data:
            continue

        # Separate declined carriers from the main table
        active_carriers = [c for c in carriers_data if c["status"] not in ("Declined", "Blocked", "Lost")]
        declined_carriers = [c for c in carriers_data if c["status"] in ("Declined", "Blocked", "Lost")]

        # Use active carriers for the main table (or all if none are active)
        table_carriers = active_carriers if active_carriers else carriers_data
        metrics = metrics_func(table_carriers, is_internal=is_internal)

        # Page break management - avoid too many tables on one page
        if not first_on_page and len(table_carriers) > 2:
            add_page_break(doc)
            first_on_page = True

        add_subsection_header(doc, display_title)
        create_carrier_comparison_table(doc, short_title, metrics, table_carriers)

        # Show declined carriers as notes below the table
        if declined_carriers:
            declined_names = []
            for dc in declined_carriers:
                name = dc["name"]
                notes = dc.get("notes", "")
                broker = dc.get("values", {}).get("Broker", "")
                parts = [name]
                if broker:
                    parts.append(f"via {broker}")
                if notes and notes != "Declined to quote":
                    # Truncate long notes
                    short_notes = notes[:80] + "..." if len(notes) > 80 else notes
                    parts.append(f"\u2014 {short_notes}")
                declined_names.append(" ".join(parts))
            declined_text = f"Declined: {'; '.join(declined_names)}"
            p_declined = doc.add_paragraph()
            p_declined.paragraph_format.space_before = Pt(2)
            p_declined.paragraph_format.space_after = Pt(4)
            p_declined.paragraph_format.left_indent = Inches(0.1)
            run_d = p_declined.add_run(declined_text)
            run_d.font.size = Pt(8)
            run_d.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run_d.font.italic = True
            run_d.font.name = "Calibri"

        add_formatted_paragraph(doc, "", size=8, space_before=0, space_after=0)
        first_on_page = False

    # Internal detail pages (Property Tower Detail, Umbrella Tower Detail, GL Detail)
    if is_internal:
        _add_internal_detail_pages(doc, by_coverage, parsed_policies)

    # ════════════════════════════════════════════════════════════════
    # NEXT STEPS & COVERAGE SUMMARY
    # ════════════════════════════════════════════════════════════════

    add_page_break(doc)
    add_subsection_header(doc, "Next Steps")

    next_steps = _generate_next_steps(by_coverage, parsed_policies)
    for i, step in enumerate(next_steps):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(14)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        run_num = p.add_run(f"{i+1}.  ")
        run_num.font.size = Pt(10)
        run_num.font.color.rgb = ELECTRIC_BLUE
        run_num.font.bold = True
        run_num.font.name = "Calibri"
        run_text = p.add_run(step)
        run_text.font.size = Pt(10)
        run_text.font.color.rgb = CLASSIC_BLUE
        run_text.font.name = "Calibri"

    # Coverage Summary at a Glance
    add_formatted_paragraph(doc, "", size=8, space_before=0, space_after=0)
    add_subsection_header(doc, "Coverage Summary at a Glance")

    summary_headers = ["Coverage", "Expiring", "Proposed / Bound", "Quoted", "Pending"]
    summary_rows = build_coverage_summary(by_coverage)

    create_premium_summary_table(
        doc, summary_headers, summary_rows,
        col_widths=[1.5, 1.5, 1.5, 1.5, 1.5],
    )

    # ── Contact & Disclaimer ──
    add_formatted_paragraph(doc, "", size=8, space_before=0, space_after=0)
    add_divider(doc)

    add_formatted_paragraph(doc, "Your HUB Hotel Franchise Team",
                           size=12, color=ELECTRIC_BLUE, bold=True, space_before=4, space_after=6)

    contact_table = doc.add_table(rows=3, cols=2)
    contact_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(contact_table)

    contacts = [
        ("Stefan Burkey", "Hotel Franchise Practice Leader  |  stefan.burkey@hubinternational.com"),
        ("Maureen Harvey", "Account Executive  |  maureen.harvey@hubinternational.com"),
        ("Sheena Callazo", "Claims Advocate  |  sheena.callazo@hubinternational.com"),
    ]

    for r_idx, (name, role) in enumerate(contacts):
        cell_name = contact_table.rows[r_idx].cells[0]
        cell_name.text = ""
        p_n = cell_name.paragraphs[0]
        p_n.paragraph_format.space_before = Pt(2)
        p_n.paragraph_format.space_after = Pt(2)
        run_n = p_n.add_run(name)
        run_n.font.size = Pt(9)
        run_n.font.color.rgb = CLASSIC_BLUE
        run_n.font.bold = True
        run_n.font.name = "Calibri"
        set_cell_width(cell_name, 1.8)
        remove_cell_borders(cell_name)

        cell_role = contact_table.rows[r_idx].cells[1]
        cell_role.text = ""
        p_r = cell_role.paragraphs[0]
        p_r.paragraph_format.space_before = Pt(2)
        p_r.paragraph_format.space_after = Pt(2)
        run_r = p_r.add_run(role)
        run_r.font.size = Pt(9)
        run_r.font.color.rgb = CHARCOAL
        run_r.font.name = "Calibri"
        set_cell_width(cell_role, 5.7)
        remove_cell_borders(cell_role)

    add_formatted_paragraph(doc, "", size=6, space_before=0, space_after=0)
    add_callout_box(doc, "This marketing update is for informational purposes only and does not constitute a binder of insurance. "
                        "Actual coverage terms, conditions, and exclusions are governed by the policies as issued. "
                        "Please review all policies carefully upon receipt.", size=8)

    # ── Save ──
    version_suffix = "Internal" if is_internal else "Client"
    safe_name = "".join(c for c in display_name if c.isalnum() or c in " _-").strip().replace(" ", "_")
    filename = f"Marketing_Update_{safe_name}_{version_suffix}.docx"
    output_path = os.path.join(tempfile.gettempdir(), filename)
    doc.save(output_path)
    logger.info(f"Marketing Update DOCX saved to {output_path}")
    return output_path


# ══════════════════════════════════════════════════════════════════════════
# AUTO-GENERATED HIGHLIGHTS & NEXT STEPS
# ══════════════════════════════════════════════════════════════════════════

def _generate_highlights(by_coverage, parsed_policies):
    """Auto-generate key highlights from the policy data."""
    highlights = []

    for ct in sorted(by_coverage.keys(), key=_coverage_sort_key):
        policies = by_coverage[ct]
        display_name = COVERAGE_SHORT.get(ct, ct)

        expiring = [p for p in policies if p["status"] == "Incumbent"]
        bound = [p for p in policies if p["status"] in ("Bound", "Proposed")]
        quoted = [p for p in policies if p["status"] == "Quoted"]

        if bound and expiring:
            exp_prem = sum(p["premium_tx"] for p in expiring)
            bound_prem = sum(p["premium_tx"] for p in bound)
            if exp_prem > 0 and bound_prem > 0:
                change = bound_prem - exp_prem
                pct = (change / exp_prem) * 100
                carrier = bound[0]["carrier"]
                if change < 0:
                    highlights.append((
                        f"{display_name} Premium Reduction",
                        f"{carrier} renewal at {_safe_currency(bound_prem)}, "
                        f"a {abs(pct):.1f}% decrease from expiring premium of {_safe_currency(exp_prem)}."
                    ))
                elif change > 0:
                    highlights.append((
                        f"{display_name} Premium Increase",
                        f"{carrier} renewal at {_safe_currency(bound_prem)}, "
                        f"a {pct:.1f}% increase from expiring premium of {_safe_currency(exp_prem)}."
                    ))

        # Note competitive quotes
        if quoted and len(quoted) >= 2:
            carriers = [q["carrier"] for q in quoted]
            highlights.append((
                f"{display_name} — Multiple Quotes",
                f"Received competitive quotes from {', '.join(set(carriers))}."
            ))

    # Check for pending coverages
    pending = []
    for ct in by_coverage:
        policies = by_coverage[ct]
        if not any(p["status"] in ("Bound", "Proposed", "Quoted") for p in policies):
            pending.append(COVERAGE_SHORT.get(ct, ct))

    if pending:
        highlights.append((
            "Pending Coverages",
            f"Awaiting quotes for: {', '.join(pending)}. Will update upon receipt."
        ))

    if not highlights:
        highlights.append((
            "Marketing in Progress",
            "Our team is actively marketing all coverage lines to ensure competitive pricing and comprehensive coverage."
        ))

    return highlights


def _generate_next_steps(by_coverage, parsed_policies):
    """Auto-generate next steps from the policy data."""
    steps = []

    # Check for pending/market status policies
    for ct in sorted(by_coverage.keys(), key=_coverage_sort_key):
        policies = by_coverage[ct]
        display_name = COVERAGE_SHORT.get(ct, ct)

        pending = [p for p in policies if p["status"] in ("Market", "Submit")]
        if pending:
            carriers = list(set(p["carrier"] for p in pending))
            steps.append(
                f"Awaiting {display_name} quotes from {', '.join(carriers)}."
            )

    # Check for quoted but not yet proposed
    for ct in sorted(by_coverage.keys(), key=_coverage_sort_key):
        policies = by_coverage[ct]
        display_name = COVERAGE_SHORT.get(ct, ct)

        quoted = [p for p in policies if p["status"] == "Quoted"]
        bound = [p for p in policies if p["status"] in ("Bound", "Proposed")]
        if quoted and not bound:
            steps.append(
                f"Finalizing {display_name} recommendation based on received quotes."
            )

    steps.append(
        "Full insurance proposal with all coverage sections, forms, and endorsements "
        "will be provided once all outstanding quotes are received."
    )

    return steps


# ══════════════════════════════════════════════════════════════════════════
# MAIN ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════

async def generate_marketing_update(client_name: str, is_internal: bool = True) -> str:
    """
    Main entry point: generate a Marketing Update DOCX for a client.

    Args:
        client_name: Client name to search for in Airtable.
        is_internal: If True, generate internal version with all financial details.

    Returns:
        Path to the generated DOCX file, or error message string.
    """
    logger.info(f"generate_marketing_update called for: '{client_name}' (internal={is_internal})")

    opp_fields, policies = await resolve_client_data(client_name)

    if not policies:
        return f"No policies found for '{client_name}'."

    parsed = parse_policies(policies)
    by_coverage = group_by_coverage(parsed)

    if not opp_fields:
        opp_fields = {}

    try:
        output_path = generate_marketing_update_docx(
            opp_fields=opp_fields,
            parsed_policies=parsed,
            by_coverage=by_coverage,
            client_name=client_name,
            is_internal=is_internal,
        )
        return output_path
    except Exception as e:
        logger.error(f"Error generating marketing update DOCX: {e}", exc_info=True)
        return f"Error generating document: {str(e)}"
